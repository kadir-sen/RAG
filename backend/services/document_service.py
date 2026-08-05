"""Document content serving for the right-panel viewer."""

import asyncio
import base64
import re
from pathlib import Path, PureWindowsPath

from backend.models.responses import DocContent, SchemaColumn
from backend.services.response_builder import _json_safe, clean_corrupted_date_string

_DATA_EXTENSIONS = {".xlsx", ".xls", ".csv"}


def _dtype_label(dtype) -> str:
    """pandas dtype → a friendly schema type for the viewer panel."""
    dt = str(dtype).lower()
    if "int" in dt:
        return "integer"
    if "float" in dt or "decimal" in dt:
        return "number"
    if "datetime" in dt or "date" in dt:
        return "date"
    if "bool" in dt:
        return "boolean"
    return "text"


def _build_schema(df, col_jargon: dict | None = None) -> list:
    """Per-column schema (name, dtype, jargon meaning) for the viewer panel."""
    cj = col_jargon or {}
    out = []
    for c in df.columns:
        name = str(c)
        out.append(SchemaColumn(
            name=name,
            dtype=_dtype_label(df[c].dtype),
            meaning=cj.get(name, "") or cj.get(name.strip(), ""),
        ))
    return out


def _clean_table_rows(rows: list) -> list:
    """Strip stray trailing letters from corrupted ISO date strings in viewer rows
    (e.g. "2027-03-23T00:00:00A" → "2027-03-23T00:00:00"). Shares the cleaner with
    the SQL artifact path so the same defect is fixed in both surfaces."""
    cleaned = []
    for row in rows:
        cleaned.append({
            str(k): _json_safe(
                clean_corrupted_date_string(v) if isinstance(v, str) else v
            )
            for k, v in row.items()
        })
    return cleaned

# Pattern to strip deduplication suffixes: "name_3.ext" -> "name.ext"
_DEDUP_SUFFIX_RE = re.compile(r'_\d+(\.[^.]+)$')

# Candidate roots searched by file_name when the stored file_path is stale
# (typical with vectors indexed on a different host: Windows or container
# paths that don't exist on the current disk).
_PROJECT_ROOT = Path(__file__).resolve().parents[2]
_DATA_FALLBACK_ROOTS = (
    _PROJECT_ROOT / "data" / "documents",
    _PROJECT_ROOT / "data" / "emails",
    _PROJECT_ROOT / "data" / "tables",
    # Bulk corpus PDFs synced to the server disk (e.g. the Edinburgh Tram set).
    # Kept in their own dir so a 7k-file corpus doesn't mix with demo uploads;
    # _resolve_path's direct `root / name` lookup stays O(1) regardless of size.
    _PROJECT_ROOT / "data" / "edinburgh_pdfs",
    _PROJECT_ROOT / "data",
)

# Shared, read-only corpora which may be referenced by more than one project.
# A file in one of these roots is NEVER resolved by name alone for a project-
# aware request: the caller must first prove membership through a project-
# scoped registry, catalog, RAG registry or chunk row.
_SHARED_CORPUS_ROOTS = (
    _PROJECT_ROOT / "data" / "documents",
    _PROJECT_ROOT / "data" / "emails",
    _PROJECT_ROOT / "data" / "tables",
    _PROJECT_ROOT / "data" / "edinburgh_pdfs",
    _PROJECT_ROOT / "data" / "edinburgh_tram",
)


class DocumentService:

    async def get_content(self, doc_id: str, anchor: str = "",
                          file_name: str = "", project_id: str = "") -> DocContent:
        return await asyncio.to_thread(
            self._get_content_sync, doc_id, anchor, file_name, project_id)

    def _is_data_file(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() in _DATA_EXTENSIONS

    @staticmethod
    def _resolve_path(file_path: str) -> str:
        """Return an existing file path, with progressively forgiving fallbacks.

        1. exact path on disk
        2. dedup-suffix removal ('letter-_3.docx' -> 'letter-.docx') in the
           same directory
        3. search by file_name across the project's data subdirs — this rescues
           vectors whose metadata was indexed on a different host (e.g.
           ``C:\\projects\\ML_project\\data\\…`` or ``/app/data/…``) but whose
           files now live in the local ``data/`` tree.

        Returns the original path when nothing matches so the caller can surface
        a meaningful error.
        """
        if not file_path:
            return file_path
        p = Path(file_path)
        if p.exists():
            return file_path

        # Extract the basename robustly: backslashes in a path string mean we
        # are looking at a Windows-style path that POSIX ``Path`` won't split.
        if "\\" in file_path and "/" not in file_path:
            base_name = PureWindowsPath(file_path).name
        else:
            base_name = p.name

        # 2) Strip dedup suffix in the same dir
        alt_name = _DEDUP_SUFFIX_RE.sub(r'\1', base_name)
        if alt_name != base_name and p.parent != Path(file_path):
            alt_path = p.parent / alt_name
            if alt_path.exists():
                return str(alt_path)

        # 3) Search the local data tree by file name. Try both the original
        # name and the dedup-stripped variant.
        candidates = {base_name}
        if alt_name != base_name:
            candidates.add(alt_name)
        for name in candidates:
            for root in _DATA_FALLBACK_ROOTS:
                hit = root / name
                if hit.exists():
                    return str(hit)
        # Deep search as a last resort (only the configured data subdirs, so
        # this stays O(small)).
        for name in candidates:
            for root in _DATA_FALLBACK_ROOTS:
                if not root.exists():
                    continue
                try:
                    for found in root.rglob(name):
                        if found.is_file():
                            return str(found)
                except OSError:
                    continue
        return file_path  # return original — caller will handle the error

    @staticmethod
    def _path_name(file_path: str) -> str:
        """Extract a basename from POSIX or Windows paths."""
        value = (file_path or "").strip()
        if not value:
            return ""
        if "\\" in value:
            return PureWindowsPath(value).name
        return Path(value).name

    @staticmethod
    def _project_file_roots(project_id: str) -> tuple[Path, ...]:
        """Return roots private to one project without accepting path traversal."""
        roots = []
        for base in (
            _PROJECT_ROOT / "data" / "projects",
            _PROJECT_ROOT / "storage" / "projects",
        ):
            try:
                resolved_base = base.resolve()
                candidate = (base / project_id).resolve()
                if candidate.is_relative_to(resolved_base):
                    roots.append(candidate)
            except (OSError, ValueError):
                continue
        return tuple(roots)

    @classmethod
    def _resolve_scoped_path(
        cls,
        file_path: str,
        file_name: str,
        project_id: str,
        *,
        allow_shared: bool,
    ) -> str:
        """Resolve a membership-proven source only inside approved roots.

        Production catalog/vector metadata can contain paths from the machine
        that performed ingestion.  We therefore recover by basename, but keep
        the search constrained to the active project's private roots and, only
        after project membership has been established, the shared corpus roots.
        """
        roots = list(cls._project_file_roots(project_id))
        if allow_shared:
            roots.extend(Path(root).resolve() for root in _SHARED_CORPUS_ROOTS)

        approved = []
        for root in roots:
            try:
                approved.append(Path(root).resolve())
            except (OSError, ValueError):
                continue

        raw = (file_path or "").strip()
        if raw:
            try:
                exact = Path(raw).resolve()
                if exact.is_file() and any(exact.is_relative_to(root) for root in approved):
                    return str(exact)
            except (OSError, ValueError):
                pass

        names = []
        for value in (file_name, file_path):
            name = cls._path_name(value)
            if name and name not in names:
                names.append(name)
            alt = _DEDUP_SUFFIX_RE.sub(r"\1", name) if name else ""
            if alt and alt not in names:
                names.append(alt)

        for name in names:
            for root in approved:
                try:
                    direct = (root / name).resolve()
                    if direct.is_file() and direct.is_relative_to(root):
                        return str(direct)
                except (OSError, ValueError):
                    continue

        for name in names:
            for root in approved:
                if not root.exists():
                    continue
                try:
                    for found in root.rglob(name):
                        resolved = found.resolve()
                        if resolved.is_file() and resolved.is_relative_to(root):
                            return str(resolved)
                except OSError:
                    continue
        return ""

    def _get_content_sync(self, doc_id: str, anchor: str,
                          file_name: str = "", project_id: str = "") -> DocContent:
        # Guard: nothing to go on at all. A file name alone is enough, though —
        # see the fallback at the end of this chain.
        if (not doc_id or not doc_id.strip()) and not (file_name or "").strip():
            return DocContent(type="text", error="No document ID provided")
        doc_id = (doc_id or "").strip()

        # Project-aware requests never enter the legacy global fallback chain:
        # resolving a guessed filename against the whole data tree would bypass
        # the membership boundary enforced by the API.
        if project_id:
            return self._get_project_content(doc_id, anchor, file_name, project_id)

        # Try data tables (Excel viewer) first — match by doc_id from file_paths
        try:
            from src.data_analyzer_sql import get_data_analyzer
            from src.document_rag import generate_doc_id
            analyzer = get_data_analyzer()
            for table_name, file_path in analyzer.file_paths.items():
                if generate_doc_id(file_path) == doc_id:
                    return self._serve_table_preview(table_name, analyzer)
        except Exception:
            pass

        # Then match by catalog source file → serve the EXTRACTED table (proper
        # detected headers + learned schema/jargon), not the raw sheet. file_paths
        # is keyed by parquet path, so a source-file doc_id (registry/sidebar) only
        # resolves here. Falls through to the raw-file path for 0-table entries.
        try:
            from src.catalog import get_catalog
            from src.data_analyzer_sql import get_data_analyzer
            from src.document_rag import generate_doc_id
            analyzer = get_data_analyzer()
            for entry in get_catalog().entries.values():
                if entry.source_type not in ("excel", "csv") or not entry.tables:
                    continue
                if generate_doc_id(entry.source_file) == doc_id:
                    tname = entry.tables[0].table_name
                    if tname in analyzer.tables:
                        return self._serve_table_preview(tname, analyzer)
        except Exception:
            pass

        # Try RAG file registry (match by doc_id hash OR by file_name)
        try:
            from src.document_rag import get_document_rag, generate_doc_id
            rag = get_document_rag()
            for fname, info in rag.file_registry.items():
                if (info.get("project_id", "") or "") != project_id:
                    continue
                stored_doc_id = info.get("doc_id", "")
                # Match by: exact doc_id, file_name, or MD5 hash of file_name
                import hashlib
                fname_hash = hashlib.md5(fname.encode()).hexdigest()[:16]
                if doc_id in (fname, stored_doc_id, fname_hash):
                    file_path = self._resolve_path(info.get("file_path", ""))
                    return self._serve_by_extension(file_path, anchor)
        except Exception:
            pass

        # Fallback: DocumentRegistry (JSON-backed, survives restarts)
        try:
            from src.document_registry import get_document_registry
            registry = get_document_registry()
            rec = registry.get(doc_id)
            # Also try matching by file_name hash if direct lookup fails
            if not rec:
                import hashlib
                for r in registry.get_all():
                    fname_hash = hashlib.md5(r.file_name.encode()).hexdigest()[:16]
                    if doc_id in (r.file_name, fname_hash):
                        rec = r
                        break
            if rec and rec.file_path:
                file_path = self._resolve_path(rec.file_path)
                return self._serve_by_extension(file_path, anchor)
        except Exception:
            pass

        # Before the text fallback: if a real file with this name is on disk, serve
        # it as a proper page image (PDF) / native preview. Bulk vectors-only docs
        # carry doc_id == file_name (see the library merge), so once their PDFs are
        # synced to the data dir this resolves them to the actual page render instead
        # of plain text. Hash-style doc_ids of older docs won't match a file → skip.
        try:
            resolved = self._resolve_path(doc_id)
            if resolved and Path(resolved).is_file():
                return self._serve_by_extension(resolved, anchor)
        except Exception:
            pass

        # Fallback: serve the chunk text from the chunk_store. This rescues the
        # right-panel viewer for corpora ingested vectors-only (no registry entry,
        # no PDF on disk) — the cited excerpt still opens, just as text not a page
        # image. Matches the source's doc_id (or file_name) against the mirrored
        # chunks; prefers the anchored page, else the whole document's text.
        try:
            from src.chunk_store import get_chunk_store
            con = get_chunk_store().connection()
            rows = con.execute(
                "SELECT file_name, page_number, text FROM chunks "
                "WHERE doc_id = ? OR file_name = ? OR file_name = ? "
                "ORDER BY page_number",
                [doc_id, doc_id, f"{doc_id}.pdf"],
            ).fetchall()
            if rows:
                page = self._parse_anchor_page(anchor)
                fname = rows[0][0]
                # Bulk-ingested corpora carry doc_id == md5(file_path) hash (see
                # generate_doc_id), so the disk probe above (keyed on doc_id) can
                # never match their PDFs. Now that the chunk row gives us the real
                # file_name, resolve THAT on disk: if the PDF is present, serve the
                # actual page image instead of falling back to raw chunk/OCR text.
                if fname:
                    resolved = self._resolve_path(fname)
                    if resolved and Path(resolved).is_file():
                        return self._serve_by_extension(resolved, anchor)
                total = max(int(r[1] or 1) for r in rows)
                page_rows = [r for r in rows if int(r[1] or 1) == page] or rows
                text = "\n\n".join((r[2] or "") for r in page_rows)[:8000]
                return DocContent(type="text", file_name=fname, page=page,
                                  total_pages=total, text=text)
        except Exception:
            pass

        # Last resort: the file name the caller already had.
        #
        # doc_id for older documents is generate_doc_id() — md5 of the file
        # *path* at ingest time. That path is a fingerprint of a moment: re-ingest
        # the corpus, move the data directory, or change host layout and every
        # id minted before it becomes unresolvable, while the documents sit
        # untouched on disk. Measured on production: 398 of 932 stored citations
        # (43%) carry such an id, 20 of 20 sampled failed to open by id, and 20
        # of 20 opened by name. Every citation carries doc_name, so a viewer that
        # gives up here is discarding the answer it was handed.
        #
        # Kept last on purpose: the id chain above is more specific (it can serve
        # an extracted table rather than a raw sheet), so this only runs once
        # that has genuinely failed.
        name = (file_name or "").strip()
        if name:
            try:
                resolved = self._resolve_path(name)
                if resolved and Path(resolved).is_file():
                    return self._serve_by_extension(resolved, anchor)
            except Exception:
                pass
            try:
                from src.chunk_store import get_chunk_store
                con = get_chunk_store().connection()
                rows = con.execute(
                    "SELECT file_name, page_number, text FROM chunks "
                    "WHERE file_name = ? ORDER BY page_number", [name],
                ).fetchall()
                if rows:
                    page = self._parse_anchor_page(anchor)
                    total = max(int(r[1] or 1) for r in rows)
                    page_rows = [r for r in rows if int(r[1] or 1) == page] or rows
                    text = "\n\n".join((r[2] or "") for r in page_rows)[:8000]
                    return DocContent(type="text", file_name=rows[0][0], page=page,
                                      total_pages=total, text=text)
            except Exception:
                pass

        return DocContent(type="text", error="Document not found")

    def _get_project_content(self, doc_id: str, anchor: str, file_name: str,
                             project_id: str) -> DocContent:
        import hashlib

        from src.document_rag import generate_doc_id
        from src.document_registry import get_document_registry

        registry = get_document_registry()
        records = registry.get_all(project_id=project_id)
        rec = next((r for r in records if doc_id in (
            r.doc_id,
            r.file_name,
            hashlib.md5(r.file_name.encode()).hexdigest()[:16],
        )), None)
        if not rec and file_name:
            rec = next((r for r in records if r.file_name == file_name), None)
        if rec and rec.file_path:
            resolved = self._resolve_scoped_path(
                rec.file_path, rec.file_name, project_id, allow_shared=True,
            )
            if resolved:
                return self._serve_by_extension(resolved, anchor)

        # Extracted spreadsheets live in the catalog rather than the document
        # registry.  Their source paths frequently point to the ingest machine,
        # so match within the active project first and then resolve the basename
        # against the managed corpus on this host.
        try:
            from src.catalog import get_catalog
            from src.data_analyzer_sql import get_data_analyzer

            requested_name = self._path_name(file_name or doc_id)
            for entry in get_catalog().entries.values():
                if entry.project_id != project_id or entry.source_type not in ("excel", "csv"):
                    continue
                source_name = self._path_name(entry.source_file)
                entry_id = generate_doc_id(entry.source_file)
                if doc_id not in (entry_id, source_name) and requested_name != source_name:
                    continue

                if entry.tables:
                    try:
                        analyzer = get_data_analyzer()
                        requested_sheet, row_from, row_to = self._parse_table_anchor(anchor)
                        selected_table = next((
                            table for table in entry.tables
                            if requested_sheet and str(table.sheet_name or table.table_name) == requested_sheet
                        ), entry.tables[0])
                        table_name = selected_table.table_name
                        if table_name in analyzer.tables:
                            preview = self._serve_table_preview(
                                table_name, analyzer, row_from=row_from, row_to=row_to,
                            )
                            if not preview.error:
                                return preview
                    except Exception:
                        # The raw source remains a valid viewer fallback when a
                        # parquet/view is temporarily unavailable after deploy.
                        pass

                resolved = self._resolve_scoped_path(
                    entry.source_file, source_name, project_id, allow_shared=True,
                )
                if resolved:
                    return self._serve_by_extension(resolved, anchor)
                break
        except Exception:
            pass

        # The in-memory RAG registry is another project-scoped membership
        # source, especially for documents uploaded since the latest restart.
        try:
            from src.document_rag import get_document_rag

            requested_name = self._path_name(file_name or doc_id)
            for fname, info in get_document_rag().file_registry.items():
                if (info.get("project_id") or "") != project_id:
                    continue
                stored_name = self._path_name(fname)
                stored_id = info.get("doc_id", "")
                fname_hash = hashlib.md5(stored_name.encode()).hexdigest()[:16]
                if doc_id not in (fname, stored_name, stored_id, fname_hash) and requested_name != stored_name:
                    continue
                resolved = self._resolve_scoped_path(
                    info.get("file_path", ""), stored_name, project_id,
                    allow_shared=True,
                )
                if resolved:
                    return self._serve_by_extension(resolved, anchor)
                break
        except Exception:
            pass

        try:
            from src.chunk_store import get_chunk_store
            con = get_chunk_store().connection()
            rows = con.execute(
                "SELECT file_name,page_number,text FROM chunks "
                "WHERE project_id=? AND (doc_id=? OR file_name=? OR file_name=?) "
                "ORDER BY page_number",
                [project_id, doc_id, file_name or doc_id, doc_id],
            ).fetchall()
            if rows:
                page = self._parse_anchor_page(anchor)
                total = max(int(r[1] or 1) for r in rows)
                source_name = self._path_name(rows[0][0])
                resolved = self._resolve_scoped_path(
                    "", source_name, project_id, allow_shared=True,
                )
                if resolved:
                    return self._serve_by_extension(resolved, anchor)
                page_rows = [r for r in rows if int(r[1] or 1) == page] or rows
                return DocContent(
                    type="text", file_name=source_name, page=page, total_pages=total,
                    text="\n\n".join((r[2] or "") for r in page_rows)[:8000],
                )
        except Exception:
            pass
        return DocContent(type="text", error="Document not found")

    def _serve_by_extension(self, file_path: str, anchor: str = "") -> DocContent:
        """Route to the right renderer based on file extension."""
        page = self._parse_anchor_page(anchor)
        lower = file_path.lower()
        if lower.endswith(".pdf"):
            return self._serve_pdf_page(file_path, page)
        elif self._is_data_file(file_path):
            return self._serve_excel_file(file_path, anchor)
        else:
            return self._serve_text_content(file_path)

    def _parse_anchor_page(self, anchor: str) -> int:
        if anchor.startswith("page_"):
            try:
                return int(anchor.replace("page_", ""))
            except ValueError:
                pass
        return 1

    @staticmethod
    def _parse_table_anchor(anchor: str) -> tuple[str, int | None, int | None]:
        match = re.fullmatch(r"sheet_(.+)_rows_(\d+)_(\d+)", anchor or "")
        if not match:
            return "", None, None
        row_from = max(2, int(match.group(2)))
        row_to = max(row_from, int(match.group(3)))
        return match.group(1), row_from, row_to

    def _serve_pdf_page(self, file_path: str, page: int) -> DocContent:
        try:
            import fitz
            doc = fitz.open(file_path)
            if page < 1 or page > len(doc):
                page = 1
            pdf_page = doc[page - 1]

            pix = pdf_page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            text = pdf_page.get_text()
            total = len(doc)
            doc.close()

            return DocContent(
                type="pdf",
                file_name=Path(file_path).name,
                page=page,
                total_pages=total,
                image_base64=base64.b64encode(img_bytes).decode(),
                text=text,
            )
        except Exception as e:
            return DocContent(type="pdf", error=str(e))

    def _serve_text_content(self, file_path: str) -> DocContent:
        try:
            fp = Path(file_path)
            ext = fp.suffix.lower()
            # Handle .msg (Outlook) files — binary format, need extract_msg
            if ext == ".msg":
                return self._serve_msg_content(file_path)
            # Handle .docx files — ZIP-based, need python-docx
            if ext == ".docx":
                return self._serve_docx_content(file_path)
            # Handle .doc files — legacy binary
            if ext == ".doc":
                return DocContent(type="text", file_name=fp.name,
                                  text="(Legacy .doc format — preview not available. Download to view.)")
            text = fp.read_text(encoding="utf-8", errors="replace")[:5000]
            return DocContent(
                type="text",
                file_name=fp.name,
                text=text,
            )
        except Exception as e:
            return DocContent(type="text", error=str(e))

    def _serve_docx_content(self, file_path: str) -> DocContent:
        """Parse .docx files and return readable text."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:5000]
            return DocContent(
                type="text",
                file_name=Path(file_path).name,
                text=text or "(Empty document)",
            )
        except Exception as e:
            return DocContent(type="text", error=f"Cannot parse docx: {e}")

    def _serve_msg_content(self, file_path: str) -> DocContent:
        """Parse .msg (Outlook email) files and return readable text."""
        try:
            import extract_msg
            msg = extract_msg.Message(file_path)
            parts = []
            if msg.subject:
                parts.append(f"Subject: {msg.subject}")
            if msg.sender:
                parts.append(f"From: {msg.sender}")
            if msg.to:
                parts.append(f"To: {msg.to}")
            if msg.date:
                parts.append(f"Date: {msg.date}")
            parts.append("")
            parts.append(msg.body or "(No body)")
            attachments = [att.longFilename or att.shortFilename for att in (msg.attachments or []) if att.longFilename or att.shortFilename]
            if attachments:
                parts.append(f"\nAttachments: {', '.join(attachments)}")
            msg.close()
            return DocContent(
                type="text",
                file_name=Path(file_path).name,
                text="\n".join(parts)[:5000],
            )
        except Exception as e:
            return DocContent(type="text", error=f"Cannot parse email: {e}")

    def _serve_excel_file(self, file_path: str, anchor: str = "") -> DocContent:
        """Read an Excel/CSV file directly with pandas and return as table."""
        try:
            import pandas as pd
            fp = Path(file_path)
            ext = fp.suffix.lower()
            requested_sheet, row_from, row_to = self._parse_table_anchor(anchor)
            row_offset = max(0, (row_from or 2) - 2)
            row_count = min(50, max(1, (row_to or (row_offset + 51)) - row_offset - 1))
            if ext == ".csv":
                df = pd.read_csv(
                    file_path, skiprows=range(1, row_offset + 1), nrows=row_count,
                )
            else:
                sheet = requested_sheet or 0
                df = pd.read_excel(
                    file_path, sheet_name=sheet,
                    skiprows=range(1, row_offset + 1), nrows=row_count,
                )
            # Get total row count without loading entire file
            if ext == ".csv":
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    total_rows = sum(1 for _ in f) - 1  # minus header
            else:
                df_full_len = len(pd.read_excel(
                    file_path, usecols=[0], sheet_name=requested_sheet or 0,
                ))
                total_rows = df_full_len
            col_jargon, description = self._catalog_schema_hints(fp.name)
            return DocContent(
                type="table",
                file_name=fp.name,
                columns=list(df.columns.astype(str)),
                rows=_clean_table_rows(df.fillna("").to_dict("records")),
                total_rows=max(total_rows, len(df)),
                schema_columns=_build_schema(df, col_jargon),
                description=description,
                sheet_name=requested_sheet,
                row_from=(row_offset + 2),
                row_to=(row_offset + len(df) + 1),
            )
        except Exception as e:
            return DocContent(type="table", error=str(e))

    def _serve_table_preview(
        self, table_name: str, analyzer, *, row_from: int | None = None,
        row_to: int | None = None,
    ) -> DocContent:
        try:
            offset = max(0, (row_from or 2) - 2)
            limit = min(50, max(1, (row_to or (offset + 51)) - offset - 1))
            df = analyzer.conn.execute(
                f'SELECT * FROM "{table_name}" LIMIT ? OFFSET ?', [limit, offset]
            ).fetchdf()
            info = analyzer.tables.get(table_name, {})
            display_name = info.get("file_name", table_name)
            col_jargon, cat_desc = self._catalog_schema_hints(display_name, table_name)
            return DocContent(
                type="table",
                file_name=display_name,
                columns=list(df.columns),
                rows=_clean_table_rows(df.to_dict("records")),
                total_rows=info.get("row_count", len(df)),
                schema_columns=_build_schema(df, col_jargon),
                description=info.get("description", "") or cat_desc,
                sheet_name=str(info.get("sheet_name", "") or ""),
                row_from=offset + 2,
                row_to=offset + len(df) + 1,
            )
        except Exception as e:
            return DocContent(type="table", error=str(e))

    @staticmethod
    def _catalog_schema_hints(file_name: str, table_name: str = "") -> tuple[dict, str]:
        """Best-effort per-column meaning (column_jargon) + table description from
        the catalog, matched by table name first then by source file name."""
        try:
            from src.catalog import get_catalog
            best = None
            for tm in get_catalog().get_all_tables():
                if table_name and tm.table_name == table_name:
                    best = tm
                    break
                if Path(tm.source_file).name == file_name and best is None:
                    best = tm
            if best:
                return (best.column_jargon or {}), (best.description or "")
        except Exception:
            pass
        return {}, ""
