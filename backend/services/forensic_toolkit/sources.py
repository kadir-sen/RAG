"""Project-scoped source catalogue and immutable forensic snapshots.

The normal COAir library and the programme store intentionally use different
pipelines.  This service is the single read boundary that joins them without
registering XER files in RAG or charging storage twice when a source is selected
for analysis.
"""

from __future__ import annotations

import hashlib
import os
import csv
import io
import shutil
from email import policy
from email.parser import BytesParser
from pathlib import Path
from typing import Any, Dict, Iterable, List

from src.config import STORAGE_DIR
from src.forensic_store import ForensicStore, get_forensic_store


_EVIDENCE_EXTENSIONS = {".pdf", ".doc", ".docx", ".txt", ".eml", ".msg",
                        ".xls", ".xlsx", ".csv"}


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _fallback_hash(*parts: object) -> str:
    return hashlib.sha256("|".join(str(value or "") for value in parts).encode()).hexdigest()


def _capabilities(extension: str, *, physical: bool, text_only: bool = False,
                  sheets: Iterable[str] = ()) -> List[str]:
    values = ["forensic_evidence"]
    if physical:
        values.append("original_download")
    if extension == ".xer":
        return ["programme_analysis", "inventory", "chain_of_custody"]
    if extension in {".pdf", ".doc", ".docx", ".txt"} or text_only:
        values.extend(["text_extraction", "tia_event_extraction", "narrative_evidence"])
    if extension in {".eml", ".msg"}:
        values.extend(["email_headers", "text_extraction", "tia_event_extraction"])
    if extension in {".xls", ".xlsx", ".csv"} or list(sheets):
        values.extend(["table_rows", "sheet_scope", "narrative_evidence"])
    if text_only:
        values.append("text_only")
    return list(dict.fromkeys(values))


class ForensicSourceService:
    def __init__(self, store: ForensicStore | None = None):
        self.store = store or get_forensic_store()

    def list_project_sources(self, project_id: str, *, include_internal: bool = False) -> List[Dict[str, Any]]:
        sources: Dict[str, Dict[str, Any]] = {}

        for programme in self.store.list_programmes(project_id, include_path=True):
            path = Path(programme["file_path"])
            item = {
                "source_id": programme["file_id"], "source_kind": "programme",
                "file_name": programme["name"], "extension": ".xer",
                "size_bytes": programme["size_bytes"], "content_hash": programme["sha256"],
                "status": "ready" if path.is_file() else "missing",
                "capabilities": _capabilities(".xer", physical=path.is_file()),
                "metadata": {"uploaded_by": programme.get("uploaded_by", "")},
            }
            if include_internal:
                item["_source_path"] = str(path)
            sources[item["source_id"]] = item

        try:
            from src.document_index import get_document_index
            indexed = {record.doc_id: record for record in
                       get_document_index().list_project(project_id)}
        except Exception:
            indexed = {}

        try:
            from src.document_registry import get_document_registry
            records = get_document_registry().get_all(project_id=project_id)
        except Exception:
            records = []
        known_names = set()
        for record in records:
            extension = str(record.extension or Path(record.file_name).suffix).casefold()
            if extension not in _EVIDENCE_EXTENSIONS:
                continue
            path = Path(record.file_path)
            index = indexed.get(record.doc_id)
            content_hash = (getattr(index, "content_hash", "") or
                            str(getattr(record, "file_hash", "")) or
                            _fallback_hash(project_id, record.doc_id, record.file_name,
                                           record.file_size_kb, record.completed_at))
            sheets = list(getattr(index, "sheet_names", []) or [])
            item = {
                "source_id": record.doc_id,
                "source_kind": "email" if extension in {".eml", ".msg"}
                    else "data" if extension in {".xls", ".xlsx", ".csv"}
                    else "document",
                "file_name": record.file_name, "extension": extension,
                "size_bytes": int(record.file_size_kb or 0) * 1024,
                "content_hash": content_hash,
                "status": "ready" if record.status == "completed" else record.status,
                "capabilities": _capabilities(extension, physical=path.is_file(), sheets=sheets),
                "metadata": {
                    "title": getattr(index, "title", "") if index else "",
                    "reference": getattr(index, "reference", "") if index else "",
                    "sheets": sheets,
                    "ocr_quality": getattr(index, "ocr_quality", "") if index else "",
                },
            }
            if include_internal:
                item["_source_path"] = str(path) if path.is_file() else ""
            sources[item["source_id"]] = item
            known_names.add(record.file_name.casefold())

        # Bulk projects may only have vector/chunk records. They remain useful
        # for evidence extraction even when an original PDF was never retained.
        try:
            from src.chunk_store import get_chunk_store
            rows = get_chunk_store().connection().execute(
                "SELECT doc_id,file_name,MAX(page_number) FROM chunks "
                "WHERE project_id=? AND file_name IS NOT NULL AND file_name<>'' "
                "GROUP BY doc_id,file_name ORDER BY file_name", [project_id],
            ).fetchall()
        except Exception:
            rows = []
        for doc_id, file_name, max_page in rows:
            source_id = str(doc_id or file_name)
            if source_id in sources or str(file_name).casefold() in known_names:
                continue
            extension = Path(str(file_name)).suffix.casefold()
            if extension not in _EVIDENCE_EXTENSIONS:
                extension = ".txt"
            index = indexed.get(source_id)
            item = {
                "source_id": source_id,
                "source_kind": "email" if extension in {".eml", ".msg"}
                    else "data" if extension in {".xls", ".xlsx", ".csv"}
                    else "document",
                "file_name": str(file_name), "extension": extension,
                "size_bytes": 0,
                "content_hash": (getattr(index, "content_hash", "") or
                                 _fallback_hash(project_id, source_id, file_name, max_page)),
                "status": "text_only",
                "capabilities": _capabilities(extension, physical=False, text_only=True,
                                                sheets=getattr(index, "sheet_names", []) if index else []),
                "metadata": {"pages": int(max_page or 0), "text_only": True,
                             "title": getattr(index, "title", "") if index else ""},
            }
            if include_internal:
                item["_source_path"] = ""
            sources[source_id] = item

        # Catalog-only spreadsheets are not necessarily represented in either
        # the registry or the chunk store.
        try:
            from src.catalog import get_catalog
            entries = list(get_catalog().entries.values())
        except Exception:
            entries = []
        for entry in entries:
            if (getattr(entry, "project_id", "") or "") != project_id:
                continue
            file_name = Path(entry.source_file).name
            if file_name.casefold() in known_names:
                continue
            extension = Path(file_name).suffix.casefold()
            if extension not in {".xls", ".xlsx", ".csv"}:
                continue
            source_id = "table:" + hashlib.sha256(
                f"{project_id}|{entry.source_file}".encode()).hexdigest()[:24]
            path = Path(entry.source_file)
            sheets = list(dict.fromkeys(str(table.sheet_name or table.table_name)
                                        for table in entry.tables))
            item = {
                "source_id": source_id, "source_kind": "data", "file_name": file_name,
                "extension": extension, "size_bytes": path.stat().st_size if path.is_file() else 0,
                "content_hash": str(entry.file_hash or _fallback_hash(source_id, sheets)),
                "status": "ready" if path.is_file() else "table_only",
                "capabilities": _capabilities(extension, physical=path.is_file(), sheets=sheets),
                "metadata": {"sheets": sheets, "tables": len(entry.tables)},
            }
            if include_internal:
                item["_source_path"] = str(path) if path.is_file() else ""
            sources[source_id] = item

        return sorted(sources.values(), key=lambda item: (item["source_kind"],
                                                           item["file_name"].casefold()))

    def resolve_selection(self, *, project_id: str, workspace_id: str,
                          selections: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        catalogue = {item["source_id"]: item for item in
                     self.list_project_sources(project_id, include_internal=True)}
        snapshot_root = (Path(STORAGE_DIR) / "projects" / project_id / "forensic" /
                         "snapshots" / workspace_id)
        snapshot_root.mkdir(parents=True, exist_ok=True)
        resolved: List[Dict[str, Any]] = []
        seen: set[str] = set()
        for requested in selections:
            source_id = str(requested.get("source_id") or "")
            if not source_id or source_id in seen or source_id not in catalogue:
                raise ValueError("forensic_source_selection_invalid")
            seen.add(source_id)
            source = dict(catalogue[source_id])
            if source["status"] in {"processing", "error", "missing"}:
                raise ValueError("forensic_source_not_ready")
            source_path = Path(source.pop("_source_path", ""))
            snapshot_path = ""
            if source_path.is_file():
                actual_hash = _sha256_file(source_path)
                target = snapshot_root / f"{actual_hash[:16]}-{Path(source['file_name']).name}"
                if not target.exists():
                    try:
                        os.link(source_path, target)
                    except OSError:
                        shutil.copy2(source_path, target)
                if _sha256_file(target) != actual_hash:
                    target.unlink(missing_ok=True)
                    raise ValueError("forensic_source_snapshot_hash_mismatch")
                source["content_hash"] = actual_hash
                source["size_bytes"] = target.stat().st_size
                snapshot_path = str(target)
            elif "text_only" in source["capabilities"]:
                text = self._chunk_text(project_id, source_id, source["file_name"])
                if not text.strip():
                    raise ValueError("forensic_source_text_unavailable")
                digest = hashlib.sha256(text.encode("utf-8")).hexdigest()
                target = snapshot_root / f"{digest[:16]}-{Path(source['file_name']).stem}.md"
                if not target.exists():
                    target.write_text(text, encoding="utf-8")
                source["content_hash"] = digest
                source["size_bytes"] = target.stat().st_size
                snapshot_path = str(target)
            source["snapshot_path"] = snapshot_path
            source["selected_scope"] = dict(requested.get("selected_scope") or {})
            source.pop("metadata", None)
            resolved.append(source)
        return resolved

    def evidence_documents(self, *, project_id: str, workspace_id: str,
                           source_ids: List[str] | None = None) -> List[tuple[str, str]]:
        """Read pinned evidence with stable page/sheet/row anchors.

        Only workspace snapshots are read, never a newly changed registry
        path.  This keeps AI proposals reproducible and ensures a selected
        source can still be explained after its live project record changes.
        """
        wanted = set(source_ids or [])
        documents: List[tuple[str, str]] = []
        for source in self.store.list_workspace_sources(project_id, workspace_id):
            if source["source_kind"] == "programme":
                continue
            if wanted and source["source_id"] not in wanted:
                continue
            path = Path(source.get("snapshot_path") or "")
            if not path.is_file():
                continue
            text = self._read_snapshot(path, source.get("extension") or path.suffix,
                                       source.get("selected_scope") or {})
            if text.strip():
                documents.append((source["file_name"], text))
        return documents

    @staticmethod
    def _read_snapshot(path: Path, extension: str, scope: Dict[str, Any]) -> str:
        extension = extension.casefold()
        raw = path.read_bytes()
        if extension == ".pdf":
            try:
                from pypdf import PdfReader
                pages = set(int(value) for value in scope.get("pages") or [])
                blocks = []
                for index, page in enumerate(PdfReader(io.BytesIO(raw)).pages, 1):
                    if pages and index not in pages:
                        continue
                    blocks.append(f"<!-- page:{index} -->\n{page.extract_text() or ''}")
                return "\n\n".join(blocks)
            except Exception:
                return ""
        if extension == ".docx":
            try:
                from docx import Document
                document = Document(io.BytesIO(raw))
                lines = [f"<!-- paragraph:{index} -->\n{paragraph.text}"
                         for index, paragraph in enumerate(document.paragraphs, 1)
                         if paragraph.text.strip()]
                for table_index, table in enumerate(document.tables, 1):
                    lines.append(f"<!-- table:{table_index} -->")
                    lines.extend(" | ".join(cell.text for cell in row.cells)
                                 for row in table.rows)
                return "\n".join(lines)
            except Exception:
                return ""
        if extension in {".xlsx", ".xls"}:
            try:
                if extension == ".xls":
                    import pandas as pd
                    selected_sheet = str(scope.get("sheet") or "") or None
                    frames = pd.read_excel(io.BytesIO(raw), sheet_name=selected_sheet,
                                           header=None, dtype=str)
                    if not isinstance(frames, dict):
                        frames = {selected_sheet or "Sheet1": frames}
                    row_from = max(1, int(scope.get("row_from") or 1))
                    row_to = int(scope.get("row_to") or 0) or None
                    lines = []
                    for sheet_name, frame in frames.items():
                        lines.append(f"<!-- sheet:{sheet_name} -->")
                        for offset, row in frame.iloc[row_from - 1:row_to].iterrows():
                            lines.append(f"<!-- row:{offset + 1} --> " + " | ".join(
                                "" if value is None else str(value) for value in row.tolist()))
                    return "\n".join(lines)
                from openpyxl import load_workbook
                workbook = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
                selected_sheet = str(scope.get("sheet") or "")
                sheets = [selected_sheet] if selected_sheet in workbook.sheetnames else workbook.sheetnames
                row_from = max(1, int(scope.get("row_from") or 1))
                row_to = int(scope.get("row_to") or 0) or None
                lines = []
                for sheet_name in sheets:
                    lines.append(f"<!-- sheet:{sheet_name} -->")
                    for row_number, row in enumerate(workbook[sheet_name].iter_rows(values_only=True), 1):
                        if row_number < row_from or (row_to and row_number > row_to):
                            continue
                        values = ["" if value is None else str(value) for value in row]
                        lines.append(f"<!-- row:{row_number} --> " + " | ".join(values))
                return "\n".join(lines)
            except Exception:
                return ""
        if extension == ".csv":
            decoded = raw.decode("utf-8-sig", errors="replace")
            row_from = max(1, int(scope.get("row_from") or 1))
            row_to = int(scope.get("row_to") or 0) or None
            return "\n".join(
                f"<!-- row:{index} --> " + " | ".join(row)
                for index, row in enumerate(csv.reader(io.StringIO(decoded)), 1)
                if index >= row_from and (row_to is None or index <= row_to)
            )
        if extension in {".eml", ".msg"}:
            try:
                message = BytesParser(policy=policy.default).parsebytes(raw)
                headers = "\n".join(f"{name}: {message.get(name, '')}"
                                    for name in ("From", "To", "Date", "Subject"))
                body = message.get_body(preferencelist=("plain",)) if message.is_multipart() else message
                content = body.get_content() if body else ""
                return f"<!-- email-headers -->\n{headers}\n\n<!-- email-body -->\n{content}"
            except Exception:
                return raw.decode("utf-8", errors="replace")
        return raw.decode("utf-8", errors="replace")

    @staticmethod
    def _chunk_text(project_id: str, doc_id: str, file_name: str) -> str:
        try:
            from src.chunk_store import get_chunk_store
            rows = get_chunk_store().connection().execute(
                "SELECT page_number,text FROM chunks WHERE project_id=? "
                "AND (doc_id=? OR file_name=?) ORDER BY page_number,chunk_id",
                [project_id, doc_id, file_name],
            ).fetchall()
        except Exception:
            rows = []
        blocks = []
        current_page = None
        for page, text in rows:
            if page != current_page:
                current_page = page
                blocks.append(f"\n\n<!-- page:{page or 1} -->\n")
            blocks.append(str(text or ""))
        return "\n".join(blocks)


__all__ = ["ForensicSourceService"]
