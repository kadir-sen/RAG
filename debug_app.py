"""
Ön İşleme Gözlem Servisi (Debug App)
Dosya işleme pipeline sonuçlarını gözlemlemek için ayrı Streamlit uygulaması.

Usage:
    streamlit run debug_app.py --server.port 8502
"""
import os
import sys
import json
import tempfile
from pathlib import Path
from datetime import datetime

import streamlit as st
import pandas as pd

# Ensure src is importable
sys.path.insert(0, str(Path(__file__).parent))

# ── Page Config ─────────────────────────────────────────────
st.set_page_config(
    page_title="Ön İşleme Gözlem Servisi",
    page_icon="🔍",
    layout="wide",
)

# ── Custom CSS ──────────────────────────────────────────────
st.markdown("""
<style>
    .stTabs [data-baseweb="tab-list"] { gap: 8px; }
    .stTabs [data-baseweb="tab"] {
        padding: 8px 16px;
        border-radius: 8px 8px 0 0;
    }
    .insight-box {
        background: #1e1e2e;
        border-radius: 10px;
        padding: 16px;
        margin: 8px 0;
        border-left: 4px solid #4285f4;
    }
    .anomaly-box {
        background: #2d1f1f;
        border-radius: 8px;
        padding: 12px;
        margin: 4px 0;
        border-left: 3px solid #ef4444;
    }
    .stat-card {
        background: #1a1a2e;
        border-radius: 8px;
        padding: 12px;
        text-align: center;
    }
    .notice-field {
        background: #1e2a1e;
        border-radius: 6px;
        padding: 8px 12px;
        margin: 4px 0;
    }
    .tag-chip {
        display: inline-block;
        padding: 2px 10px;
        border-radius: 12px;
        font-size: 0.85em;
        margin: 2px;
    }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════
# Helper Functions (defined before tabs so they are available)
# ══════════════════════════════════════════════════════════════

def _render_insight(insight: dict):
    """Render table insight in a nice format."""
    if not insight:
        st.info("Insight çıkarılamadı")
        return

    narrative = insight.get("narrative", "")
    if narrative:
        st.markdown(f'<div class="insight-box">📋 <b>{narrative}</b></div>',
                    unsafe_allow_html=True)

    stats = insight.get("stats", {})
    if stats:
        st.markdown("**📊 İstatistikler:**")
        stat_cols = st.columns(min(len(stats), 3))
        for i, (key, val) in enumerate(stats.items()):
            col_idx = i % min(len(stats), 3)
            with stat_cols[col_idx]:
                display_key = key.replace("_", " ").title()
                if isinstance(val, dict):
                    st.markdown(f"**{display_key}:**")
                    for k, v in list(val.items())[:5]:
                        st.markdown(f"- {k}: {v}")
                elif isinstance(val, float):
                    st.metric(display_key, f"{val:,.2f}")
                else:
                    st.metric(display_key, str(val))

    anomalies = insight.get("anomalies", [])
    if anomalies:
        st.markdown("**⚠️ Anomaliler:**")
        for a in anomalies:
            st.markdown(f'<div class="anomaly-box">⚠️ {a}</div>',
                        unsafe_allow_html=True)

    completeness = insight.get("completeness_score", 0)
    if completeness:
        st.progress(completeness, text=f"Veri Tamlığı: %{completeness*100:.0f}")

    jargon_map = insight.get("jargon_map", {})
    if jargon_map:
        st.markdown("**📖 Jargon:**")
        for abbr, meaning in jargon_map.items():
            st.markdown(f"- **{abbr}** = {meaning}")


def _extract_doc_text(file_path: str) -> dict:
    """Extract text from a document file, returns {page_num: text}."""
    ext = Path(file_path).suffix.lower()
    doc_text_by_page = {}

    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    if text.strip():
                        doc_text_by_page[i] = text
        except Exception:
            try:
                import fitz
                doc = fitz.open(file_path)
                for i, page in enumerate(doc, 1):
                    text = page.get_text()
                    if text.strip():
                        doc_text_by_page[i] = text
            except Exception:
                pass

    elif ext in [".docx", ".doc"]:
        try:
            from docx import Document
            doc = Document(file_path)
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if full_text:
                chunk_size = 2000
                for i in range(0, len(full_text), chunk_size):
                    doc_text_by_page[i // chunk_size + 1] = full_text[i:i + chunk_size]
        except Exception:
            pass

    elif ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            if text.strip():
                doc_text_by_page[1] = text
        except Exception:
            pass

    return doc_text_by_page


def _extract_and_display_notice(file_path: str, doc_text_by_page: dict):
    """Extract notice metadata and display it."""
    try:
        from src.notice_extractor import NoticeExtractor

        extractor = NoticeExtractor()
        doc_id = Path(file_path).stem
        notice = extractor.extract_notice(
            doc_id=doc_id,
            file_path=file_path,
            doc_text_by_page=doc_text_by_page,
            use_llm=False,
        )

        if notice:
            fields = {
                "Doküman Tipi": notice.doc_type or "—",
                "Tarih": notice.date or "—",
                "Gönderen": notice.sender or "—",
                "Gönderen Ünvan": getattr(notice, 'sender_title', '') or "—",
                "Gönderen Şirket": getattr(notice, 'sender_company', '') or "—",
                "Alıcı": notice.recipient or "—",
                "Proje": getattr(notice, 'project_name', '') or "—",
                "Sözleşme Ref": getattr(notice, 'contract_ref', '') or "—",
                "Yön": getattr(notice, 'direction', '') or "—",
            }
            field_df = pd.DataFrame([
                {"Alan": k, "Değer": v}
                for k, v in fields.items()
                if v != "—"
            ])
            st.dataframe(field_df, use_container_width=True, hide_index=True)

            refs = getattr(notice, 'ref_numbers', [])
            if refs:
                st.markdown(f"**📎 Referans Numaraları:** {', '.join(refs)}")

            actions = getattr(notice, 'actions', [])
            if actions:
                action_html = " ".join(
                    f'<span class="tag-chip" style="background:#166534;color:#fff;">{a}</span>'
                    for a in actions
                )
                st.markdown(f"**🎯 Aksiyonlar:** {action_html}", unsafe_allow_html=True)

            deadlines = getattr(notice, 'deadlines', [])
            if deadlines:
                st.markdown("**⏰ Deadline'lar:**")
                for dl in deadlines:
                    if isinstance(dl, dict):
                        st.markdown(f"- {dl.get('date', '?')}: {dl.get('context', '')}")
                    else:
                        st.markdown(f"- {dl}")

            jargon_found = getattr(notice, 'jargon_found', [])
            if jargon_found:
                st.markdown("**📖 Jargon:**")
                for j in jargon_found:
                    if isinstance(j, dict):
                        st.markdown(f"- **{j.get('abbreviation', '')}** = {j.get('meaning', '')}")

            evidence = getattr(notice, 'evidence_spans', [])
            if evidence:
                with st.expander(f"📍 Evidence ({len(evidence)} kaynak)"):
                    for ev in evidence:
                        if isinstance(ev, dict):
                            field_name = ev.get('field_name', '?')
                            page = ev.get('page', '?')
                            snippet = ev.get('snippet', '')[:200]
                            conf = ev.get('confidence', 0)
                            st.markdown(
                                f"**{field_name}** → Sayfa {page}: "
                                f"*\"{snippet}\"* ({conf:.2f})"
                            )
        else:
            st.info("Bu dokümandan notice çıkarılamadı")

    except Exception as e:
        st.warning(f"Notice extraction hatası: {e}")


def _review_and_display_categories(file_path: str, doc_text_by_page: dict):
    """Review document and display issue categories."""
    try:
        from src.document_reviewer import DocumentReviewer, IssueCategorizer

        full_text = "\n\n".join(
            doc_text_by_page[k] for k in sorted(doc_text_by_page.keys())
        )
        doc_id = Path(file_path).stem

        categorizer = IssueCategorizer()
        cat_result = categorizer.categorize(
            doc_id=doc_id,
            file_name=Path(file_path).name,
            full_text=full_text,
        )

        if cat_result:
            ISSUE_COLORS = {
                "delay": "#dc2626",
                "payment_dispute": "#d97706",
                "quality_concern": "#2563eb",
                "safety_issue": "#7c3aed",
                "scope_change": "#059669",
                "communication_gap": "#6366f1",
                "contractual_dispute": "#be185d",
            }

            cats = getattr(cat_result, 'categories', [])
            if cats:
                chips_html = " ".join(
                    f'<span class="tag-chip" style="background:{ISSUE_COLORS.get(c, "#666")};">'
                    f'{c}</span>'
                    for c in cats
                )
                st.markdown(f"**Kategoriler:** {chips_html}", unsafe_allow_html=True)

            score = getattr(cat_result, 'key_document_score', 0)
            is_key = getattr(cat_result, 'is_key_document', False)
            if is_key:
                st.success(f"⭐ Key Document Score: **{score:.2f}** (Yüksek önem)")
            else:
                st.info(f"Key Document Score: {score:.2f}")

            reasons = getattr(cat_result, 'key_reasons', [])
            if reasons:
                for r in reasons:
                    st.markdown(f"- {r}")

    except Exception as e:
        st.warning(f"Review hatası: {e}")


def _show_existing_notices():
    """Show notices already extracted and stored on disk."""
    notices_dir = Path("data/notices")
    if not notices_dir.exists():
        st.info("Henüz notice verisi yok (data/notices/ boş)")
        return

    notice_files = sorted(notices_dir.glob("*.json"), key=lambda p: p.stat().st_mtime,
                          reverse=True)

    if not notice_files:
        st.info("Henüz notice verisi yok")
        return

    st.markdown(f"**{len(notice_files)} notice dosyası bulundu**")

    for nf in notice_files[:20]:
        try:
            with open(nf, "r", encoding="utf-8") as f:
                data = json.load(f)

            doc_type = data.get("doc_type", "?")
            date = data.get("date", "?")
            sender = data.get("sender", "?")
            recipient = data.get("recipient", "?")
            subject = data.get("subject", "")[:60]

            with st.expander(f"📄 {nf.stem} — {doc_type} ({date})"):
                st.markdown(f"**Gönderen:** {sender}")
                st.markdown(f"**Alıcı:** {recipient}")
                if subject:
                    st.markdown(f"**Konu:** {subject}")
                actions = data.get("actions", [])
                if actions:
                    st.markdown(f"**Aksiyonlar:** {', '.join(actions)}")
        except Exception:
            continue


def _show_project_profile():
    """Show project profile from document_agent if available."""
    try:
        from src.document_agent import get_document_agent
        agent = get_document_agent()

        if hasattr(agent, 'profile') and agent.profile:
            profile = agent.profile

            st.markdown(f"**Proje:** {profile.project_name or '—'}")
            st.markdown(f"**Dönem:** {profile.date_range or '—'}")
            st.markdown(f"**Doküman Sayısı:** {profile.doc_count}")

            if profile.parties:
                rows = []
                for name, info in profile.parties.items():
                    rows.append({
                        "Taraf": name,
                        "Rol": info.get("role", "—"),
                        "Gönderen": info.get("sent_count", 0),
                        "Alınan": info.get("received_count", 0),
                    })
                st.dataframe(
                    pd.DataFrame(rows).sort_values("Gönderen", ascending=False),
                    use_container_width=True,
                    hide_index=True,
                )

            if profile.insights:
                st.markdown("**İçgörüler:**")
                SEVERITY_ICONS = {"high": "🔴", "medium": "🟡", "low": "🟢"}
                for ins in profile.insights:
                    icon = SEVERITY_ICONS.get(ins.severity, "⚪")
                    st.markdown(f"{icon} **{ins.insight_type}**: {ins.description}")

        else:
            st.info("Proje profili henüz oluşturulmadı. "
                    "Ana uygulamadan dokümanlar yükleyip analiz çalıştırın.")

    except Exception as e:
        st.info(f"Proje profili yüklenemedi: {e}")


def _show_timeline(graph):
    """Show document timeline."""
    try:
        timeline = graph.timeline(limit=30)
        if timeline:
            for doc in timeline:
                date = doc.get("date", "?")
                doc_type = doc.get("doc_type", "📄")
                sender = doc.get("sender", "?")
                recipient = doc.get("recipient", "?")
                subject = doc.get("subject", "")[:60]

                icon = {"letter": "📄", "email": "📧", "notice": "📋",
                        "report": "📊", "invoice": "💰"}.get(
                    str(doc_type).lower(), "📄"
                )

                st.markdown(
                    f"**{date}** {icon} {doc_type}: {subject or '—'} "
                    f"({sender} → {recipient})"
                )

                edges = doc.get("edges", [])
                for edge in edges[:3]:
                    st.markdown(
                        f"    └─ {edge.get('edge_type', '?')} → {edge.get('target', '?')}"
                    )
        else:
            st.info("Timeline boş")
    except Exception as e:
        st.warning(f"Timeline hatası: {e}")


def _show_communication_flow(graph):
    """Show communication flow between parties."""
    try:
        flow = graph.communication_flow()
        if flow:
            rows = []
            for item in flow:
                rows.append({
                    "Gönderen": item.get("sender", "?"),
                    "Alıcı": item.get("recipient", "?"),
                    "Sayı": item.get("count", 0),
                })
            flow_df = pd.DataFrame(rows).sort_values("Sayı", ascending=False)
            st.dataframe(flow_df, use_container_width=True, hide_index=True)
        else:
            st.info("İletişim verisi bulunamadı")
    except Exception as e:
        st.warning(f"İletişim akışı hatası: {e}")


def _render_extractor_result(df: pd.DataFrame, extractor_name: str,
                             target_schema: str, file_path: str):
    """Render a single extractor result with insight."""
    SCHEMA_LABELS = {
        "manpower_production": "Manpower Production Log",
        "equipment_log": "Equipment Log",
        "ipc_sample": "IPC (Interim Progress Certificate)",
    }
    label = SCHEMA_LABELS.get(target_schema, target_schema)

    st.success(
        f"✅ **{label}** — {len(df)} satır | "
        f"Extractor: `{extractor_name}` (pre-built, LLM kullanılmadı)"
    )

    st.dataframe(df.head(20), use_container_width=True, hide_index=True)
    st.caption(f"Sütunlar: {', '.join(df.columns)}")

    # Insight
    try:
        from src.table_insight_extractor import extract_table_insight
        insight = extract_table_insight(df, file_path, target_schema)
        _render_insight(insight)
    except Exception as e:
        st.warning(f"Insight hatası: {e}")


# ── Title ───────────────────────────────────────────────────
st.title("🔍 Ön İşleme Gözlem Servisi")
st.caption("Dosya işleme pipeline sonuçlarını gözlemle — Excel dönüşüm, doküman notice, ilişki grafı")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📊 Excel Dönüşüm",
    "📄 Doküman Notice",
    "🔗 İlişki & Kronoloji",
    "📧 Email Analiz",
    "🔍 SQL Test",
])


# ══════════════════════════════════════════════════════════════
# TAB 1: Excel Dönüşüm
# ══════════════════════════════════════════════════════════════
with tab1:
    st.header("📊 Excel Format Dönüşüm Gözlemi")

    # Show available target schemas
    with st.expander("📋 Hedef Formatlar (formatlar/ klasörü)", expanded=False):
        try:
            from src.schema_converter import get_target_schemas
            schema_registry = get_target_schemas()
            schema_list = schema_registry.list_schemas()

            cols = st.columns(len(schema_list)) if schema_list else []
            for i, schema_info in enumerate(schema_list):
                schema_obj = schema_registry.get_schema(schema_info["schema_id"])
                if not schema_obj:
                    continue
                with cols[i]:
                    st.markdown(f"**{schema_obj.name}**")
                    st.caption(schema_obj.description[:80])
                    st.markdown(f"Sütun: **{len(schema_obj.columns)}** "
                                f"(zorunlu: {sum(1 for c in schema_obj.columns if c.required)})")
                    for col_def in schema_obj.columns:
                        icon = "🔴" if col_def.required else "⚪"
                        st.markdown(f"{icon} `{col_def.name}` ({col_def.dtype})")
        except Exception as e:
            st.warning(f"Şemalar yüklenemedi: {e}")

    # File upload
    uploaded_excel = st.file_uploader(
        "Excel dosyası yükle",
        type=["xlsx", "xls", "csv"],
        key="excel_upload",
    )

    if uploaded_excel:
        # Save to temp file
        tmp_dir = tempfile.mkdtemp()
        tmp_path = os.path.join(tmp_dir, uploaded_excel.name)
        with open(tmp_path, "wb") as f:
            f.write(uploaded_excel.getbuffer())

        # Show source sheets info
        st.subheader("📄 Kaynak Veri")
        try:
            ext = Path(uploaded_excel.name).suffix.lower()
            if ext in (".xlsx", ".xls"):
                xls = pd.ExcelFile(tmp_path)
                st.caption(f"Sheet sayısı: {len(xls.sheet_names)} — {', '.join(xls.sheet_names)}")

                # Show each sheet preview
                for sheet_name in xls.sheet_names:
                    try:
                        sdf = pd.read_excel(tmp_path, sheet_name=sheet_name, header=None, nrows=10)
                        if sdf.empty or sdf.dropna(how="all").empty:
                            continue
                        with st.expander(f"Sheet: **{sheet_name}** ({sdf.shape[0]}+ satır, {sdf.shape[1]} sütun)"):
                            st.dataframe(sdf, use_container_width=True, hide_index=True)
                    except Exception:
                        continue
            elif ext == ".csv":
                source_df = pd.read_csv(tmp_path, nrows=20)
                st.dataframe(source_df, use_container_width=True, hide_index=True)
                st.caption(f"Sütunlar: {', '.join(source_df.columns)}")
        except Exception as e:
            st.error(f"Dosya okunamadı: {e}")

        # Step 1: Try pre-built extractors (no LLM)
        st.subheader("🔄 Dönüşüm Sonucu")

        extractor_success = False
        try:
            from src.extractors import match_extractor, run_extractor

            matches = match_extractor(tmp_path)
            if matches:
                st.info(f"📌 Pre-built extractor eşleşti: {len(matches)} tablo tespit edildi")

                for extractor_name, target_schema in matches:
                    with st.spinner(f"{extractor_name} çalıştırılıyor..."):
                        df = run_extractor(extractor_name, tmp_path)

                    if df is not None and not df.empty:
                        _render_extractor_result(df, extractor_name, target_schema, tmp_path)
                        extractor_success = True
                    else:
                        st.warning(f"⚠️ {extractor_name} → veri çıkarılamadı")

        except Exception as e:
            st.warning(f"Extractor hatası: {e}")

        # Step 2: If no extractor matched, try LLM-based FormatConverter
        if not extractor_success:
            try:
                from src.schema_converter import FormatConverter

                with st.spinner("LLM ile dönüştürülüyor... (şema eşleştirme + kod üretimi)"):
                    converter = FormatConverter()
                    conv_result = converter.process_excel(tmp_path)

                if conv_result and conv_result.success and conv_result.df is not None:
                    st.success(
                        f"✅ Eşleşen Format: **{conv_result.target_schema}** | "
                        f"Satır: {len(conv_result.df)} | "
                        f"{'Yeni converter üretildi (LLM)' if conv_result.generated else 'Kayıtlı converter'}"
                    )

                    st.markdown("**Dönüştürülmüş Tablo:**")
                    st.dataframe(conv_result.df.head(20), use_container_width=True, hide_index=True)

                    if conv_result.validation_errors:
                        st.warning("Validasyon uyarıları:")
                        for err in conv_result.validation_errors:
                            st.markdown(f"- {err}")

                    st.subheader("💡 Tablo Insight")
                    from src.table_insight_extractor import extract_table_insight
                    insight = extract_table_insight(
                        conv_result.df, tmp_path, conv_result.target_schema
                    )
                    _render_insight(insight)

                elif conv_result and not conv_result.success:
                    st.error("❌ Dönüşüm başarısız")
                    if conv_result.validation_errors:
                        for err in conv_result.validation_errors:
                            st.markdown(f"- {err}")
                else:
                    st.warning("⚠️ Hiçbir hedef şemaya eşleştirilemedi. "
                               "Dosya formatlar/ klasöründeki 3 formattan birine uymuyor olabilir.")

            except Exception as e:
                st.error(f"Dönüşüm hatası: {e}")
                import traceback
                with st.expander("Hata detayı"):
                    st.code(traceback.format_exc())


# ══════════════════════════════════════════════════════════════
# TAB 2: Doküman Notice
# ══════════════════════════════════════════════════════════════
with tab2:
    st.header("📄 Doküman Notice Çıkarma Gözlemi")

    uploaded_doc = st.file_uploader(
        "Doküman yükle (PDF, DOCX, TXT)",
        type=["pdf", "docx", "doc", "txt"],
        key="doc_upload",
    )

    if uploaded_doc:
        tmp_dir = tempfile.mkdtemp()
        tmp_path = os.path.join(tmp_dir, uploaded_doc.name)
        with open(tmp_path, "wb") as f:
            f.write(uploaded_doc.getbuffer())

        with st.spinner("Doküman işleniyor..."):
            try:
                doc_text_by_page = _extract_doc_text(tmp_path)

                if doc_text_by_page:
                    st.success(f"✅ {len(doc_text_by_page)} sayfa okundu")

                    with st.expander("📖 Metin Önizleme"):
                        for page_num, text in sorted(doc_text_by_page.items())[:5]:
                            st.markdown(f"**Sayfa {page_num}:**")
                            st.text(text[:500] + ("..." if len(text) > 500 else ""))

                    st.subheader("📋 Notice Metadata")
                    _extract_and_display_notice(tmp_path, doc_text_by_page)

                    st.subheader("🏷️ Issue Kategorileri")
                    _review_and_display_categories(tmp_path, doc_text_by_page)

                else:
                    st.warning("Dokümandan metin çıkarılamadı")

            except Exception as e:
                st.error(f"İşleme hatası: {e}")
                import traceback
                with st.expander("Hata detayı"):
                    st.code(traceback.format_exc())

    st.divider()
    st.subheader("📂 Mevcut Notice Verileri")
    _show_existing_notices()


# ══════════════════════════════════════════════════════════════
# TAB 3: İlişki & Kronoloji
# ══════════════════════════════════════════════════════════════
with tab3:
    st.header("🔗 Doküman İlişkileri & Kronoloji")

    try:
        from src.light_graph import get_document_graph
        graph = get_document_graph()
        stats = graph.get_statistics()

        if stats.get("node_count", 0) > 0:
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Doküman", stats.get("node_count", 0))
            c2.metric("İlişki", stats.get("edge_count", 0))
            c3.metric("Gönderen", stats.get("unique_senders", 0))
            c4.metric("Alıcı", stats.get("unique_recipients", 0))

            st.subheader("🏗️ Proje Profili")
            _show_project_profile()

            st.subheader("📅 Kronoloji")
            _show_timeline(graph)

            st.subheader("📨 İletişim Akışı")
            _show_communication_flow(graph)

            st.subheader("🔗 İlişki Tipleri")
            edge_types = stats.get("edge_types", {})
            if edge_types:
                edge_df = pd.DataFrame([
                    {"Tip": k, "Sayı": v} for k, v in edge_types.items()
                ])
                st.dataframe(edge_df, use_container_width=True, hide_index=True)
        else:
            st.info("Henüz doküman grafiğinde veri yok. "
                    "Ana uygulamadan doküman yükleyerek grafı oluşturun.")

    except Exception as e:
        st.warning(f"Graf yüklenemedi: {e}")
        st.info("Doküman grafiği oluşturulmamış olabilir. "
                "Ana uygulamadan doküman yükleyin.")


# ══════════════════════════════════════════════════════════════
# TAB 4: Email Analiz
# ══════════════════════════════════════════════════════════════
with tab4:
    st.header("📧 Email Analiz Gözlemi")

    uploaded_email = st.file_uploader(
        "Email dosyası yükle",
        type=["eml", "msg"],
        key="email_upload",
    )

    if uploaded_email:
        tmp_dir = tempfile.mkdtemp()
        tmp_path = os.path.join(tmp_dir, uploaded_email.name)
        with open(tmp_path, "wb") as f:
            f.write(uploaded_email.getbuffer())

        try:
            from src.email_parser import EmailParser
            parser = EmailParser()
            parsed = parser.parse(tmp_path)

            st.subheader("📋 Email Metadata")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"**From:** {parsed.sender}")
                if hasattr(parsed, 'sender_email') and parsed.sender_email:
                    st.caption(parsed.sender_email)
                st.markdown(f"**To:** {', '.join(parsed.recipients) if parsed.recipients else '—'}")
                if parsed.cc:
                    st.markdown(f"**CC:** {', '.join(parsed.cc)}")
            with col2:
                st.markdown(f"**Date:** {parsed.date or '—'}")
                st.markdown(f"**Subject:** {parsed.subject or '—'}")

            if parsed.body_text:
                with st.expander("📖 Email İçeriği"):
                    st.text(parsed.body_text[:2000])

            if parsed.attachments:
                st.subheader(f"📎 Ekler ({len(parsed.attachments)})")
                for att in parsed.attachments:
                    ext = Path(att.filename).suffix.lower()
                    icon = "📊" if ext in [".xlsx", ".xls", ".csv"] else "📄"
                    st.markdown(f"{icon} **{att.filename}** ({att.size / 1024:.1f} KB)")

            if parsed.body_text:
                page_texts = parser.to_document_text(parsed)
                if page_texts:
                    st.subheader("📋 Email Notice")
                    _extract_and_display_notice(tmp_path, page_texts)

        except Exception as e:
            st.error(f"Email parse hatası: {e}")
            import traceback
            with st.expander("Hata detayı"):
                st.code(traceback.format_exc())


# ══════════════════════════════════════════════════════════════
# TAB 5: SQL Test
# ══════════════════════════════════════════════════════════════
with tab5:
    st.header("🔍 DuckDB SQL Test")

    try:
        from src.data_analyzer_sql import DataAnalyzerSQL
        from src.catalog import get_catalog

        if "debug_analyzer" not in st.session_state:
            analyzer = DataAnalyzerSQL()
            analyzer.load_from_catalog()
            st.session_state.debug_analyzer = analyzer
        else:
            analyzer = st.session_state.debug_analyzer

        if analyzer.tables:
            st.subheader("📋 Yüklü Tablolar")
            for name, info in analyzer.tables.items():
                if name.endswith("_raw") or name.endswith("_clean"):
                    continue

                cols_preview = ", ".join(info.get("columns", [])[:6])
                if len(info.get("columns", [])) > 6:
                    cols_preview += "..."

                narrative = info.get("insight", {}).get("narrative", "")
                desc_text = narrative or info.get("description", "")

                with st.expander(f"**{name}** — {info.get('row_count', 0)} satır"):
                    if desc_text:
                        st.info(desc_text)
                    st.markdown(f"**Sütunlar:** `{cols_preview}`")

                    insight = info.get("insight", {})
                    if insight.get("stats"):
                        st.markdown("**İstatistikler:**")
                        for key, val in insight["stats"].items():
                            if isinstance(val, dict):
                                st.markdown(f"- {key}: {val}")
                            else:
                                st.markdown(f"- {key}: **{val}**")

                    try:
                        preview = analyzer.conn.execute(
                            f"SELECT * FROM {name} LIMIT 5"
                        ).fetchdf()
                        st.dataframe(preview, use_container_width=True, hide_index=True)
                    except Exception:
                        pass

            st.divider()
            st.subheader("💻 SQL Sorgusu Çalıştır")
            sql_input = st.text_area(
                "SQL",
                value="SELECT * FROM ... LIMIT 10",
                height=100,
                key="sql_input",
            )

            col1, col2 = st.columns([1, 3])
            with col1:
                run_sql = st.button("▶️ Çalıştır", type="primary")
            with col2:
                nl_query = st.text_input(
                    "Veya doğal dil ile soru sor:",
                    placeholder="Toplam vinç saati nedir?",
                    key="nl_query",
                )
                run_nl = st.button("🤖 Soru Sor")

            if run_sql and sql_input and sql_input != "SELECT * FROM ... LIMIT 10":
                try:
                    result = analyzer.conn.execute(sql_input).fetchdf()
                    st.success(f"✅ {len(result)} satır döndü")
                    st.dataframe(result, use_container_width=True, hide_index=True)
                except Exception as e:
                    st.error(f"SQL hatası: {e}")

            if run_nl and nl_query:
                with st.spinner("Sorgu oluşturuluyor..."):
                    result = analyzer.query(nl_query)
                    st.markdown(f"**Cevap:** {result.get('answer', '—')}")
                    if result.get("sql"):
                        st.code(result["sql"], language="sql")
                    if result.get("result_data"):
                        st.dataframe(
                            pd.DataFrame(result["result_data"]),
                            use_container_width=True,
                            hide_index=True,
                        )
        else:
            st.info("Henüz DuckDB'de tablo yok. "
                    "Ana uygulamadan Excel/CSV yükleyin veya Excel Dönüşüm tab'ını kullanın.")

        st.divider()
        st.subheader("📊 Katalog İstatistikleri")
        catalog = get_catalog()
        cat_stats = catalog.get_stats()
        c1, c2, c3 = st.columns(3)
        c1.metric("Dosya", cat_stats.get("total_entries", 0))
        c2.metric("Tablo", cat_stats.get("total_tables", 0))
        c3.metric("Toplam Satır", cat_stats.get("total_rows", 0))

    except Exception as e:
        st.error(f"DuckDB bağlantı hatası: {e}")
        import traceback
        with st.expander("Hata detayı"):
            st.code(traceback.format_exc())


# ── Run ─────────────────────────────────────────────────────
if __name__ == "__main__":
    pass
