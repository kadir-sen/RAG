"""Report Assembler.

Assembles the per-module narratives and key findings into one Word document:
title page, one chapter per included module, a single deduplicated
Limitations section aggregated from every module's caveats and warnings, and
a "Basis of Analysis" appendix recording exactly what the report was built
from (files, SHA-256 hashes, data dates, settings used).

The aggregation is pure concatenation/deduplication — no LLM. Narrative
chapters reproduce the module narratives the analyst generated (and
reviewed) in the app; modules without a generated narrative contribute their
deterministic key findings only.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

# python-docx is imported lazily inside build_assembled_report(), not here,
# so a missing/broken install of this one optional dependency can't take
# down `import programme` (and therefore the whole app) at startup — only
# the report-assembly feature itself.
ACCENT_RGB = (0x1F, 0x38, 0x64)

STANDING_REPORT_CAVEAT = (
    "This document is a preliminary factual screening assembled from the "
    "programme files listed in the Basis of Analysis appendix. It describes "
    "movement and change recorded in those files; it does not attribute "
    "cause or responsibility and is not a cause-linked delay analysis."
)


@dataclass
class ReportSection:
    title: str
    narrative_md: str | None = None       # analyst-generated narrative
    key_findings: list[str] = field(default_factory=list)
    caveats: list[str] = field(default_factory=list)   # feeds Limitations
    images: list[tuple[bytes, str]] = field(default_factory=list)  # (png, caption)


@dataclass
class SourceFile:
    file_name: str
    sha256: str
    data_date: datetime | None
    role: str                             # Baseline / Update / Current
    activity_count: int


@dataclass
class BasisOfAnalysis:
    files: list[SourceFile] = field(default_factory=list)
    settings: list[str] = field(default_factory=list)  # "Module — setting: x"
    generated_at: datetime = field(default_factory=datetime.now)
    tool_note: str = ("Deterministic engines computed all figures; where an "
                      "AI narrative is included it was generated from those "
                      "figures under fixed objectivity rules and reviewed "
                      "by the analyst before inclusion.")


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _stamp_docx(doc) -> None:
    """Build provenance in the page footer of every section — the same
    stamp the Excel exports carry, so any document can testify which
    code revision produced it (audit F-04, answered per-export)."""
    from docx.shared import Pt as _Pt_
    from buildinfo import build_stamp
    stamp = build_stamp()
    for section in doc.sections:
        para = section.footer.paragraphs[0]
        para.text = stamp
        for run in para.runs:
            run.font.size = _Pt_(7)
            run.font.italic = True
    try:
        doc.core_properties.comments = stamp
    except Exception:  # noqa: BLE001 - never block the document
        pass


def _add_md_paragraph(doc: Document, text: str, style: str | None = None):
    """Add one paragraph, converting **bold** spans."""
    p = doc.add_paragraph(style=style)
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        p.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])
    return p


def _md_table(doc: Document, lines: list[str]) -> None:
    """A buffered run of |-delimited markdown lines as a real Word
    table — narratives are full of tables, and rendering them as text
    left the reader a wall of pipes and dashes."""
    from docx.shared import Pt
    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        # the |---|---| separator row carries no content
        if cells and all(re.fullmatch(r":?-{2,}:?", c) for c in cells
                         if c):
            continue
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.style = "Table Grid"
    for i, r in enumerate(rows):
        for j in range(ncols):
            cell = t.cell(i, j)
            cell.text = _BOLD_RE.sub(r"\1", r[j]) if j < len(r) else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if i == 0:
                        run.bold = True
    doc.add_paragraph()


def _add_markdown(doc: Document, md: str, base_heading_level: int) -> None:
    """Minimal markdown renderer: #/##/### headings, - bullets,
    | tables |, paragraphs."""
    table_buf: list[str] = []

    def flush_table() -> None:
        if table_buf:
            _md_table(doc, table_buf)
            table_buf.clear()

    for raw in md.split("\n"):
        line = raw.rstrip()
        if line.strip().startswith("|") and line.count("|") >= 2:
            table_buf.append(line)
            continue
        flush_table()
        if not line.strip():
            continue
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            hashes, text = m.groups()
            # Narrative "##" becomes a level under the chapter heading.
            level = min(base_heading_level + max(len(hashes) - 2, 0), 4)
            doc.add_heading(_BOLD_RE.sub(r"\1", text), level=level)
        elif re.match(r"^\s*[-*•]\s+", line):
            _add_md_paragraph(doc, re.sub(r"^\s*[-*•]\s+", "", line),
                              style="List Bullet")
        elif re.match(r"^\s*\d+\.\s+", line):
            _add_md_paragraph(doc, re.sub(r"^\s*\d+\.\s+", "", line),
                              style="List Number")
        else:
            _add_md_paragraph(doc, line)
    flush_table()


def _dedupe(items: list[str]) -> list[str]:
    seen, out = set(), []
    for i in items:
        key = i.strip().lower()
        if key and key not in seen:
            seen.add(key)
            out.append(i.strip())
    return out


def _data_table(doc, rows: list[dict], max_rows: int = 400) -> None:
    """A list of row-dicts as a real Word table (appendix use)."""
    from docx.shared import Pt
    if not rows:
        return
    headers = list(rows[0].keys())
    shown = rows[:max_rows]
    t = doc.add_table(rows=len(shown) + 1, cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8.5)
    for i, row in enumerate(shown, start=1):
        for j, h in enumerate(headers):
            v = row.get(h)
            cell = t.cell(i, j)
            cell.text = "" if v is None else str(v)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    if len(rows) > max_rows:
        p = doc.add_paragraph()
        run = p.add_run(f"({len(rows) - max_rows} further rows omitted "
                        f"— the complete set is in the Excel workbook.)")
        run.italic = True
    doc.add_paragraph()


def build_narrative_docx(title: str, narrative_md: str,
                         images: list[tuple[str, bytes]] | None = None,
                         ) -> bytes:
    """One module's AI narrative as a Word document.

    Analysts deliver in Word; a .md file forces a paste-and-restyle
    detour. Reuses the assembled-report markdown renderer so headings,
    bullets and numbering carry through. ``images`` are (caption, PNG)
    figures placed at the FRONT of the document, before the narrative —
    the final gantt is the exhibit that tells the whole delay story, so
    it leads and the words follow.

    The narrative stays SHORT by design: the complete tables ship as a
    separate appendix workbook (report_xlsx.build_appendix_xlsx). Word
    tables cost roughly 5ms per row to build, so a 7,000-row appendix
    took 37 seconds inline against 0.4 in Excel — the reader gets a
    document that opens instantly and a workbook they can filter."""
    from docx import Document as _Document
    from docx.shared import Inches as _Inches
    doc = _Document()
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"Generated {datetime.now():%d %b %Y %H:%M} — "
                    "AI-drafted from the module's computed figures; "
                    "review before inclusion in any deliverable.")
    run.italic = True
    for cap, png in (images or []):
        if not png:
            continue
        doc.add_heading(cap, level=2)
        doc.add_picture(io.BytesIO(png), width=_Inches(6.4))
    _add_markdown(doc, narrative_md, base_heading_level=2)
    _stamp_docx(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_assembled_report(
    report_title: str,
    project_name: str,
    author: str,
    sections: list[ReportSection],
    basis: BasisOfAnalysis,
) -> bytes:
    """Assemble the full preliminary report as a .docx (returns bytes)."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    accent = RGBColor(*ACCENT_RGB)
    doc = Document()
    styles = doc.styles["Normal"]
    styles.font.name = "Calibri"
    styles.font.size = Pt(10.5)

    # --- title page -------------------------------------------------------
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(report_title or "Preliminary Delay Analysis Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = accent
    for line in (project_name, author,
                 f"Issued {basis.generated_at:%d %B %Y}",
                 "PRELIMINARY — for review and discussion"):
        if line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(line).font.size = Pt(12)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run(STANDING_REPORT_CAVEAT).italic = True
    doc.add_page_break()

    # --- contents (static list; Word can replace with a live TOC) ---------
    doc.add_heading("Contents", level=1)
    for i, s in enumerate(sections, start=1):
        doc.add_paragraph(f"{i}. {s.title}", style="List Number")
    doc.add_paragraph(f"{len(sections) + 1}. Limitations",
                      style="List Number")
    doc.add_paragraph(f"{len(sections) + 2}. Appendix A — Basis of Analysis",
                      style="List Number")
    doc.add_page_break()

    # --- module chapters ---------------------------------------------------
    fig_no = 0
    for i, s in enumerate(sections, start=1):
        doc.add_heading(f"{i}. {s.title}", level=1)
        if s.key_findings:
            doc.add_heading("Key figures", level=2)
            for kf in s.key_findings:
                _add_md_paragraph(doc, kf, style="List Bullet")
        for png, caption in s.images:
            fig_no += 1
            doc.add_picture(io.BytesIO(png), width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(f"Figure {fig_no} — {caption}")
            run.italic = True
            run.font.size = Pt(9)
        if s.narrative_md:
            _add_markdown(doc, s.narrative_md, base_heading_level=2)
        else:
            p = doc.add_paragraph()
            p.add_run(
                "No narrative was generated for this module; the key "
                "figures above are reported without commentary."
            ).italic = True

    # --- limitations (aggregated, deduplicated) ----------------------------
    doc.add_heading(f"{len(sections) + 1}. Limitations", level=1)
    doc.add_paragraph(
        "The following limitations apply to this report. They are "
        "aggregated from every analysis included and are reproduced in "
        "full; none has been omitted."
    )
    all_caveats = _dedupe(
        [STANDING_REPORT_CAVEAT]
        + [c for s in sections for c in s.caveats])
    for c in all_caveats:
        _add_md_paragraph(doc, c, style="List Bullet")

    # --- basis of analysis --------------------------------------------------
    doc.add_page_break()
    doc.add_heading(f"{len(sections) + 2}. Appendix A — Basis of Analysis",
                    level=1)
    doc.add_paragraph(
        "This appendix records the source files, their integrity hashes, "
        "and the analysis settings from which every figure in this report "
        "was computed, so the analysis can be independently reproduced."
    )
    doc.add_heading("Source programme files", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(["File", "Role", "Data date", "Activities",
                           "SHA-256 (first 16)"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    for f in basis.files:
        row = table.add_row().cells
        row[0].text = f.file_name
        row[1].text = f.role
        row[2].text = (f"{f.data_date:%Y-%m-%d}" if f.data_date else "—")
        row[3].text = str(f.activity_count)
        row[4].text = f.sha256[:16]

    if basis.settings:
        doc.add_heading("Analysis settings", level=2)
        for s in basis.settings:
            _add_md_paragraph(doc, s, style="List Bullet")

    doc.add_heading("Method note", level=2)
    doc.add_paragraph(basis.tool_note)
    doc.add_paragraph(
        f"Report assembled {basis.generated_at:%Y-%m-%d %H:%M}."
    )

    _stamp_docx(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
