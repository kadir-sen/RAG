"""Report section exports (PDF/DOCX) + artifact persistence.

The PDF/DOCX contract IS the HTML contract, so this reuses the same section
fixture shape as tests/test_orchestration_html.py.
"""

import io

import pytest

from backend.models.blocks import validate_blocks
from src.export import artifacts as export_artifacts
from src.export.section import render_section, section_artifacts
from src.orchestration.helpers import artifact_blocks


def section(**kw):
    base = dict(
        title="Delayed Blockwork",
        narrative_markdown=("6.1.1 On 19 July 2023, **JAMED** raised concerns. "
                            "(L1.pdf, p.3)\n\nSecond paragraph."),
        tables=[{"title": "Chronology", "columns": ["Date", "Actor"],
                 "rows": [["19 July 2023", "JAMED"]]}],
        sources=[{"file_name": "L1.pdf", "page_number": 3}],
        caveats=["Statements are the parties' claims."],
        validation={"computation_guard": {"post": "passed"}},
        section_number="6.1",
    )
    base.update(kw)
    return base


def _body(**kw):
    s = section(**kw)
    return s.pop("title"), s.pop("narrative_markdown"), s


def _pdf_text(data):
    import fitz
    with fitz.open(stream=data, filetype="pdf") as doc:
        return "".join(p.get_text() for p in doc)


def _docx_doc(data):
    import docx
    return docx.Document(io.BytesIO(data))


class TestPdf:
    def test_renders_every_part_of_the_section(self):
        title, narrative, kw = _body()
        data, reason = render_section("pdf", title, narrative, **kw)
        assert data, reason
        assert data[:5] == b"%PDF-"
        text = _pdf_text(data)
        for probe in ["6.1. Delayed Blockwork", "JAMED", "19 July 2023",
                      "L1.pdf", "parties' claims", "computation_guard"]:
            assert probe in text, f"{probe!r} missing from the PDF"

    def test_hostile_narrative_stays_text(self):
        """Same payload matrix as TestSanitizerXSS: the templater escapes, so
        a script tag must arrive as characters, never as markup."""
        title, narrative, kw = _body(
            narrative_markdown="<script>alert(1)</script> **bold** text")
        data, reason = render_section("pdf", title, narrative, **kw)
        assert data, reason
        text = _pdf_text(data)
        assert "<script>" in text        # rendered as literal text
        assert "bold" in text

    def test_large_table_renders(self):
        rows = [[f"{i} Jan 2024", f"Party {i}"] for i in range(500)]
        title, narrative, kw = _body(
            tables=[{"title": "Chronology", "columns": ["Date", "Actor"],
                     "rows": rows}])
        data, reason = render_section("pdf", title, narrative, **kw)
        assert data, reason
        assert data[:5] == b"%PDF-"


class TestDocx:
    def test_renders_every_part_of_the_section(self):
        title, narrative, kw = _body()
        data, reason = render_section("docx", title, narrative, **kw)
        assert data, reason
        assert data[:4] == b"PK\x03\x04"
        doc = _docx_doc(data)
        assert doc.paragraphs[0].text == "6.1. Delayed Blockwork"
        assert len(doc.tables) == 1
        assert [c.text for c in doc.tables[0].rows[1].cells] == [
            "19 July 2023", "JAMED"]
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "L1.pdf, p.3" in text
        assert "parties' claims" in text
        assert "computation_guard: passed" in text

    def test_bold_and_paragraph_number_become_runs_not_literals(self):
        title, narrative, kw = _body()
        doc = _docx_doc(render_section("docx", title, narrative, **kw)[0])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "**" not in text          # markdown consumed, not printed
        assert "6.1.1" in text
        bolded = [r.text for p in doc.paragraphs for r in p.runs if r.bold]
        assert "JAMED" in bolded

    def test_hostile_narrative_stays_text(self):
        title, narrative, kw = _body(
            narrative_markdown="<script>alert(1)</script> plain")
        doc = _docx_doc(render_section("docx", title, narrative, **kw)[0])
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "<script>alert(1)</script>" in text


class TestCapParity:
    """The DOCX renderer re-implements the section layout, so its caps must
    equal the templater's — this is the drift that file risks."""

    def test_caps_match_html_section(self):
        import inspect

        from src.export import section_docx
        from src.orchestration import html_section

        src = inspect.getsource(html_section.compose_section)
        assert f"[:{section_docx.MAX_ROWS}]" in src
        assert f"[:{section_docx.MAX_SOURCES}]" in src
        assert f"[:{section_docx.MAX_CAVEATS}]" in src

    def test_docx_truncates_rows_at_the_cap(self):
        from src.export.section_docx import MAX_ROWS
        rows = [[f"{i}", "x"] for i in range(MAX_ROWS + 50)]
        title, narrative, kw = _body(
            tables=[{"title": "T", "columns": ["A", "B"], "rows": rows}])
        doc = _docx_doc(render_section("docx", title, narrative, **kw)[0])
        assert len(doc.tables[0].rows) == MAX_ROWS + 1     # +1 header


class TestSectionArtifacts:
    @pytest.fixture(autouse=True)
    def tmp_artifacts(self, tmp_path, monkeypatch):
        monkeypatch.setattr("src.programme_tools.config_paths.artifacts_dir",
                            lambda: tmp_path)
        self.dir = tmp_path

    def test_writes_both_formats_and_returns_artifact_dicts(self):
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, run_id="r1", **kw)
        assert {a["kind"] for a in arts} == {"pdf", "docx"}
        for a in arts:
            assert a["url"].startswith("/api/artifacts/r1/")
            assert (self.dir / "r1" / a["filename"]).is_file()

    def test_artifacts_survive_the_response_block_guard(self):
        """Closes the loop on ArtifactLinkBlock's /api/artifacts/ validator."""
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, run_id="r1", **kw)
        blocks = artifact_blocks({"artifacts": arts})
        assert len(blocks) == 2
        valid, dropped = validate_blocks(blocks)
        assert dropped == []
        assert len(valid) == 2

    def test_formats_can_be_narrowed(self):
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, run_id="r1",
                                 formats=("docx",), **kw)
        assert [a["kind"] for a in arts] == ["docx"]

    def test_unknown_format_is_skipped_not_raised(self):
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, run_id="r1",
                                 formats=("exe",), **kw)
        assert arts == []

    def test_title_cannot_escape_the_run_directory(self):
        title, narrative, kw = _body(title="../../etc/passwd")
        arts = section_artifacts(title, narrative, run_id="r1", **kw)
        assert arts
        for a in arts:
            # What matters is where the bytes land, not whether the name still
            # contains dots — the separators are gone, so it cannot traverse.
            assert "/" not in a["filename"] and "\\" not in a["filename"]
            written = (self.dir / "r1" / a["filename"]).resolve()
            assert written.parent == (self.dir / "r1").resolve()
            assert written.is_file()

    @pytest.mark.parametrize("name", ["..", ".", "...", "/", "../..", ""])
    def test_dot_only_names_cannot_address_a_directory(self, name):
        assert export_artifacts.safe_filename(name) == "artifact.bin"

    def test_run_id_defaults_to_a_fresh_id(self):
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, **kw)
        assert arts and arts[0]["url"].startswith("/api/artifacts/")

    def test_disk_failure_costs_the_download_not_the_section(self, monkeypatch):
        monkeypatch.setattr(export_artifacts, "persist_blobs",
                            lambda *a, **k: (_ for _ in ()).throw(OSError("full")))
        monkeypatch.setattr("src.export.section.persist_blobs",
                            lambda *a, **k: (_ for _ in ()).throw(OSError("full")))
        title, narrative, kw = _body()
        assert section_artifacts(title, narrative, run_id="r1", **kw) == []

    def test_render_failure_costs_only_that_format(self, monkeypatch):
        monkeypatch.setattr("src.export.section_pdf.render_section_pdf",
                            lambda *a, **k: (_ for _ in ()).throw(RuntimeError("boom")))
        title, narrative, kw = _body()
        arts = section_artifacts(title, narrative, run_id="r1", **kw)
        assert [a["kind"] for a in arts] == ["docx"]


class TestArtifactMediaTypes:
    def test_pdf_and_docx_are_served_with_real_media_types(self):
        from backend.api.artifacts import _MEDIA_TYPES
        assert _MEDIA_TYPES[".pdf"] == "application/pdf"
        assert "wordprocessingml" in _MEDIA_TYPES[".docx"]


class TestPackAggregation:
    def test_packs_carry_sub_workflow_downloads_through(self):
        """delay_claim_pack aggregates sub-results via content_blocks; a
        section's download links must survive that filter."""
        from src.workflows.blocks import CONTENT_TYPES, content_blocks
        from src.workflows.types import WorkflowId, WorkflowResult

        assert "artifact_link" in CONTENT_TYPES
        wr = WorkflowResult(
            workflow_id=WorkflowId.DELAY_CHRONOLOGY_SECTION, status="success",
            blocks=[
                {"type": "html_report_section", "block_id": "section"},
                {"type": "artifact_link", "block_id": "a1",
                 "url": "/api/artifacts/r1/x.pdf", "filename": "x.pdf",
                 "kind": "pdf"},
                {"type": "caveats", "block_id": "cav"},
                {"type": "validation_status", "block_id": "val"},
            ])
        kept = [b["type"] for b in content_blocks(wr)]
        assert kept == ["html_report_section", "artifact_link"]
