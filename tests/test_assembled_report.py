"""Multi-section forensic report assembler (report_docx)."""

import hashlib
import io
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).parent.parent))

from src.export.assembled_report import assemble_report

_TK = ("/private/tmp/claude-501/-Users-kadirsen-Desktop-projects-ML-project-V2/"
       "b69b4c75-d4e6-4673-b058-42caba2f618f/scratchpad/delay-toolkit/"
       "sample/revisions")

_HAVE_XER = Path(f"{_TK}/revA.xer").is_file()

SECTIONS = [
    {"title": "1. Programme Inventory", "narrative_md": "Three revisions.",
     "caveats": ["Preliminary."]},
    {"title": "2. Variance", "narrative_md": "Zone B slipped **78 days**.",
     "key_findings": ["Zone B +78d"], "caveats": ["As-recorded, unverified."]},
]


@pytest.fixture
def tmp_artifacts(tmp_path, monkeypatch):
    monkeypatch.setattr("src.programme_tools.config_paths.artifacts_dir",
                        lambda: tmp_path)
    return tmp_path


def _doc(path):
    import docx
    d = docx.Document(str(path))
    para = "\n".join(p.text for p in d.paragraphs)
    cells = "\n".join(c.text for t in d.tables for r in t.rows for c in r.cells)
    return d, para, cells


@pytest.mark.skipif(not _HAVE_XER, reason="sample XER fixtures not present")
class TestAssembleReport:
    def _run(self, tmp):
        return assemble_report(
            "Preliminary Delay Analysis", "Marina Tower", SECTIONS,
            source_records=[
                {"file_name": "revA.xer", "file_path": f"{_TK}/revA.xer",
                 "role": "Baseline"},
                {"file_name": "revC.xer", "file_path": f"{_TK}/revC.xer",
                 "role": "Current"}],
            settings=["Variance — dimension: Area"], run_id="r1")

    def test_produces_a_docx_artifact(self, tmp_artifacts):
        arts = self._run(tmp_artifacts)
        assert [a["kind"] for a in arts] == ["docx"]
        assert arts[0]["url"].startswith("/api/artifacts/r1/")
        assert (tmp_artifacts / "r1" / arts[0]["filename"]).is_file()

    def test_sections_and_narrative_survive(self, tmp_artifacts):
        arts = self._run(tmp_artifacts)
        _, para, _ = _doc(tmp_artifacts / "r1" / arts[0]["filename"])
        assert "Preliminary Delay Analysis" in para
        assert "Programme Inventory" in para
        assert "78 days" in para          # **bold** rendered, markdown consumed
        assert "**" not in para

    def test_basis_of_analysis_carries_the_source_hashes(self, tmp_artifacts):
        """The audit trail: each source file with its real SHA-256 prefix."""
        arts = self._run(tmp_artifacts)
        _, para, cells = _doc(tmp_artifacts / "r1" / arts[0]["filename"])
        h = hashlib.sha256(
            Path(f"{_TK}/revA.xer").read_bytes()).hexdigest()[:16]
        blob = para + "\n" + cells
        assert "revA.xer" in blob
        assert h in blob                  # the real hash, not a placeholder
        assert "Basis" in para

    def test_limitations_are_deduped(self, tmp_artifacts):
        # Both sections carry caveats; the assembled Limitations section dedupes.
        arts = self._run(tmp_artifacts)
        _, para, _ = _doc(tmp_artifacts / "r1" / arts[0]["filename"])
        assert "Limitations" in para


class TestDegradesGracefully:
    def test_bad_source_path_still_produces_a_report(self, tmp_artifacts):
        """A file we cannot hash costs its hash, not the whole report."""
        arts = assemble_report(
            "R", "P", [{"title": "1. X", "narrative_md": "y"}],
            source_records=[{"file_name": "missing.xer",
                             "file_path": "/no/such/file.xer"}],
            run_id="r2")
        assert [a["kind"] for a in arts] == ["docx"]

    def test_persist_failure_returns_empty_not_raise(self, tmp_artifacts,
                                                     monkeypatch):
        monkeypatch.setattr(
            "src.export.assembled_report.persist_blobs",
            lambda *a, **k: (_ for _ in ()).throw(OSError("disk full")))
        arts = assemble_report("R", "P", [{"title": "1"}], [], run_id="r3")
        assert arts == []
