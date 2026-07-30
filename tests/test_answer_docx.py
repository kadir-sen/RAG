"""A chat answer must export as a document someone can hand on.

Before this, an answer could only leave the browser as clipboard text (losing the
citations, the page anchors and any table) or, for a SQL answer, as a CSV of its
preview rows.

Two of these tests are about honesty rather than formatting, and both come from
mistakes this codebase already made:

  * A restored conversation keeps only the first 20 preview rows and no total
    (chat_orchestrator persists `preview_rows`; the count is recomputed as their
    length on reload, so a 4,000-row query comes back claiming 20). The document
    must not launder that into a row count it does not have.
  * When a cap bites, the document says so. A silently truncated export reads as
    complete — the same defect as the chronology export's old 5,000-row ceiling.
"""

import io
import sys
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.answer_docx import (MAX_COLS, MAX_ROWS, answer_filename,  # noqa: E402
                             build_answer_docx)


def _read(blob: bytes):
    from docx import Document
    return Document(io.BytesIO(blob))


def _text(doc) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


QUESTION = "What caused the delay to construction of the depot?"
ANSWER = """The delay was caused by several factors:

1.  **Water main:** access restricted from the planned start date
    of 01/08/08 until 18/02/09.
2.  **MUDFA works:** outstanding diversions in the same area.

| Cause | Period |
|---|---|
| Water main | 01/08/08 – 18/02/09 |
| MUDFA | ongoing |
"""
CITATIONS = [
    {"doc_name": "CEC00443401_PART3.pdf", "anchor": "page_1",
     "snippet": "a water main restricted access"},
    {"doc_name": "WED00000533.pdf", "anchor": "page_25", "snippet": ""},
]


class TestTheDocumentCarriesWhatMakesItCheckable:
    def test_question_answer_and_stamp(self):
        doc = _read(build_answer_docx(QUESTION, ANSWER))
        text = _text(doc)
        assert QUESTION in text
        assert "The delay was caused by several factors" in text
        assert "Exported from COAir on" in text

    def test_citations_keep_their_page_anchor(self):
        """The anchor is what makes a citation checkable, so it travels."""
        doc = _read(build_answer_docx(QUESTION, ANSWER, citations=CITATIONS))
        text = _text(doc)
        assert "SOURCES · 2" in text      # label() draws tracked-out caps
        assert "CEC00443401_PART3.pdf  p.1" in text
        assert "WED00000533.pdf  p.25" in text
        assert "a water main restricted access" in text

    def test_a_markdown_pipe_table_becomes_a_real_table(self):
        """The chat itself cannot draw this — no remark-gfm — so the file is
        better than the screen here, not a lossy copy of it."""
        doc = _read(build_answer_docx(QUESTION, ANSWER))
        assert len(doc.tables) == 1
        head = [c.text for c in doc.tables[0].rows[0].cells]
        assert head == ["CAUSE", "PERIOD"]
        assert [c.text for c in doc.tables[0].rows[1].cells] == [
            "Water main", "01/08/08 – 18/02/09"]

    def test_sql_and_its_rows_render_as_a_table(self):
        doc = _read(build_answer_docx(
            "How many workers by trade?", "See the table below.",
            sql="SELECT trade, n FROM manpower",
            table_columns=["trade", "n"],
            table_rows=[["Joiner", 12], ["Electrician", 7]]))
        text = _text(doc)
        assert "SELECT trade, n FROM manpower" in text
        assert len(doc.tables) == 1
        assert [c.text for c in doc.tables[0].rows[0].cells] == ["TRADE", "N"]
        assert [c.text for c in doc.tables[0].rows[1].cells] == ["Joiner", "12"]

    def test_an_empty_answer_does_not_crash(self):
        doc = _read(build_answer_docx(QUESTION, ""))
        assert "This answer had no text." in _text(doc)

    def test_no_question_is_fine(self):
        assert build_answer_docx("", "Just an answer.")


class TestItNeverClaimsARowCountItDoesNotHave:
    def test_without_a_total_it_says_what_it_retained(self):
        """A restored answer has 20 rows and no memory of the query's total."""
        doc = _read(build_answer_docx(
            "q", "a", table_columns=["a"], table_rows=[[i] for i in range(20)]))
        text = _text(doc)
        assert "Showing the first 20 rows this answer retained." in text
        assert "of 20 rows" not in text          # no invented total

    def test_with_a_known_total_it_states_it(self):
        doc = _read(build_answer_docx(
            "q", "a", table_columns=["a"], table_rows=[[i] for i in range(20)],
            total_rows=4000))
        assert "Showing 20 of 4000 rows." in _text(doc)

    def test_a_total_equal_to_what_is_shown_is_not_announced(self):
        doc = _read(build_answer_docx(
            "q", "a", table_columns=["a"], table_rows=[[1], [2]], total_rows=2))
        assert "Showing" not in _text(doc)

    def test_one_row_reads_singular(self):
        doc = _read(build_answer_docx("q", "a", table_columns=["a"],
                                      table_rows=[[1]]))
        assert "the first 1 row this answer retained" in _text(doc)


class TestCapsAreStatedNotSilent:
    def test_too_many_rows_says_so(self):
        doc = _read(build_answer_docx(
            "q", "a", table_columns=["a"],
            table_rows=[[i] for i in range(MAX_ROWS + 50)]))
        text = _text(doc)
        assert f"only the first {MAX_ROWS} rows are included" in text
        assert len(doc.tables[0].rows) == MAX_ROWS + 1     # + header

    def test_too_many_columns_says_so(self):
        cols = [f"c{i}" for i in range(MAX_COLS + 10)]
        doc = _read(build_answer_docx("q", "a", table_columns=cols,
                                      table_rows=[list(range(len(cols)))]))
        assert f"only the first {MAX_COLS} columns are included" in _text(doc)
        assert len(doc.tables[0].columns) == MAX_COLS

    def test_a_very_long_answer_says_it_was_cut(self):
        doc = _read(build_answer_docx("q", "A" * 300_000))
        assert "longer than this export carries" in _text(doc)

    def test_a_short_answer_carries_no_note(self):
        assert "Note —" not in _text(_read(build_answer_docx(QUESTION, ANSWER)))


class TestFilename:
    def test_it_is_taken_from_the_question(self):
        assert answer_filename(QUESTION) == (
            "COAir-answer-What-caused-the-delay-to-construction-of-the-depot.docx")

    def test_a_blank_question_still_yields_a_name(self):
        assert answer_filename("").endswith(".docx")
        assert answer_filename("   ") == "COAir-answer.docx"

    def test_path_separators_cannot_escape(self):
        assert "/" not in answer_filename("../../etc/passwd")
