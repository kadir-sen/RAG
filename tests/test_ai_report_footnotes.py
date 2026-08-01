import io
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document

from src.evidence_model import ChronologyEntry, EvidenceItem, VerifiedClaim
from src.report_docx import build_ai_chronology_docx


W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def test_ai_chronology_uses_real_claim_level_footnotes_and_canonical_layout():
    evidence = [
        EvidenceItem(
            source_id="src-doc", doc_id="DOC-17", file_name="notice.pdf",
            title="Notice of Delay", document_date="2025-01-07", page=4,
        ),
        EvidenceItem(
            source_id="src-mail", doc_id="MAIL-2", file_name="mail.msg",
            kind="email", sender="Engineer", recipient="Contractor",
            subject="Access", document_date="2025-01-08",
        ),
    ]
    entries = [ChronologyEntry(
        entry_ref="MODEL.MUST.NOT.CONTROL.NUMBERING",
        event_date="2025-01-08", date_precision="exact", parties=["Engineer"],
        event_type="notice", claims=[
            VerifiedClaim("The contractor issued notice.", ["src-doc"]),
            VerifiedClaim("The engineer replied the following day.", ["src-mail"]),
        ],
    )]

    blob, audit = build_ai_chronology_docx(
        project_name="Demo Project", issue_number=3, title="Access",
        entries=entries, evidence=evidence,
    )

    # A normal OOXML consumer can still open the package.
    doc = Document(io.BytesIO(blob))
    assert not doc.tables
    assert any(p.text.startswith("6.3.1  2025-01-08") for p in doc.paragraphs)
    assert all("MODEL.MUST" not in p.text for p in doc.paragraphs)
    section = doc.sections[0]
    assert abs(section.top_margin.cm - 2.54) < 0.02
    assert abs(section.left_margin.cm - 2.54) < 0.02

    with zipfile.ZipFile(io.BytesIO(blob)) as package:
        document_xml = ET.fromstring(package.read("word/document.xml"))
        footnotes_xml = ET.fromstring(package.read("word/footnotes.xml"))
        rels = package.read("word/_rels/document.xml.rels").decode()
        content_types = package.read("[Content_Types].xml").decode()
        styles = package.read("word/styles.xml").decode()

    refs = document_xml.findall(f".//{{{W}}}footnoteReference")
    records = [
        node for node in footnotes_xml.findall(f"{{{W}}}footnote")
        if int(node.attrib[f"{{{W}}}id"]) > 0
    ]
    note_text = " ".join("".join(node.itertext()) for node in records)
    assert len(refs) == len(records) == 2
    assert "Exhibit 1" in note_text and "p.4" in note_text
    assert "from Engineer" in note_text and "subject “Access”" in note_text
    assert "relationships/footnotes" in rels
    assert "/word/footnotes.xml" in content_types
    assert "FootnoteText" in styles and "FootnoteReference" in styles
    assert audit.footnote_references == audit.footnote_records == 2
    assert audit.unique_source_ids == 2
    assert audit.unresolved_source_ids == []


def test_unresolved_sources_are_reported_and_not_turned_into_fake_footnotes():
    blob, audit = build_ai_chronology_docx(
        project_name="Demo", issue_number=1, title="Missing source",
        entries=[ChronologyEntry(
            entry_ref="", event_date="2025-01-01", date_precision="exact",
            claims=[VerifiedClaim("A claim.", ["does-not-exist"])],
        )],
        evidence=[],
    )
    assert audit.unresolved_source_ids == ["does-not-exist"]
    assert audit.footnote_records == 0
    with zipfile.ZipFile(io.BytesIO(blob)) as package:
        assert "word/footnotes.xml" not in package.namelist()
