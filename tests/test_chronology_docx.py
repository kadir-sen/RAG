"""Chronology → .docx, and the date column that makes it a chronology.

Two things are pinned here:

  * the date lead parser, against the phrasings the authored files actually use
    (an earlier version took only a bare date after a short list of
    prepositions and left 22 of 80 entries showing "—");
  * the Word export, which has to carry the date beside its entry — that is the
    whole reason to export rather than copy out of the browser.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import pytest

from src.chronology_docx import (build_events_docx, build_subject_docx,
                                 describe_filters, safe_filename)
from src.chronology_library import _split_date


# ── the date column ─────────────────────────────────────────────────────
@pytest.mark.parametrize("text,when", [
    # already worked
    ("On 29 March 2005, tie issued a letter.", "29 March 2005"),
    ("In March 2007, the parties met.", "March 2007"),
    # qualifiers — the preposition is kept, because "late 2008" alone would
    # claim a point in time where the source gave a boundary
    ("By late 2008, after the Infraco Contract was let, BSC established its team.",
     "By late 2008"),
    ("From late 2009, tie's management produced disputes.", "From late 2009"),
    ("By the end of 2010, works had stopped.", "By the end of 2010"),
    ("In mid-2009, the board met.", "mid-2009"),
    # ranges and pairs
    ("In 2010–2011, tie issued a Remediable Termination Notice.", "2010–2011"),
    ("In June–July 2010, tie pursued an RTN strategy.", "June–July 2010"),
    ("In December 2008 and January 2009, the TPB considered options.",
     "December 2008 and January 2009"),
    ("Between June and December 2010, the report was presented.",
     "Between June and December 2010"),
    # a meeting that ran over consecutive days
    ("On 28–30 April 2008, the draft Report was reviewed.", "28–30 April 2008"),
    ("On 8–9 September 2011, the inquiry sat.", "8–9 September 2011"),
])
def test_date_leads_are_recognised(text, when):
    assert _split_date(text)[0] == when


@pytest.mark.parametrize("text", [
    "In addition, the parties disagreed.",
    "On the other hand, nothing happened.",
    "Under the procurement of the Edinburgh Tram Network, tie engaged PB.",
    "The strategy for delivering the tram rested on transferring risk.",
    "Throughout, the incomplete diversions were intertwined with construction.",
])
def test_context_openers_are_not_mistaken_for_dates(text):
    """A sentence that merely starts with a preposition is not a dated entry —
    the date column must stay empty rather than invent a period."""
    assert _split_date(text) == ("", text)


def test_the_sentence_survives_and_keeps_its_lowercase_names():
    when, rest = _split_date("On 29 March 2005, tie issued the SDS Contract.")
    assert when == "29 March 2005"
    # "tie" is lowercase by design in these documents; recapitalising is wrong
    assert rest == "tie issued the SDS Contract."


def test_the_real_files_are_mostly_dated_now():
    """Guards the parser against a regression that would quietly empty the date
    column. Measured: 58/80 before the range and qualifier support, 71/80 after."""
    from src.chronology_library import get_entries, list_docs

    entries = [e for d in list_docs() for e in get_entries(d.ref)]
    dated = [e for e in entries if (e.get("date") or "").strip()]
    assert len(entries) >= 70
    assert len(dated) >= 69, f"only {len(dated)}/{len(entries)} entries carry a date"


# ── the export ──────────────────────────────────────────────────────────
def _read(blob: bytes):
    import io

    from docx import Document

    return Document(io.BytesIO(blob))


_SUBJECT = {"ref": "03", "title": "Utility Diversion Failures (MUDFA)",
            "summary": "Scope growth and the diversion programme."}
_ENTRIES = [
    {"ref": "6.3.1", "date": "", "text": "MUDFA was the contract by which tie procured diversions.", "sub": []},
    {"ref": "6.3.2", "date": "28 March 2006", "text": "DLA Piper produced a report.",
     "sub": ["i) scope was not fixed", "ii) apparatus was unknown"]},
]


def test_subject_export_puts_the_date_beside_its_entry():
    doc = _read(build_subject_docx(_SUBJECT, _ENTRIES, "Edinburgh Tram Network"))

    heads = [p.text for p in doc.paragraphs[:3]]
    assert "CHRONOLOGY · 03" in heads
    assert _SUBJECT["title"] in heads

    table = doc.tables[0]
    assert len(table.rows) == len(_ENTRIES)
    # left cell carries date AND ref, so an entry stays citable in either medium
    assert "28 March 2006" in table.rows[1].cells[0].text
    assert "6.3.2" in table.rows[1].cells[0].text
    # an undated entry shows the dash rather than borrowing its neighbour's date
    assert "—" in table.rows[0].cells[0].text
    # sub-points travel with their entry
    assert "scope was not fixed" in table.rows[1].cells[1].text


def test_subject_export_states_its_provenance():
    doc = _read(build_subject_docx(_SUBJECT, _ENTRIES, "Edinburgh Tram Network"))
    tail = " ".join(p.text for p in doc.paragraphs[-3:])
    assert "2 entries" in tail
    assert "Exported from COAir" in tail
    assert "Edinburgh Tram Network" in tail


def test_subject_export_survives_an_empty_chronology():
    doc = _read(build_subject_docx(_SUBJECT, [], ""))
    assert any("no readable entries" in p.text for p in doc.paragraphs)


_ROWS = [
    {"date": "1970 to 1975", "event_type": "milestone", "actor": "Alex Macaulay",
     "description": "Chartered Engineer at Stirling County Council", "file_name": "CVS00000004.pdf"},
    {"date": "28 March 2006", "event_type": "delay", "actor": "tie",
     "description": "MUDFA works delayed", "file_name": "X.pdf"},
]


def test_event_export_keeps_free_form_dates_whole():
    """Laid out as a hanging indent rather than a table — date, tab, event —
    which is what makes exporting the whole 28k-event record affordable."""
    text = "\n".join(p.text for p in _read(build_events_docx(_ROWS, {})).paragraphs)
    # "1970 to 1975" must not be trimmed to a year: the store keeps what the
    # document said, and so does the export
    assert "1970 to 1975\tChartered Engineer at Stirling County Council" in text
    assert "28 March 2006\tMUDFA works delayed" in text
    # type, party and source travel under their entry
    assert "milestone · Alex Macaulay · CVS00000004.pdf" in text


def test_event_export_has_no_row_ceiling():
    """The ceiling was removed on purpose: a document that stops at a round
    number and tells the reader to narrow their filters is refusing to hand over
    the record while sounding helpful about it."""
    many = [dict(_ROWS[0], description=f"event {i}") for i in range(6000)]
    text = "\n".join(p.text for p in _read(build_events_docx(many, {})).paragraphs)
    assert "event 5999" in text                  # past the old 5,000 cap
    assert "6000 events on file" in text
    assert "6000 events · end of chronology" in text
    # and nothing that reads like an apology for a missing remainder
    assert "narrow the filters" not in text.lower()
    assert "continues beyond" not in text.lower()


def test_a_filtered_export_says_it_is_filtered():
    """An extract that doesn't state its filters reads as the whole record."""
    doc = _read(build_events_docx(_ROWS, {"event_type": "delay", "actor": "tie"}))
    text = " ".join(p.text for p in doc.paragraphs)
    assert "Filters" in text and "type = delay" in text
    assert "matching the filters" in text

    plain = _read(build_events_docx(_ROWS, {}))
    plain_text = " ".join(p.text for p in plain.paragraphs)
    assert "Filters" not in plain_text and "on file" in plain_text


def test_describe_filters():
    assert describe_filters({}) == ""
    assert describe_filters({"event_type": "delay"}) == "type = delay"
    assert "from 2006-01-01" in describe_filters({"date_from": "2006-01-01"})


def test_event_export_survives_an_empty_list():
    doc = _read(build_events_docx([], {"actor": "nobody"}))
    assert any("No events match" in p.text for p in doc.paragraphs)


@pytest.mark.parametrize("parts,expected", [
    (("Chronology", "03", "Utility Diversion Failures (MUDFA)"),
     "Chronology-03-Utility-Diversion-Failures-MUDFA"),
    (("Chronology", "project-record"), "Chronology-project-record"),
    # The all-empty fallback moved from "chronology" to "coair-export" when this
    # helper moved into src/docx_kit and became shared with the chat-answer
    # export — a file named "chronology.docx" would be wrong for half its callers
    # now. Unreachable from the chronology path either way: ref and title are
    # always present there.
    (("", ""), "coair-export"),
    (("../../etc/passwd",), "etc-passwd"),
])
def test_safe_filename(parts, expected):
    assert safe_filename(*parts) == expected
