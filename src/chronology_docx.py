"""Chronology → .docx.

The chronology area is something people read and then have to hand on — to a
solicitor, into a bundle, as an appendix. Copying it out of a browser loses the
one thing that makes it a chronology, which is the date set against the entry,
so the export is built to the same shape as the page: a date column on the left,
the entry beside it, sub-points indented under their entry.

Two documents, because the area shows two things:

  * an authored chronology (subject narrative) — numbered entries, refs kept, so
    a reader can cite 6.3.4 and it means the same thing in both places;
  * the event list — what the corpus produced, with its filters stated at the
    top so an extract can never be mistaken for the whole record.

Everything here is deterministic string-to-document work. No LLM, same as the
rest of the chronology layer.
"""

import io
import re
from datetime import datetime
from typing import Any, Dict, List, Optional

# python-docx is already a dependency (the authored chronologies are parsed from
# .docx). Imported lazily inside the builders so a missing optional install can
# never take the whole chronology API down with it.



# The layout primitives moved to docx_kit so the chat-answer export can reach
# them without importing a module named "chronology". Aliased back to their old
# private names: this file's two builders and tests/test_chronology_docx.py are
# written against them, and renaming would be churn with no reader benefit.
from .docx_kit import (  # noqa: F401
    ACCENT as _ACCENT,
    MUTED as _MUTED,
    RULE as _RULE,
    body as _body,
    borderless as _borderless,
    fit as _fit,
    footer_note as _footer_note,
    label as _label,
    mono as _mono,
    new_document as _new_document,
    paragraph_style_ids,
    rule as _rule,
    safe_filename,
    save as _save,
    shared as _shared,
    stamp as _stamp,
    title as _title,
)


# ── the authored chronology ─────────────────────────────────────────────
def build_subject_docx(subject: Dict[str, Any], entries: List[Dict[str, Any]],
                       collection: str = "",
                       evidence: Optional[Dict[str, Any]] = None) -> bytes:
    """One authored chronology, laid out as the page lays it out."""
    Pt, RGBColor, _ = _shared()
    doc = _new_document()

    _label(doc, f"Chronology · {subject.get('ref', '')}")
    _title(doc, str(subject.get("title") or "Chronology"))
    if subject.get("summary"):
        _body(doc, str(subject["summary"]), size=10, colour="4B5563")
    _rule(doc)

    if not entries:
        _body(doc, "This chronology has no readable entries.", colour=_MUTED)
        _stamp(doc, collection)
        return _save(doc)

    table = doc.add_table(rows=0, cols=2)
    _borderless(table)
    for e in entries:
        row = table.add_row().cells

        # Left: the date, and the reference beneath it — the same pairing the
        # page shows, so "6.3.4" cites identically in either medium.
        left = row[0].paragraphs[0]
        left.paragraph_format.space_after = Pt(0)
        _mono(left.add_run(str(e.get("date") or "—")), size=9,
              colour="111827", bold=True)
        ref_p = row[0].add_paragraph()
        ref_p.paragraph_format.space_before = Pt(1)
        _mono(ref_p.add_run(str(e.get("ref") or "")), size=8, colour=_MUTED)

        right = row[1].paragraphs[0]
        right.paragraph_format.space_after = Pt(0)
        run = right.add_run(str(e.get("text") or ""))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string("111827")

        for sub in e.get("sub") or []:
            sp = row[1].add_paragraph(style="List Bullet")
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after = Pt(0)
            sr = sp.add_run(str(sub))
            sr.font.size = Pt(9)
            sr.font.color.rgb = RGBColor.from_string("4B5563")

        # Breathing room between entries, matching the page's row padding.
        row[1].paragraphs[-1].paragraph_format.space_after = Pt(8)

    _fit(table, [3.6, 12.6])
    _rule(doc)
    n = len(entries)
    _footer_note(doc, f"{n} entr{'y' if n == 1 else 'ies'} · end of chronology")

    # The supporting documents, when an evidence pass has resolved them. The
    # narrative names no source files, so without this the exported chronology
    # asserts a history with nothing a reader can check it against.
    if evidence and evidence.get("documents"):
        _rule(doc)
        docs = evidence["documents"]
        _label(doc, f"Supporting documents · {len(docs)}")
        _body(doc,
              "Resolved from the project record for this subject — ranked by the "
              "search, not authored alongside the narrative.",
              size=9, colour=_MUTED)
        Pt, _, _ = _shared()
        for d in docs:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            page = d.get("page_number") or 1
            _mono(p.add_run(f"{d.get('file_name', '')}  p.{page}"),
                  size=9, colour="111827")
        note = (f"{evidence.get('elapsed_ms', 0)} ms · "
                f"{evidence.get('passes', 0)} searches · "
                f"{evidence.get('units_searched', 0)} entries · "
                f"{evidence.get('corpus_searched', 0)} documents searched")
        if evidence.get("budget_exhausted"):
            note += " · budget reached, later entries not searched"
        _footer_note(doc, note)

    _stamp(doc, collection)
    return _save(doc)


# ── the extracted event list ────────────────────────────────────────────
def describe_filters(f: Dict[str, Optional[str]]) -> str:
    """Human-readable filter line, or "" when nothing was narrowed.

    Stated at the top of the export on purpose: a filtered extract that does not
    say so reads as the complete record.
    """
    bits = []
    if f.get("event_type"):
        bits.append(f"type = {f['event_type']}")
    if f.get("actor"):
        bits.append(f"party contains \"{f['actor']}\"")
    if f.get("date_from"):
        bits.append(f"from {f['date_from']}")
    if f.get("date_to"):
        bits.append(f"to {f['date_to']}")
    return " · ".join(bits)


_DATE_COL_CM = 3.4


def _entry_styles(doc):
    """Two paragraph styles, defined once and reused by every event.

    This is the difference between an export that finishes and one that times
    out. Setting indents, a tab stop and fonts on each paragraph individually
    rebuilds the same XML tens of thousands of times.

    Returns style *ids*, not style objects, and the caller writes them straight
    onto the paragraph. Assigning `paragraph.style = <style>` looks like the
    obvious way and is a trap at this size: the setter resolves the id through
    `styles.default_for()`, which walks every style in the document on every
    call. Profiling 5,000 events put 67% of the run inside that lookup alone.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT

    Pt, RGBColor, Cm = _shared()

    def entry(style):
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor.from_string("111827")
        pf = style.paragraph_format
        pf.left_indent = Cm(_DATE_COL_CM)
        pf.first_line_indent = Cm(-_DATE_COL_CM)   # the date hangs in its column
        pf.space_after = Pt(1)
        pf.tab_stops.add_tab_stop(Cm(_DATE_COL_CM), WD_TAB_ALIGNMENT.LEFT)

    def meta(style):
        style.font.size = Pt(8)
        style.font.name = "Consolas"
        style.font.color.rgb = RGBColor.from_string(_MUTED)
        pf = style.paragraph_format
        pf.left_indent = Cm(_DATE_COL_CM)
        pf.space_after = Pt(6)

    ids = paragraph_style_ids(doc, {"ChronologyEntry": entry, "ChronologyMeta": meta})
    return ids["ChronologyEntry"], ids["ChronologyMeta"]


def build_events_docx(rows: List[Dict[str, Any]],
                      filters: Optional[Dict[str, Optional[str]]] = None) -> bytes:
    """The event list, whole, with whatever filters produced it stated up front.

    There is no row ceiling. There was one, and it was the wrong shape: a
    document that stops at a round number and tells the reader to narrow their
    filters is refusing to hand over the record while sounding helpful about it.

    Laid out as a hanging indent — date in a left column, event beside it —
    rather than a table, which is what makes "whole" affordable. The store holds
    ~28k events; as a Word table that took 82 seconds to build and opens badly
    in Word, where very large tables are slow to paginate. The same content as
    indented paragraphs takes 8 seconds and a third of the file size, and reads
    the same: date on the left, event on the right, exactly as on the page.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT

    Pt, RGBColor, Cm = _shared()
    doc = _new_document()
    described = describe_filters(filters or {})
    n = len(rows)

    _label(doc, "Chronology · project record")
    _title(doc, "Chronology")
    _body(doc, f"{n} event{'' if n == 1 else 's'}"
               f"{' matching the filters below' if described else ' on file'}",
          size=10, colour="4B5563")
    if described:
        p = doc.add_paragraph()
        _mono(p.add_run(f"Filters — {described}"), size=8, colour=_MUTED)
    _rule(doc)

    if not rows:
        _body(doc, "No events match these filters.", colour=_MUTED)
        _stamp(doc)
        return _save(doc)

    entry_id, meta_id = _entry_styles(doc)
    add = doc.add_paragraph          # bound once; called ~2x per event

    for r in rows:
        p = add()
        p._p.get_or_add_pPr().style = entry_id
        # Shown exactly as extracted — the store's dates are free-form because
        # the ingest kept whatever the document said ("1905", "1970 to 1975").
        # The date and the meta line are the only runs that need their own
        # formatting; the body takes the style's.
        _mono(p.add_run(f"{r.get('date') or '—'}\t"), size=8,
              colour="111827", bold=True)
        p.add_run(str(r.get("description") or r.get("reason") or ""))

        # Type, party and source under the entry, in the same muted register the
        # page uses. Carried on a line break inside the same paragraph rather
        # than a paragraph of its own: at 28k events that halves the paragraph
        # count, and the hanging indent keeps it aligned under the entry either
        # way.
        meta = " · ".join(str(x) for x in (
            r.get("event_type"), r.get("actor"), r.get("file_name")) if x)
        if meta:
            run = p.add_run()
            run.add_break()
            run.add_text(meta)
            _mono(run, size=8, colour=_MUTED)

    _rule(doc)
    _footer_note(doc, f"{n} event{'' if n == 1 else 's'} · end of chronology")
    _stamp(doc)
    return _save(doc)


