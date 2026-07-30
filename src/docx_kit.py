"""Shared Word-layout primitives.

These started life inside src/chronology_docx.py, which is where the first .docx
writer in this codebase was built. The chat-answer export needs exactly the same
page furniture — margins, the tracked-out caps label, the hairline rule, the
provenance stamp — and importing them from a module named "chronology" to render
a chat answer would be a lie about what the code is for. So they live here and
chronology_docx aliases them back under their old private names.

Everything is deliberately small and dependency-light: python-docx is imported
inside the functions, so a module that only wants `safe_filename` does not pay
for it.
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Callable, Dict

# COAir's ink. ACCENT is the orange used by the UI's --accent.
ACCENT = "E89517"
MUTED = "6B7280"
RULE = "D1D5DB"
INK = "111827"
BODY = "374151"


def shared():
    """python-docx's unit helpers, imported on demand."""
    from docx.shared import Pt, RGBColor, Cm
    return Pt, RGBColor, Cm


def mono(run, size: int = 9, colour: str = MUTED, bold: bool = False):
    Pt, RGBColor, _ = shared()
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(colour)


def label(doc, text: str):
    """The small tracked-out caps the page uses above a heading."""
    Pt, _, _ = shared()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    mono(p.add_run(text.upper()), size=8, colour=MUTED)
    return p


def rule(doc):
    """A hairline, drawn as a bottom border on an empty paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    Pt, _, _ = shared()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)
    pPr.append(borders)
    return p


def borderless(table):
    """python-docx has no 'no borders' guarantee across templates, so the table is
    told explicitly. Without this the export picks up whatever the default
    template defines and a narrative reads as a spreadsheet."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    tblPr.append(borders)


def fit(table, widths_cm):
    """Column widths only stick when set on every cell."""
    _, _, Cm = shared()
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def footer_note(doc, text: str):
    Pt, _, _ = shared()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    mono(p.add_run(text), size=8, colour=MUTED)
    return p


def stamp(doc, collection: str = ""):
    """Who produced this and when. An undated extract that turns up in a bundle
    six months later is worth very little."""
    when = datetime.now().strftime("%d %B %Y")
    tail = f" · {collection}" if collection else ""
    return footer_note(doc, f"Exported from COAir on {when}{tail}")


def new_document():
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    return doc


def title(doc, text: str):
    Pt, RGBColor, _ = shared()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    return p


def body(doc, text: str, size: int = 10, colour: str = BODY):
    Pt, RGBColor, _ = shared()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(colour)
    return p


def save(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def paragraph_style_ids(doc, specs: Dict[str, Callable]) -> Dict[str, str]:
    """Define paragraph styles once and return their *ids*.

    Ids, not style objects, because the caller writes them straight onto the
    paragraph::

        p = doc.add_paragraph()
        p._p.get_or_add_pPr().style = ids["entry"]

    `paragraph.style = <style>` looks like the obvious way and is a trap at size:
    the setter resolves the id through `styles.default_for()`, which walks every
    style in the document on every call. Profiling the 28k-event chronology export
    put 67% of the run inside that lookup alone. A chat answer is a few hundred
    paragraphs and would not notice, but both builders should be reaching for the
    same kit rather than one of them keeping a private fast path.

    `specs` maps a name to a callable that receives the freshly created style.
    """
    from docx.enum.style import WD_STYLE_TYPE

    styles = doc.styles
    out: Dict[str, str] = {}
    for name, configure in specs.items():
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        configure(style)
        out[name] = styles.get_style_id(style, WD_STYLE_TYPE.PARAGRAPH)
    return out


_UNSAFE = re.compile(r"[^A-Za-z0-9]+")


def safe_filename(*parts: str) -> str:
    """A predictable, filesystem-safe name. Callers add the extension."""
    joined = "-".join(_UNSAFE.sub("-", str(p or "")).strip("-") for p in parts if p)
    joined = re.sub(r"-{2,}", "-", joined).strip("-")
    return joined or "coair-export"
