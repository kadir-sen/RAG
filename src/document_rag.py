"""
Document RAG module with proper page-level citations.
Each PDF page is stored as a separate document with metadata.
Supports OCR for scanned PDFs.
"""
import os
import hashlib
from pathlib import Path
from typing import List, Optional, Dict, Any

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from pinecone import Pinecone, ServerlessSpec

from llama_index.core import (
    VectorStoreIndex,
    Document,
    StorageContext,
    Settings,
)
from llama_index.core.node_parser import SentenceSplitter
from llama_index.vector_stores.pinecone import PineconeVectorStore

from .config import (
    GOOGLE_API_KEY,
    PINECONE_API_KEY,
    GEMINI_MODEL,
    GEMINI_MODEL_LITE,
    ENABLE_LITE_TIER,
    EMBEDDING_MODEL,
    EMBEDDING_DIMENSION,
    EMBEDDING_PROVIDER,
    LOCAL_EMBEDDING_MODEL,
    EMBEDDING_DEVICE,
    PINECONE_INDEX_NAME,
    PINECONE_DIMENSION,
    VECTOR_STORE_BACKEND,
    QDRANT_URL,
    QDRANT_API_KEY,
    QDRANT_COLLECTION,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    OCR_MODE,
    OCR_LANG,
)
from .logger import logger, log_document_processing, log_pinecone, log_llm, log_separator, log_ocr_summary
from .ocr import get_ocr_pipeline, OCRPipeline


def generate_doc_id(file_path: str) -> str:
    """Generate a stable document ID from file path."""
    return hashlib.md5(file_path.encode()).hexdigest()[:16]


from contextvars import ContextVar

# Set by the chat orchestrator (in the request's async task, BEFORE the query is
# dispatched to a worker thread) so it reliably propagates via asyncio.to_thread.
# We don't read backend's auth contextvar directly because FastAPI runs the sync
# auth dependency in a separate threadpool whose context isn't the request task's,
# so that value doesn't reach the query thread.
corpus_var: ContextVar[str] = ContextVar("corpus", default="")


def _current_user_corpus() -> str:
    """The active request user's document corpus ('edinburgh' | 'demo' | '').
    Isolates retrieval per account: admin2 → edinburgh, admin → demo. Empty for
    scripts / unauthenticated work → no corpus filter (unchanged behaviour)."""
    try:
        return (corpus_var.get() or "").lower()
    except Exception:
        return ""


_EDIN_NAMES_CACHE = {"set": None, "ts": 0.0}


def edinburgh_filenames():
    """Cached set of file_names in the bulk (edinburgh) corpus — i.e. the distinct
    file names mirrored in the chunk store. Used as a corpus allow-list to drop any
    demo source (light_graph notices / SQL tables) that leaks into an edinburgh
    user's citations or related docs. 5-min TTL; empty set on failure (= no filter)."""
    import time
    now = time.time()
    c = _EDIN_NAMES_CACHE
    if c["set"] is not None and (now - c["ts"]) < 300:
        return c["set"]
    s = set()
    try:
        from .chunk_store import get_chunk_store
        con = get_chunk_store().connection()
        rows = con.execute(
            "SELECT DISTINCT file_name FROM chunks WHERE file_name IS NOT NULL "
            "AND file_name <> ''"
        ).fetchall()
        s = {r[0] for r in rows if r[0]}
    except Exception:
        pass
    c["set"], c["ts"] = s, now
    return s



class DocumentRAG:
    """
    Handles document indexing and retrieval using Pinecone.
    Each PDF page is stored separately with proper metadata for citations.
    """

    def __init__(self):
        """Initialize the Document RAG system."""
        log_separator("Initializing Document RAG")
        # Backend handle attributes (one of the two will remain None depending on backend).
        self.backend: str = VECTOR_STORE_BACKEND
        self.pinecone_index = None
        self.qdrant_client = None
        self._setup_llm()
        self._setup_vector_store()
        self.index: Optional[VectorStoreIndex] = None
        self.documents: List[Document] = []
        self.file_registry: Dict[str, Dict[str, Any]] = {}  # Track indexed files
        logger.info(f"✅ Document RAG initialized successfully (backend={self.backend})")

    DOCUMENT_SYSTEM_PROMPT = (
        "You are a construction project document analyst for a project management intelligence system. "
        "Answer questions based ONLY on the provided document excerpts.\n\n"
        "RULES:\n"
        "1. Always cite the specific document name and page number when referencing information.\n"
        "2. If the information is NOT in the provided excerpts, say so explicitly — do NOT guess or fabricate.\n"
        "3. Use construction industry terminology accurately (BOQ, IPC, EOT, VO, RFI, etc.).\n"
        "4. For numerical questions (quantities, counts, measurements), only state a number "
        "that appears verbatim in the excerpts, and name the document it came from "
        "(e.g. \"1,048 cameras, per Security Control Room Report p.3\"). NEVER invent or "
        "guess a figure to satisfy a request for an 'exact number'.\n"
        "5. If a total/count is NOT written as a single figure in one place and you must add up "
        "or infer it across documents, say so explicitly: prefix with 'approximately', show which "
        "documents you combined, and note it is an estimate — do not present an inferred total as exact.\n"
        "6. For yes/no questions, provide the answer first, then supporting evidence.\n"
        "7. For contract clause questions, quote the exact clause text when available.\n"
        "8. When multiple documents discuss the same topic, synthesize across all sources.\n"
        "9. Always answer in English with professional tone suitable for a project manager.\n"
        "10. Highlight any contradictions or discrepancies between different document sources.\n"
        "11. If a question relates to data/numbers that would be in Excel tables (manpower counts, "
        "equipment hours, progress percentages), note that the answer may be more accurately found "
        "in the project data tables rather than document text."
    )

    @staticmethod
    def _build_embed_model():
        """Build the embedding model. Local (free, sentence-transformers) when
        EMBEDDING_PROVIDER=local, else Gemini (cloud). The same model must embed
        both documents and queries, so this single factory is used everywhere via
        Settings.embed_model."""
        if EMBEDDING_PROVIDER == "api":
            # Hosted bge over an OpenAI-compatible endpoint — zero model RAM on
            # the server, wire-compatible with the fastembed/local bge corpus.
            from .config import EMBEDDING_API_URL, EMBEDDING_API_KEY
            from .api_embedding import ApiEmbedding
            log_llm("Setting up API embeddings (remote bge)", LOCAL_EMBEDDING_MODEL)
            model = ApiEmbedding(
                url=EMBEDDING_API_URL,
                api_key=EMBEDDING_API_KEY,
                model_name=LOCAL_EMBEDDING_MODEL,
            )
            dim = len(model.get_text_embedding("dimension probe"))
            if dim != EMBEDDING_DIMENSION:
                raise RuntimeError(
                    f"Embedding API model '{LOCAL_EMBEDDING_MODEL}' is {dim}-dim but "
                    f"EMBEDDING_DIMENSION={EMBEDDING_DIMENSION}.")
            return model

        if EMBEDDING_PROVIDER == "fastembed":
            # ONNX bge via fastembed — low RAM, no torch (fits a 2 GB server).
            # Wire-compatible with the sentence-transformers bge corpus.
            try:
                from .fastembed_embedding import FastEmbedEmbedding
            except ImportError as e:
                raise RuntimeError(
                    "EMBEDDING_PROVIDER=fastembed requires `fastembed`. "
                    "Install via `pip install fastembed`."
                ) from e
            log_llm("Setting up FASTEMBED (ONNX, low-RAM) embeddings", LOCAL_EMBEDDING_MODEL)
            model = FastEmbedEmbedding(model_name=LOCAL_EMBEDDING_MODEL)
            dim = len(model.get_text_embedding("dimension probe"))
            if dim != EMBEDDING_DIMENSION:
                raise RuntimeError(
                    f"fastembed model '{LOCAL_EMBEDDING_MODEL}' is {dim}-dim but "
                    f"EMBEDDING_DIMENSION={EMBEDDING_DIMENSION}.")
            return model

        if EMBEDDING_PROVIDER == "local":
            try:
                from llama_index.embeddings.huggingface import HuggingFaceEmbedding
            except ImportError as e:
                raise RuntimeError(
                    "EMBEDDING_PROVIDER=local requires `llama-index-embeddings-huggingface` "
                    "and `sentence-transformers`. Install via `pip install -r requirements.txt`."
                ) from e
            device = None if EMBEDDING_DEVICE == "auto" else EMBEDDING_DEVICE
            log_llm("Setting up LOCAL embeddings (free)", LOCAL_EMBEDDING_MODEL)
            # bge models recommend a query-side instruction; docs are embedded plain.
            model = HuggingFaceEmbedding(
                model_name=LOCAL_EMBEDDING_MODEL,
                device=device,
                query_instruction="Represent this sentence for searching relevant passages: ",
            )
            # Guard the #1 migration foot-gun: the embedding dimension MUST match
            # EMBEDDING_DIMENSION (and the vector-store collection). A mismatch
            # silently breaks retrieval (query vectors land in a different space).
            dim = len(model.get_text_embedding("dimension probe"))
            if dim != EMBEDDING_DIMENSION:
                raise RuntimeError(
                    f"LOCAL_EMBEDDING_MODEL '{LOCAL_EMBEDDING_MODEL}' produces {dim}-dim "
                    f"vectors but EMBEDDING_DIMENSION={EMBEDDING_DIMENSION}. Set "
                    f"EMBEDDING_DIMENSION={dim} (and recreate the vector-store collection) "
                    f"or pick a {EMBEDDING_DIMENSION}-dim model."
                )
            return model
        from llama_index.embeddings.google_genai import GoogleGenAIEmbedding
        return GoogleGenAIEmbedding(
            api_key=GOOGLE_API_KEY,
            model_name=EMBEDDING_MODEL,
            embedding_config={"output_dimensionality": EMBEDDING_DIMENSION},
        )

    def _setup_llm(self):
        """Configure LLM and embedding models."""
        log_llm("Setting up Gemini LLM and embeddings", GEMINI_MODEL)
        # Lazy + guarded: the LlamaIndex Gemini wrapper pulls the legacy
        # google.generativeai SDK, which is fragile on newer Pythons. It's only
        # needed for LlamaIndex query_engine synthesis; embeddings are independent
        # (and now local by default), and provider answers go through llm_client.
        try:
            from llama_index.llms.gemini import Gemini

            def _mk_gemini(m: str):
                return Gemini(api_key=GOOGLE_API_KEY, model=m,
                              system_prompt=self.DOCUMENT_SYSTEM_PROMPT)

            # Robust to the "Model names should start with `models/`" wrapper
            # check (see llm_client.create_llm). Getting a WORKING Settings.llm
            # here matters: when it's unset, LlamaIndex's dense query_engine
            # silently falls back to its OpenAI default — which, on a dead OpenAI
            # quota, just burns retries/timeouts.
            try:
                Settings.llm = _mk_gemini(GEMINI_MODEL)
            except Exception as e:
                if "models/" in str(e) and not GEMINI_MODEL.startswith(("models/", "tunedModels/")):
                    Settings.llm = _mk_gemini(f"models/{GEMINI_MODEL}")
                else:
                    raise
        except Exception as e:
            logger.warning(f"   Gemini LLM unavailable ({e}); Settings.llm unset "
                           f"(provider llm_client synthesis still works)")
        Settings.embed_model = self._build_embed_model()
        Settings.node_parser = SentenceSplitter(
            chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
        )
        logger.info(f"   Chunk size: {CHUNK_SIZE}, Overlap: {CHUNK_OVERLAP}")

    def _setup_vector_store(self):
        """Dispatch to the configured backend. Pinecone path is unchanged."""
        if self.backend == "qdrant":
            self._setup_qdrant()
        else:
            # Default + safety net: any unknown value falls back to Pinecone
            # so existing deployments keep working.
            self._setup_pinecone()

    def _setup_pinecone(self):
        """Initialize Pinecone vector store."""
        log_pinecone("Connecting to Pinecone...")
        self.pc = Pinecone(api_key=PINECONE_API_KEY)

        existing_indexes = [idx.name for idx in self.pc.list_indexes()]
        logger.info(f"   Existing indexes: {existing_indexes}")

        if PINECONE_INDEX_NAME not in existing_indexes:
            log_pinecone(f"Creating new index: {PINECONE_INDEX_NAME}")
            self.pc.create_index(
                name=PINECONE_INDEX_NAME,
                dimension=PINECONE_DIMENSION,
                metric="cosine",
                spec=ServerlessSpec(cloud="aws", region="us-east-1"),
            )
            logger.info(f"   Dimension: {PINECONE_DIMENSION}, Metric: cosine")
        else:
            log_pinecone(f"Using existing index: {PINECONE_INDEX_NAME}")

        self.pinecone_index = self.pc.Index(PINECONE_INDEX_NAME)
        self.vector_store = PineconeVectorStore(pinecone_index=self.pinecone_index)

        stats = self.pinecone_index.describe_index_stats()
        logger.info(f"   Total vectors: {stats.get('total_vector_count', 0)}")

    def _setup_qdrant(self):
        """Initialize Qdrant vector store. Lazy imports so the package is only
        required when this backend is actually selected."""
        try:
            from qdrant_client import QdrantClient
            from qdrant_client.http import models as qmodels
            from llama_index.vector_stores.qdrant import QdrantVectorStore
        except ImportError as e:
            raise RuntimeError(
                "VECTOR_STORE_BACKEND=qdrant requires `qdrant-client` and "
                "`llama-index-vector-stores-qdrant`. Install via "
                "`pip install -r requirements.txt`."
            ) from e

        logger.info(f"[Qdrant] Connecting to {QDRANT_URL} ...")
        self.qdrant_client = QdrantClient(
            url=QDRANT_URL,
            api_key=QDRANT_API_KEY or None,
            prefer_grpc=False,
            timeout=30.0,
            check_compatibility=False,  # tolerate minor client/server version skew
        )

        existing = {c.name for c in self.qdrant_client.get_collections().collections}
        logger.info(f"   Existing collections: {sorted(existing)}")

        if QDRANT_COLLECTION not in existing:
            logger.info(f"[Qdrant] Creating collection: {QDRANT_COLLECTION}")
            # On-disk vectors + int8 scalar quantization so a large corpus fits a
            # small (2 GB) box: original vectors live on disk, the int8-quantized
            # copy stays in RAM (~1/4 size) for fast search.
            try:
                self.qdrant_client.create_collection(
                    collection_name=QDRANT_COLLECTION,
                    vectors_config=qmodels.VectorParams(
                        size=EMBEDDING_DIMENSION,
                        distance=qmodels.Distance.COSINE,
                        on_disk=True,
                    ),
                    quantization_config=qmodels.ScalarQuantization(
                        scalar=qmodels.ScalarQuantizationConfig(
                            type=qmodels.ScalarType.INT8,
                            always_ram=True,
                        )
                    ),
                )
                logger.info(f"   Dimension: {EMBEDDING_DIMENSION}, Metric: cosine, on_disk+int8")
            except Exception as e:
                # Concurrent creation race (parallel ingest workers): another
                # process already created it — tolerate and proceed.
                if not self.qdrant_client.collection_exists(QDRANT_COLLECTION):
                    raise
                logger.info(f"[Qdrant] Collection created concurrently: {QDRANT_COLLECTION}")
        else:
            logger.info(f"[Qdrant] Using existing collection: {QDRANT_COLLECTION}")

        self.vector_store = QdrantVectorStore(
            client=self.qdrant_client,
            collection_name=QDRANT_COLLECTION,
        )

        try:
            info = self.qdrant_client.get_collection(QDRANT_COLLECTION)
            logger.info(f"   Total vectors: {info.points_count or 0}")
        except Exception as e:
            logger.warning(f"   Could not read collection info: {e}")

    def _delete_file_vectors(self, file_name: str):
        """Delete existing vectors for a file before re-indexing."""
        doc_id = generate_doc_id(file_name)
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                flt = qmodels.Filter(must=[
                    qmodels.FieldCondition(
                        key="doc_id",
                        match=qmodels.MatchValue(value=doc_id),
                    )
                ])
                self.qdrant_client.delete(
                    collection_name=QDRANT_COLLECTION,
                    points_selector=qmodels.FilterSelector(filter=flt),
                )
            else:
                # Pinecone metadata filter delete (unchanged behavior).
                self.pinecone_index.delete(
                    filter={"doc_id": {"$eq": doc_id}}
                )
            logger.info(f"   Cleared existing vectors for: {file_name}")
        except Exception as e:
            # Pinecone free tier might not support metadata filtering for delete
            logger.info(f"   Could not filter-delete (new file or free tier): {file_name}")

    def parse_pdf_by_pages(
        self,
        file_path: str,
        ocr_mode: Optional[str] = None,
        ocr_language: Optional[str] = None,
    ) -> List[Document]:
        """
        Parse PDF page-by-page with OCR support for scanned pages.
        Each page becomes a separate Document with metadata for citations.

        Args:
            file_path: Path to PDF file
            ocr_mode: Override OCR mode ("auto", "force", "off")
            ocr_language: Override OCR language ("eng", "eng+tur", etc.)

        Returns:
            List of Document objects, one per page
        """
        documents = []
        path = Path(file_path)
        file_name = path.name

        # Initialize OCR pipeline with settings
        ocr = get_ocr_pipeline(
            mode=ocr_mode or OCR_MODE,
            language=ocr_language or OCR_LANG,
        )

        try:
            logger.info(f"   Parsing PDF: {file_name}")
            logger.info(f"   OCR mode: {ocr.mode}, language: {ocr.language}")
            logger.info(f"   ----------------------------------------")

            # Use OCR pipeline for extraction
            page_results = ocr.extract_text_auto(file_path)
            total_pages = len(page_results)

            # Stats for summary
            native_count = 0
            ocr_count = 0
            failed_count = 0
            total_ocr_time = 0

            for page_text in page_results:
                text = page_text.text
                method = page_text.extraction_method

                # Track stats
                if method == "native":
                    native_count += 1
                elif method in ("ocr", "native+ocr"):
                    ocr_count += 1
                    if page_text.ocr_time_ms:
                        total_ocr_time += page_text.ocr_time_ms

                if not text or len(text) < 10:
                    failed_count += 1
                    continue

                # Build metadata
                page_metadata = {
                    "file_name": file_name,
                    "file_path": str(file_path),
                    "file_type": ".pdf",
                    "page_number": page_text.page_number,
                    "total_pages": total_pages,
                    "doc_id": generate_doc_id(file_path),
                    "extraction_method": method,
                }

                # Add OCR metadata if applicable
                if page_text.ocr_engine:
                    page_metadata["ocr_engine"] = page_text.ocr_engine
                if page_text.ocr_language:
                    page_metadata["ocr_language"] = page_text.ocr_language
                if page_text.ocr_confidence is not None:
                    page_metadata["ocr_confidence"] = page_text.ocr_confidence

                # Check for table-like content
                if ocr.detect_table_like_structure(text):
                    page_metadata["table_hint"] = True

                page_doc = Document(
                    text=text,
                    metadata=page_metadata,
                )
                documents.append(page_doc)

            # Log summary
            log_ocr_summary(
                total_pages=total_pages,
                native_pages=native_count,
                ocr_pages=ocr_count,
                failed_pages=failed_count,
                total_time_ms=total_ocr_time,
            )

            logger.info(f"   RESULT: {len(documents)} page documents created")

            if len(documents) == 0:
                logger.warning(f"   WARNING: No text extracted from PDF!")
                if ocr.mode == "off":
                    logger.warning(f"   TIP: Try enabling OCR mode for scanned PDFs")

        except Exception as e:
            logger.error(f"   Error parsing PDF: {e}")
            import traceback
            logger.error(traceback.format_exc())

        return documents

    def parse_docx(self, file_path: str) -> List[Document]:
        """Parse Word document into documents."""
        documents = []
        path = Path(file_path)
        file_name = path.name

        try:
            doc = DocxDocument(file_path)
            full_text = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            if full_text:
                documents.append(Document(
                    text=full_text,
                    metadata={
                        "file_name": file_name,
                        "file_path": str(file_path),
                        "file_type": ".docx",
                        "page_number": 1,
                        "total_pages": 1,
                        "doc_id": generate_doc_id(file_path),
                    },
                ))
            logger.info(f"   Parsed DOCX: {len(full_text)} characters")

        except Exception as e:
            logger.error(f"   Error parsing DOCX: {e}")

        return documents

    def parse_txt(self, file_path: str) -> List[Document]:
        """Parse text file into documents."""
        documents = []
        path = Path(file_path)
        file_name = path.name

        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read().strip()

            if text:
                documents.append(Document(
                    text=text,
                    metadata={
                        "file_name": file_name,
                        "file_path": str(file_path),
                        "file_type": ".txt",
                        "page_number": 1,
                        "total_pages": 1,
                        "doc_id": generate_doc_id(file_path),
                    },
                ))
            logger.info(f"   Parsed TXT: {len(text)} characters")

        except Exception as e:
            logger.error(f"   Error parsing TXT: {e}")

        return documents

    def add_document(
        self,
        file_path: str,
        ocr_mode: Optional[str] = None,
        ocr_language: Optional[str] = None,
    ) -> Optional[List[Document]]:
        """
        Parse a document and return new Document objects.

        Args:
            file_path: Path to document file
            ocr_mode: OCR mode ("auto", "force", "off") - only for PDFs
            ocr_language: OCR language ("eng", "eng+tur") - only for PDFs

        Returns:
            List of new Document objects, or None if failed
        """
        path = Path(file_path)
        file_name = path.name
        extension = path.suffix.lower()

        log_document_processing(file_name, "Processing", f"Type: {extension}")

        # Delete existing vectors for this file (prevent duplicates)
        self._delete_file_vectors(file_name)

        # Parse based on file type
        if extension == ".pdf":
            new_docs = self.parse_pdf_by_pages(file_path, ocr_mode, ocr_language)
        elif extension in [".docx", ".doc"]:
            new_docs = self.parse_docx(file_path)
        elif extension == ".txt":
            new_docs = self.parse_txt(file_path)
        else:
            logger.warning(f"   Unsupported file type: {extension}")
            return None

        if new_docs:
            self.documents.extend(new_docs)

            # Count OCR pages
            ocr_pages = sum(1 for d in new_docs if d.metadata.get("extraction_method") in ("ocr", "native+ocr"))

            self.file_registry[file_name] = {
                "file_path": str(file_path),
                "file_type": extension,
                "page_count": len(new_docs),
                "ocr_pages": ocr_pages,
                "doc_id": generate_doc_id(file_path),
            }
            log_document_processing(file_name, "Added", f"{len(new_docs)} pages ({ocr_pages} OCR)")
            return new_docs

        return None

    def add_document_from_pages(
        self,
        file_path: str,
        page_texts: Dict[int, str],
        metadata: Optional[Dict] = None,
    ) -> Optional[List[Document]]:
        """
        Create Document objects from pre-parsed page texts (e.g., email body).

        Args:
            file_path: Original file path (for metadata)
            page_texts: Dict mapping page_number -> text content
            metadata: Additional metadata to attach to each document

        Returns:
            List of Document objects, or None if empty
        """
        path = Path(file_path)
        file_name = path.name

        # Delete existing vectors for this file
        self._delete_file_vectors(file_name)

        new_docs = []
        for page_num, text in sorted(page_texts.items()):
            if not text or not text.strip():
                continue
            doc_meta = {
                "file_name": file_name,
                "file_path": str(file_path),
                "file_type": path.suffix.lower(),
                "page_number": page_num,
                "total_pages": len(page_texts),
                "doc_id": generate_doc_id(file_path),
            }
            if metadata:
                doc_meta.update(metadata)

            new_docs.append(Document(text=text.strip(), metadata=doc_meta))

        if new_docs:
            self.documents.extend(new_docs)
            self.file_registry[file_name] = {
                "file_path": str(file_path),
                "file_type": path.suffix.lower(),
                "page_count": len(new_docs),
                "ocr_pages": 0,
                "doc_id": generate_doc_id(file_path),
            }
            log_document_processing(file_name, "Added", f"{len(new_docs)} pages (from pages)")
            return new_docs

        return None

    def add_documents_from_folder(self, folder_path: str) -> int:
        """Add all supported documents from a folder."""
        count = 0
        supported = {".pdf", ".docx", ".doc", ".txt"}
        folder = Path(folder_path)

        if not folder.exists():
            logger.warning(f"Folder not found: {folder_path}")
            return 0

        log_separator(f"Scanning: {folder.name}")

        for file_path in folder.rglob("*"):
            if file_path.suffix.lower() in supported:
                if self.add_document(str(file_path)):
                    count += 1

        logger.info(f"📁 Added {count} documents from {folder.name}")
        return count

    def build_index(self) -> bool:
        """Build vector index from documents."""
        if not self.documents:
            logger.warning("No documents to index")
            return False

        log_separator("Building Vector Index")
        logger.info(f"Indexing {len(self.documents)} document(s)...")

        try:
            # Use default namespace for all documents (simpler, allows cross-file search)
            storage_context = StorageContext.from_defaults(vector_store=self.vector_store)

            # Build index from all documents
            self.index = VectorStoreIndex.from_documents(
                self.documents,
                storage_context=storage_context,
                show_progress=True
            )

            log_pinecone("Index built successfully")
            logger.info(f"   Indexed {len(self.documents)} documents")
            return True
        except Exception as e:
            logger.error(f"Error building index: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False

    def insert_documents(self, new_docs: List[Document]) -> bool:
        """
        Add new documents to existing Pinecone index incrementally.
        Uses batch embedding + batch upsert for performance.
        """
        if not new_docs:
            return False

        # Ensure index exists (load from Pinecone or build fresh)
        if not self.index:
            if not self.load_index():
                # No existing index - need full build for the first time
                return self.build_index()

        log_separator("Incremental Index Update")
        logger.info(f"Inserting {len(new_docs)} new document(s) into existing index...")

        try:
            from llama_index.core.schema import TextNode

            # 1. Parse documents into nodes using the configured node parser
            node_parser = Settings.node_parser
            nodes = []
            for doc in new_docs:
                parsed = node_parser.get_nodes_from_documents([doc])
                nodes.extend(parsed)

            if not nodes:
                logger.warning("   No nodes generated from documents")
                return False

            logger.info(f"   Generated {len(nodes)} nodes from {len(new_docs)} documents")

            # Mirror chunk text to the local chunk store for the lexical retrieval
            # lane (Phase 1). Non-fatal: retrieval degrades to dense-only on failure.
            try:
                from .chunk_store import get_chunk_store
                rows = []
                for node in nodes:
                    meta = node.metadata or {}
                    rows.append({
                        "doc_id": meta.get("doc_id", ""),
                        "file_name": meta.get("file_name", "Unknown"),
                        "page_number": meta.get("page_number", 1),
                        "text": node.get_content(),
                    })
                get_chunk_store().add_chunks(rows)
            except Exception as e:
                logger.warning(f"   Chunk store mirror skipped: {e}")

            # 2. Batch embed all nodes at once
            embed_model = Settings.embed_model
            texts = [node.get_content() for node in nodes]

            EMBED_BATCH = 50
            all_embeddings = []
            total_batches = (len(texts) + EMBED_BATCH - 1) // EMBED_BATCH
            for i in range(0, len(texts), EMBED_BATCH):
                batch = texts[i:i + EMBED_BATCH]
                batch_num = i // EMBED_BATCH + 1
                # Retry embedding with backoff (handles rate limits and transient errors)
                for attempt in range(3):
                    try:
                        batch_emb = embed_model.get_text_embedding_batch(batch)
                        all_embeddings.extend(batch_emb)
                        logger.info(f"   Embedded batch {batch_num}/{total_batches}")
                        break
                    except Exception as emb_err:
                        if attempt < 2:
                            import time
                            wait = 2 ** (attempt + 1)
                            logger.warning(f"   Embedding batch {batch_num} failed (attempt {attempt + 1}): {emb_err}. Retrying in {wait}s...")
                            time.sleep(wait)
                        else:
                            raise

            for node, emb in zip(nodes, all_embeddings):
                node.embedding = emb

            # 3. Batch upsert to Pinecone via vector store (with retry)
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    self.vector_store.add(nodes)
                    logger.info(f"   Inserted {len(nodes)} vectors (batch mode)")
                    return True
                except Exception as upsert_err:
                    logger.warning(f"   Batch upsert attempt {attempt + 1}/{max_retries} failed: {upsert_err}")
                    if attempt < max_retries - 1:
                        import time
                        wait = 2 ** (attempt + 1)
                        logger.info(f"   Retrying in {wait}s...")
                        time.sleep(wait)
                    else:
                        raise

        except Exception as e:
            logger.error(f"Error in batch insert: {e}")
            # Fallback to sequential insert with retry per document
            logger.info("Falling back to sequential insert...")
            inserted = 0
            for doc in new_docs:
                for attempt in range(3):
                    try:
                        self.index.insert(doc)
                        inserted += 1
                        break
                    except Exception as e2:
                        if attempt < 2:
                            import time
                            time.sleep(2 ** attempt)
                        else:
                            logger.error(f"Failed to insert doc after 3 attempts: {e2}")
            if inserted > 0:
                logger.info(f"   Sequential insert: {inserted}/{len(new_docs)} docs inserted")
                return True
            logger.error("All insert attempts failed")
            return False

    def embed_file_to_vectors(self, file_path: str) -> int:
        """Parse a file (with OCR), embed its chunks, and upsert ONLY to the vector
        store — NO chunk_store / registry / light_graph writes. This makes it safe
        to run from parallel worker processes (Qdrant upserts are concurrent-safe),
        unlike route_file()/insert_documents() which touch shared local state.

        Returns the number of chunks indexed (0 on parse failure)."""
        docs = self.add_document(file_path)  # parse + OCR + delete-existing-vectors
        # add_document accumulates into self.documents; clear it so long-lived
        # workers don't leak memory across thousands of files.
        self.documents.clear()
        if not docs:
            return 0
        node_parser = Settings.node_parser
        nodes = []
        for doc in docs:
            nodes.extend(node_parser.get_nodes_from_documents([doc]))
        if not nodes:
            return 0
        embed_model = Settings.embed_model
        texts = [n.get_content() for n in nodes]
        embeddings: List[List[float]] = []
        EMBED_BATCH = 50
        for i in range(0, len(texts), EMBED_BATCH):
            embeddings.extend(embed_model.get_text_embedding_batch(texts[i:i + EMBED_BATCH]))
        for n, e in zip(nodes, embeddings):
            n.embedding = e
        self.vector_store.add(nodes)
        return len(nodes)

    def load_index(self) -> bool:
        """Load existing index from the configured vector store backend."""
        try:
            logger.info(f"[{self.backend}] Loading existing index...")
            # self.vector_store is already initialised by _setup_vector_store().
            # Re-creating it would be a no-op for Pinecone but break for Qdrant.
            self.index = VectorStoreIndex.from_vector_store(self.vector_store)
            if self.backend == "qdrant":
                try:
                    info = self.qdrant_client.get_collection(QDRANT_COLLECTION)
                    logger.info(f"   Loaded {info.points_count or 0} vectors")
                except Exception:
                    pass
            else:
                stats = self.pinecone_index.describe_index_stats()
                logger.info(f"   Loaded {stats.get('total_vector_count', 0)} vectors")
            return True
        except Exception as e:
            logger.error(f"Could not load index: {e}")
            return False

    def query(
        self,
        question: str,
        top_k: int = 10,
        doc_ids: Optional[List[str]] = None,
        file_names: Optional[List[str]] = None,
        payload_filters: Optional[Dict[str, Any]] = None,
        synthesize: bool = True,
        rerank: bool = True,
    ) -> dict:
        """Query documents with proper page-level citations.

        rerank=False drops the LLM *rerank* call as well, leaving retrieval
        entirely model-free: dense + BM25 + RRF and nothing else. synthesize=False
        alone is not that — ENABLE_RERANK defaults on, so the retrieve-only path
        still spends one lite-model call ordering the fused pool. The chronology
        evidence build needs a retrieval it can promise has no model in it,
        because that promise is printed on the page, so it has to be a parameter
        rather than a hope about an env var.
        If doc_ids is provided, only search within those documents.
        payload_filters scope retrieval by scoped-metadata (e.g. {"doc_type":
        "delay notice"}); when a scoped query returns nothing we retry once
        unscoped so over-tight scope (or not-yet-backfilled payloads) never turns
        a real answer into "not found".

        synthesize=False is "retrieve-only": return the ranked chunks/sources
        WITHOUT spending an LLM synthesis call. Used by callers that re-synthesize
        downstream (the hybrid handler, the multi-step planner, and the ReAct
        agent's document tool) so the per-source answer isn't paid for and then
        discarded. ``answer`` is "" in that mode; ``sources`` are always populated.
        """
        log_separator("Document Query")
        logger.info(f"🔍 Question: {question[:100]}...")

        original_question = question
        try:
            from .jargon_manager import get_jargon_manager
            jm = get_jargon_manager()
            semantic_query = jm.replace_query_terms_with_meanings(question)
            if semantic_query and semantic_query != question:
                logger.info(f"   Jargon semantic query: {semantic_query[:100]}")
                question = semantic_query

            concept_terms = jm.expand_domain_concepts(question)
            if concept_terms:
                terms_to_add = [t for t in concept_terms if t.lower() not in question.lower()][:6]
                if terms_to_add:
                    question = f"{question} (related: {', '.join(terms_to_add)})"
                    logger.info(f"   Concept-boosted retrieval with {len(terms_to_add)} synonyms")
        except Exception as e:
            logger.debug(f"   Jargon normalization skipped: {e}")

        if not self.index:
            if not self.load_index():
                return {
                    "answer": "No documents indexed. Please upload documents first.",
                    "sources": [],
                }

        # ── Hybrid retrieval: dense + lexical → RRF → LLM rerank (Phase 1) ──
        # Flag-gated; on any failure or no candidates we fall through to the
        # original pure-dense LlamaIndex path below (fully backward-compatible).
        from .config import ENABLE_HYBRID_RETRIEVAL
        if ENABLE_HYBRID_RETRIEVAL:
            try:
                hybrid = self._hybrid_query(
                    question=question,               # jargon/concept-expanded (dense)
                    raw_question=original_question,   # cleaner text for lexical + rerank
                    top_k=top_k, doc_ids=doc_ids, file_names=file_names,
                    payload_filters=payload_filters, synthesize=synthesize,
                    rerank=rerank,
                )
                # Empty-scope safety net: a scoped query that found nothing retries
                # once unscoped rather than reporting "not found".
                if hybrid is not None and not hybrid.get("sources") and payload_filters:
                    logger.info("   Scoped hybrid empty → retrying unscoped")
                    hybrid = self._hybrid_query(
                        question=question, raw_question=original_question,
                        top_k=top_k, doc_ids=doc_ids, file_names=file_names,
                        synthesize=synthesize, rerank=rerank,
                    )
                if hybrid is not None:
                    return hybrid
            except Exception as e:
                logger.warning(f"   Hybrid retrieval failed → dense fallback: {e}")

        logger.info(f"   Retrieving top {top_k} matches...")
        filters = self._build_metadata_filters(doc_ids, file_names, payload_filters)

        if not synthesize:
            # Retrieve-only fast path: use a retriever (no synthesizing query
            # engine, no LLM call). Mirrors the empty-scope corpus-only retry.
            retriever = (self.index.as_retriever(similarity_top_k=top_k, filters=filters)
                         if filters is not None
                         else self.index.as_retriever(similarity_top_k=top_k))
            nodes = list(retriever.retrieve(question))
            if payload_filters and not nodes:
                logger.info("   Scoped dense retrieve empty → retrying (corpus-only)")
                retry_filters = self._build_metadata_filters(doc_ids, file_names, None)
                retriever = (self.index.as_retriever(similarity_top_k=top_k, filters=retry_filters)
                             if retry_filters is not None
                             else self.index.as_retriever(similarity_top_k=top_k))
                nodes = list(retriever.retrieve(question))
            sources = self._sources_from_scored_nodes(nodes)
            self._report_reading(sources)
            logger.info(f"✅ Retrieve-only: {len(sources)} unique sources")
            return {"answer": "", "sources": sources}

        if filters is not None:
            query_engine = self.index.as_query_engine(
                similarity_top_k=top_k, filters=filters,
            )
            scope_desc = []
            if doc_ids:
                scope_desc.append(f"{len(doc_ids)} doc_ids")
            if file_names:
                scope_desc.append(f"{len(file_names)} file_names")
            if payload_filters:
                scope_desc.append(f"scope={list(payload_filters.keys())}")
            logger.info(f"   Scoped to {' + '.join(scope_desc)}")
        else:
            query_engine = self.index.as_query_engine(similarity_top_k=top_k)
        if Settings.llm is None:
            # No LlamaIndex LLM configured (e.g. Gemini SDK unavailable locally).
            # Provider-specific synthesis goes through query_with_provider/llm_client;
            # this default engine needs an LLM, so fail clearly instead of crashing deep.
            raise RuntimeError(
                "Default query engine needs Settings.llm but it is unset "
                "(Gemini LLM unavailable). Use query_with_provider() / the provider "
                "synthesis path, or run where the Gemini SDK works (prod Python 3.11)."
            )
        response = query_engine.query(question)
        # Empty-scope safety net (dense fallback path): retry if a scoped query
        # produced no source nodes — but KEEP the per-user corpus filter (drop only
        # the doc_type/project scope), so the retry never leaks another corpus.
        if payload_filters and not getattr(response, "source_nodes", None):
            logger.info("   Scoped dense query empty → retrying (corpus-only)")
            retry_filters = self._build_metadata_filters(doc_ids, file_names, None)
            engine = (self.index.as_query_engine(similarity_top_k=top_k, filters=retry_filters)
                      if retry_filters is not None
                      else self.index.as_query_engine(similarity_top_k=top_k))
            response = engine.query(question)

        sources = self._sources_from_scored_nodes(response.source_nodes)
        self._report_reading(sources)
        logger.info(f"✅ Found {len(sources)} unique sources")
        return {"answer": str(response), "sources": sources}

    @staticmethod
    def _report_reading(sources: List[dict]) -> None:
        """Publish 'reading <file>' activity steps for the live feed (best-effort,
        no-op outside a chat request). Deduped, capped so a wide retrieval doesn't
        flood the feed."""
        try:
            from backend.tasks.query_progress import report_step
            seen = set()
            for s in (sources or []):
                fn = s.get("file_name")
                if not fn or fn in seen:
                    continue
                seen.add(fn)
                report_step("reading", f"reading {fn}", f"p.{s.get('page_number', '?')}")
                if len(seen) >= 5:
                    break
        except Exception:
            pass

    @staticmethod
    def _sources_from_scored_nodes(scored_nodes) -> List[dict]:
        """Build deduped page-level source dicts from LlamaIndex NodeWithScore
        objects (from a query engine's source_nodes OR a retriever's results).
        Shared by the synthesizing dense path and the retrieve-only fast path."""
        sources: List[dict] = []
        seen = set()
        for i, node in enumerate(scored_nodes or [], 1):
            meta = node.metadata
            file_name = meta.get("file_name", "Unknown")
            file_path = meta.get("file_path", "")
            score = round(node.score, 3) if node.score else None

            try:
                page_num = int(meta.get("page_number", 1))
            except (ValueError, TypeError):
                page_num = 1
            try:
                total_pages = int(meta.get("total_pages", 1))
            except (ValueError, TypeError):
                total_pages = 1

            # Dedupe by file+page
            key = f"{file_name}_{page_num}"
            if key in seen:
                continue
            seen.add(key)

            # Extract highlight (first 2-3 sentences)
            text = node.text.strip()
            sentences = text.replace('\n', ' ').split('. ')
            highlight = '. '.join(sentences[:3])
            if len(sentences) > 3:
                highlight += '...'

            sources.append({
                "doc_id": meta.get("doc_id", generate_doc_id(file_path)) if file_path else "",
                "file_name": file_name,
                "file_path": file_path,
                "page_number": page_num,
                "total_pages": total_pages,
                "score": score,
                "text_snippet": text[:500] + "..." if len(text) > 500 else text,
                "highlight_text": highlight[:300],
            })
        return sources

    # ── Hybrid retrieval helpers (Phase 1) ──────────────────────────────
    def _hybrid_query(self, question: str, raw_question: str, top_k: int,
                      doc_ids: Optional[List[str]] = None,
                      file_names: Optional[List[str]] = None,
                      payload_filters: Optional[Dict[str, Any]] = None,
                      synthesize: bool = True,
                      rerank: bool = True) -> Optional[dict]:
        """Dense + lexical candidate fusion (RRF) + optional LLM rerank, then
        synthesize an answer from the final chunks. Returns None when there are
        no candidates at all (caller falls back to the dense path).

        payload_filters scope the DENSE lane (its payload carries doc_type/project);
        the lexical lane stays unscoped because the chunk store doesn't mirror those
        keys — fusion + rerank still surface the in-scope chunks, and an empty-result
        unfiltered retry in query() protects recall."""
        from .config import RAG_CANDIDATE_K, RAG_RERANK_K, RAG_FINAL_K, ENABLE_RERANK

        final_k = top_k or RAG_FINAL_K

        try:
            from backend.tasks.query_progress import report_step
            report_step("searching", "searching documents…")
        except Exception:
            pass

        # 1. Dense candidates (existing vector lane, widened pool)
        dense = self._dense_candidates(question, RAG_CANDIDATE_K, doc_ids, file_names,
                                       payload_filters)

        # 2. Lexical candidates (BM25 over chunk store) + doc-keyword boost.
        # The chunk store mirrors only the bulk (edinburgh) corpus, so skip the
        # lexical lane for demo users — otherwise it would leak edinburgh chunks
        # past the dense corpus filter. dense already honours the corpus filter.
        lexical: List[Dict] = []
        doc_boost: Dict[str, float] = {}
        if _current_user_corpus() != "demo":
            try:
                from .lexical_index import get_lexical_index
                li = get_lexical_index()
                lexical = li.search_chunks(raw_question, RAG_CANDIDATE_K)
                if doc_ids:
                    ids = set(doc_ids)
                    lexical = [c for c in lexical if c.get("doc_id") in ids]
                if file_names:
                    fns = set(file_names)
                    lexical = [c for c in lexical if c.get("file_name") in fns]
                for c in lexical:
                    c["key"] = f"{c.get('file_name')}::{c.get('page_number')}"
                doc_boost = li.match_docs(self._lexical_terms(raw_question))
            except Exception as e:
                logger.debug(f"   Lexical lane skipped: {e}")

        if not dense and not lexical:
            return None

        # 3. Reciprocal Rank Fusion (pure, unit-tested in lexical_index)
        from .config import RRF_K
        from .lexical_index import rrf_fuse
        fused = rrf_fuse(dense, lexical, doc_boost, rrf_k=RRF_K)

        # 4. LLM rerank the top of the fused pool → final_k
        if ENABLE_RERANK and rerank and len(fused) > final_k:
            final = self._llm_rerank(raw_question, fused[:RAG_RERANK_K], final_k)
        else:
            final = fused[:final_k]

        # 6. Build sources (dedupe by file+page, preserve order)
        sources = []
        seen = set()
        for nd in final:
            key = f"{nd.get('file_name')}_{nd.get('page_number')}"
            if key in seen:
                continue
            seen.add(key)
            sources.append(self._node_to_source(nd))
        self._report_reading(sources)

        # 5. Synthesize the answer from the final chunks — skipped in retrieve-only
        # mode (hybrid/agent/planner callers re-synthesize downstream), so we don't
        # pay for a synthesis whose text is then discarded.
        if synthesize:
            try:
                from backend.tasks.query_progress import report_step
                report_step("analysing", "analysing…")
            except Exception:
                pass
            answer = self._synthesize_from_nodes(raw_question, final)
        else:
            answer = ""

        logger.info(
            f"   Hybrid: dense={len(dense)} lexical={len(lexical)} "
            f"fused={len(fused)} → {len(sources)} sources"
        )
        return {"answer": answer, "sources": sources}

    @staticmethod
    def _build_metadata_filters(doc_ids: Optional[List[str]],
                                file_names: Optional[List[str]],
                                payload_filters: Optional[Dict[str, Any]] = None):
        """Compose a LlamaIndex MetadataFilters from (doc_ids OR file_names) AND
        scope (doc_type / project / …). The document selector stays OR (either an
        explicit doc_id or a filename-resolved file matches), but scope is AND-ed
        on top so 'only delay notices' genuinely narrows the pool. Returns None
        when nothing constrains (today's behaviour). Scope values that are empty
        are dropped so a blank scope never over-filters."""
        from llama_index.core.vector_stores.types import (
            MetadataFilter, MetadataFilters, FilterOperator, FilterCondition,
        )
        doc_file_specs = []
        if doc_ids:
            doc_file_specs.append(
                MetadataFilter(key="doc_id", value=doc_ids, operator=FilterOperator.IN))
        if file_names:
            doc_file_specs.append(
                MetadataFilter(key="file_name", value=file_names, operator=FilterOperator.IN))

        scope_specs = []
        for key, val in (payload_filters or {}).items():
            if val in (None, "", [], {}):
                continue
            if isinstance(val, (list, tuple, set)):
                scope_specs.append(MetadataFilter(key=key, value=list(val),
                                                  operator=FilterOperator.IN))
            else:
                scope_specs.append(MetadataFilter(key=key, value=val,
                                                  operator=FilterOperator.EQ))

        # Always-on per-user corpus isolation. Read here (not from payload_filters)
        # so it survives the unfiltered-retry path and applies to every caller —
        # admin2 → edinburgh vectors only, admin → demo vectors only.
        corpus = _current_user_corpus()
        if corpus:
            scope_specs.append(MetadataFilter(key="corpus", value=corpus,
                                              operator=FilterOperator.EQ))

        if not doc_file_specs and not scope_specs:
            return None
        if doc_file_specs and not scope_specs:
            # Pure document selector — preserve the original OR semantics.
            return MetadataFilters(
                filters=doc_file_specs,
                condition=FilterCondition.OR if len(doc_file_specs) > 1 else None,
            )
        if scope_specs and not doc_file_specs:
            return MetadataFilters(filters=scope_specs, condition=FilterCondition.AND)
        # Both present: (doc_id OR file_name) AND scope1 AND scope2 …
        doc_group = MetadataFilters(
            filters=doc_file_specs,
            condition=FilterCondition.OR if len(doc_file_specs) > 1 else None,
        )
        return MetadataFilters(filters=[doc_group] + scope_specs,
                               condition=FilterCondition.AND)

    def _dense_candidates(self, question: str, candidate_k: int,
                          doc_ids: Optional[List[str]],
                          file_names: Optional[List[str]],
                          payload_filters: Optional[Dict[str, Any]] = None) -> List[Dict]:
        """Retrieve dense candidates as normalized node dicts."""
        retriever_kwargs = {"similarity_top_k": candidate_k}
        flt = self._build_metadata_filters(doc_ids, file_names, payload_filters)
        if flt is not None:
            retriever_kwargs["filters"] = flt
        try:
            nodes = self.index.as_retriever(**retriever_kwargs).retrieve(question)
        except Exception as e:
            logger.warning(f"   Dense retrieval failed: {e}")
            return []
        out = []
        for n in nodes:
            meta = n.metadata or {}
            text = (n.text or "").strip()
            fname = meta.get("file_name", "Unknown")
            try:
                page = int(meta.get("page_number", 1))
            except (TypeError, ValueError):
                page = 1
            try:
                total = int(meta.get("total_pages", 1))
            except (TypeError, ValueError):
                total = 1
            out.append({
                "key": f"{fname}::{page}",
                "doc_id": meta.get("doc_id", ""),
                "file_name": fname,
                "file_path": meta.get("file_path", ""),
                "page_number": page,
                "total_pages": total,
                "text": text,
                "dense_score": round(n.score, 3) if n.score else None,
            })
        return out

    def _llm_rerank(self, question: str, candidates: List[Dict], final_k: int) -> List[Dict]:
        """Rerank candidates by relevance with a single cached LLM call."""
        if not candidates:
            return []
        from . import llm_client
        import hashlib
        passages = []
        for i, nd in enumerate(candidates):
            snippet = (nd.get("text") or "")[:500].replace("\n", " ")
            passages.append(f"[{i}] ({nd.get('file_name')} p.{nd.get('page_number')}) {snippet}")
        prompt = (
            "Rank the passages by how well they help answer the QUESTION.\n\n"
            f"QUESTION: {question}\n\n"
            "PASSAGES:\n" + "\n\n".join(passages) + "\n\n"
            f"Return JSON {{\"order\": [passage numbers, most relevant first]}} "
            f"with the {final_k} most relevant passage numbers only."
        )
        try:
            key = "rerank:" + hashlib.sha256(
                (question + "|" + "|".join(str(c.get("key")) for c in candidates)).encode()
            ).hexdigest()[:32]
            resp = llm_client.generate_json(
                prompt,
                system="You are a precise passage reranker. Output JSON only.",
                cache_key=key,
                model=GEMINI_MODEL_LITE if ENABLE_LITE_TIER else "",
            )
            order = resp.raw.get("order", []) if isinstance(resp.raw, dict) else []
            picked, used = [], set()
            for i in order:
                if isinstance(i, int) and 0 <= i < len(candidates) and i not in used:
                    picked.append(candidates[i])
                    used.add(i)
                if len(picked) >= final_k:
                    break
            # Backfill if the model returned too few
            for i, c in enumerate(candidates):
                if len(picked) >= final_k:
                    break
                if i not in used:
                    picked.append(c)
            return picked[:final_k]
        except Exception as e:
            logger.warning(f"   LLM rerank failed → fused order: {e}")
            return candidates[:final_k]

    def _synthesize_from_nodes(self, question: str, nodes: List[Dict]) -> str:
        """Synthesize an answer from final chunks using the document system prompt."""
        from . import llm_client
        if not nodes:
            return "No relevant document excerpts were found for this question."
        parts = []
        for i, nd in enumerate(nodes, 1):
            parts.append(
                f"[Source {i}: {nd.get('file_name')}, p.{nd.get('page_number')}]\n"
                f"{nd.get('text', '')}"
            )
        context = "\n\n---\n\n".join(parts)
        prompt = (
            f"DOCUMENT EXCERPTS:\n{context}\n\n"
            f"QUESTION: {question}\n\nANSWER:"
        )
        resp = llm_client.generate_text(
            prompt, system=self.DOCUMENT_SYSTEM_PROMPT, max_tokens=1024,
        )
        # Telemetry
        try:
            from .telemetry import get_current_trace
            tr = get_current_trace()
            if tr:
                tr.record_llm_call(resp.usage)
        except Exception:
            pass
        return resp.text.strip()

    def _node_to_source(self, nd: Dict) -> Dict:
        """Convert a fused node dict into the standard source dict shape."""
        text = (nd.get("text") or "").strip()
        sentences = text.replace("\n", " ").split(". ")
        highlight = ". ".join(sentences[:3])
        if len(sentences) > 3:
            highlight += "..."
        doc_id = nd.get("doc_id") or (
            generate_doc_id(nd.get("file_path")) if nd.get("file_path") else ""
        )
        return {
            "doc_id": doc_id,
            "file_name": nd.get("file_name", "Unknown"),
            "file_path": nd.get("file_path", ""),
            "page_number": nd.get("page_number", 1),
            "total_pages": nd.get("total_pages", 1),
            "score": nd.get("dense_score"),
            "text_snippet": text[:500] + "..." if len(text) > 500 else text,
            "highlight_text": highlight[:300],
        }

    def _lexical_terms(self, question: str) -> List[str]:
        """Build keyword terms for the document-level keyword signal."""
        from .lexical_index import _tokenize
        terms = [t for t in _tokenize(question) if len(t) >= 3]
        try:
            from .jargon_manager import get_jargon_manager
            terms += get_jargon_manager().expand_domain_concepts(question) or []
        except Exception:
            pass
        seen, out = set(), []
        for t in terms:
            tl = t.lower()
            if tl not in seen:
                seen.add(tl)
                out.append(t)
        return out[:12]

    def _extract_sources(self, response) -> list:
        """Extract sources from a LlamaIndex query response."""
        sources = []
        seen = set()

        for i, node in enumerate(response.source_nodes, 1):
            meta = node.metadata
            file_name = meta.get("file_name", "Unknown")
            file_path = meta.get("file_path", "")
            score = round(node.score, 3) if node.score else None

            try:
                page_num = int(meta.get("page_number", 1))
            except (ValueError, TypeError):
                page_num = 1

            try:
                total_pages = int(meta.get("total_pages", 1))
            except (ValueError, TypeError):
                total_pages = 1

            key = f"{file_name}_{page_num}"
            if key in seen:
                continue
            seen.add(key)

            text = node.text.strip()
            sentences = text.replace('\n', ' ').split('. ')
            highlight = '. '.join(sentences[:3])
            if len(sentences) > 3:
                highlight += '...'

            sources.append({
                "doc_id": meta.get("doc_id", generate_doc_id(file_path)) if file_path else "",
                "file_name": file_name,
                "file_path": file_path,
                "page_number": page_num,
                "total_pages": total_pages,
                "score": score,
                "text_snippet": text[:500] + "..." if len(text) > 500 else text,
                "highlight_text": highlight[:300],
            })

        return sources

    def query_with_provider(self, question: str, provider: str, top_k: int = 10,
                            doc_ids: Optional[List[str]] = None,
                            file_names: Optional[List[str]] = None) -> dict:
        """Query documents using a specific LLM provider for answer synthesis."""
        from .llm_client import create_llm

        logger.info(f"[DocumentRAG] Query with provider={provider}: {question[:80]}...")

        try:
            from .jargon_manager import get_jargon_manager
            semantic_query = get_jargon_manager().replace_query_terms_with_meanings(question)
            if semantic_query and semantic_query != question:
                logger.info(f"   Jargon semantic query: {semantic_query[:100]}")
                question = semantic_query
        except Exception:
            pass

        if not self.index:
            if not self.load_index():
                return {
                    "answer": "No documents indexed. Please upload documents first.",
                    "sources": [],
                }

        llm, model_name = create_llm(provider)
        logger.info(f"   Using {provider}/{model_name} for synthesis...")

        # For Claude: LlamaIndex query_engine doesn't support custom wrappers,
        # so we retrieve chunks via default engine and synthesize with Claude directly.
        if provider == "claude":
            return self._query_with_direct_synthesis(question, llm, top_k, doc_ids, file_names)

        kwargs = {"similarity_top_k": top_k, "llm": llm}
        filter_specs = []
        if doc_ids:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, MetadataFilter, FilterOperator,
            )
            filter_specs.append(
                MetadataFilter(key="doc_id", value=doc_ids, operator=FilterOperator.IN)
            )
        if file_names:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, MetadataFilter, FilterOperator,
            )
            filter_specs.append(
                MetadataFilter(key="file_name", value=file_names, operator=FilterOperator.IN)
            )
        if filter_specs:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, FilterCondition,
            )
            kwargs["filters"] = MetadataFilters(
                filters=filter_specs,
                condition=FilterCondition.OR if len(filter_specs) > 1 else None,
            )
        query_engine = self.index.as_query_engine(**kwargs)
        response = query_engine.query(question)
        sources = self._extract_sources(response)

        logger.info(f"   [{provider}] Found {len(sources)} sources")
        return {"answer": str(response), "sources": sources}

    def _query_with_direct_synthesis(self, question: str, llm, top_k: int = 10,
                                     doc_ids: Optional[List[str]] = None,
                                     file_names: Optional[List[str]] = None) -> dict:
        """Retrieve chunks via Pinecone, then synthesize answer with a non-LlamaIndex LLM."""
        from llama_index.core import VectorStoreIndex

        # Step 1: Retrieve relevant chunks
        retriever_kwargs = {"similarity_top_k": top_k}
        filter_specs = []
        if doc_ids:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, MetadataFilter, FilterOperator,
            )
            filter_specs.append(
                MetadataFilter(key="doc_id", value=doc_ids, operator=FilterOperator.IN)
            )
        if file_names:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, MetadataFilter, FilterOperator,
            )
            filter_specs.append(
                MetadataFilter(key="file_name", value=file_names, operator=FilterOperator.IN)
            )
        if filter_specs:
            from llama_index.core.vector_stores.types import (
                MetadataFilters, FilterCondition,
            )
            retriever_kwargs["filters"] = MetadataFilters(
                filters=filter_specs,
                condition=FilterCondition.OR if len(filter_specs) > 1 else None,
            )
        retriever = self.index.as_retriever(**retriever_kwargs)
        nodes = retriever.retrieve(question)

        # Step 2: Build context from retrieved chunks
        context_parts = []
        for i, node in enumerate(nodes, 1):
            meta = node.metadata
            fname = meta.get("file_name", "Unknown")
            page = meta.get("page_number", "?")
            context_parts.append(f"[Source {i}: {fname}, p.{page}]\n{node.text}")
        context = "\n\n---\n\n".join(context_parts)

        # Step 3: Synthesize with LLM
        prompt = (
            "Based on the following document excerpts, answer the user's question.\n"
            "Only use information from the provided sources.\n"
            "Cite the source document name for every claim, and especially for every number.\n"
            "Only state a numeric figure that appears verbatim in a source; name that source. "
            "If you must add up or infer a total across sources, say 'approximately', list which "
            "sources you combined, and mark it as an estimate — never present an inferred number as exact.\n"
            "If the information is not available, say so.\n\n"
            f"SOURCES:\n{context}\n\n"
            f"QUESTION: {question}\n\n"
            "ANSWER:"
        )
        response = llm.complete(prompt)
        answer = response.text

        # Step 4: Extract sources metadata
        sources = []
        seen = set()
        for node in nodes:
            meta = node.metadata
            file_name = meta.get("file_name", "Unknown")
            page_num = int(meta.get("page_number", 1))
            key = f"{file_name}_{page_num}"
            if key in seen:
                continue
            seen.add(key)
            text = node.text.strip()
            sentences = text.replace('\n', ' ').split('. ')
            highlight = '. '.join(sentences[:3])
            if len(sentences) > 3:
                highlight += '...'
            sources.append({
                "doc_id": meta.get("doc_id", ""),
                "file_name": file_name,
                "file_path": meta.get("file_path", ""),
                "page_number": page_num,
                "total_pages": int(meta.get("total_pages", 1)),
                "score": round(node.score, 3) if node.score else None,
                "text_snippet": text[:500],
                "highlight_text": highlight[:300],
            })

        logger.info(f"   [direct] Found {len(sources)} sources, answer len={len(answer)}")
        return {"answer": answer, "sources": sources}

    # ── Named-doc retrieval (filename-resolved queries) ──────────────────
    #
    # When the user names a specific document by file name, we bypass the
    # LlamaIndex query engine entirely and fetch chunks directly via the
    # Pinecone API with a metadata filter. This avoids two failure modes
    # observed in the LlamaIndex path: (a) doc_id IN filter never matches
    # because Pinecone's doc_id is a per-chunk UUID, not a doc-level handle;
    # (b) MetadataFilter on file_name occasionally returns empty source_nodes
    # even when the chunks exist. Direct fetch + manual synthesis is more
    # deterministic and lets us combine named-doc chunks with semantic
    # neighbours in a single context.

    @staticmethod
    def _chunk_text_from_payload(md: Dict[str, Any]) -> str:
        """Extract chunk text from a vector-store payload/metadata dict.
        LlamaIndex stores the node as JSON under _node_content."""
        node_content = md.get("_node_content")
        if node_content:
            try:
                import json as _json
                return _json.loads(node_content).get("text", "")
            except Exception:
                return node_content if isinstance(node_content, str) else ""
        return md.get("text", "") or ""

    def update_payload_scope(self, file_name: str, scope: Dict[str, Any]) -> bool:
        """Stamp scoped-metadata keys (doc_type / project / date / …) onto every
        vector of a document, in-place, without re-embedding. This is what makes
        the `payload_filters` lane of retrieval usable — at ingest we know a
        document's doc_type/date, so we write them to the payload here.

        Filters by file_name, which is the reliable, INDEXED payload key (the
        per-document doc_id is not unique across chunks and is unindexed, so a
        doc_id filter would full-scan and time out on a large collection).
        Best-effort and idempotent (set_payload merges). Empty values dropped.
        """
        clean = {k: v for k, v in (scope or {}).items() if v not in (None, "", [], {})}
        if not file_name or not clean:
            return False
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                flt = qmodels.Filter(must=[qmodels.FieldCondition(
                    key="file_name", match=qmodels.MatchValue(value=file_name))])
                self.qdrant_client.set_payload(
                    collection_name=QDRANT_COLLECTION,
                    payload=clean,
                    points=qmodels.FilterSelector(filter=flt),
                )
            else:
                # Pinecone has no filter-based metadata update; fetch the doc's
                # chunk ids by file_name then update each.
                ids = self._pinecone_ids_for_file(file_name)
                for cid in ids:
                    self.pinecone_index.update(id=cid, set_metadata=clean,
                                               namespace="__default__")
            logger.info(f"[ScopePayload] {file_name[:40]} += {list(clean.keys())}")
            return True
        except Exception as e:
            logger.warning(f"[ScopePayload] update skipped for {file_name[:40]}: {e}")
            return False

    def _pinecone_ids_for_file(self, file_name: str) -> List[str]:
        """Best-effort list of Pinecone vector ids for a document (metadata-only
        query). Used by update_payload_scope on the Pinecone backend."""
        try:
            res = self.pinecone_index.query(
                vector=[0.0] * EMBEDDING_DIMENSION, top_k=10000,
                include_metadata=False, filter={"file_name": {"$eq": file_name}},
                namespace="__default__",
            )
            matches = res.get("matches") if isinstance(res, dict) else getattr(res, "matches", [])
            out = []
            for m in matches or []:
                mid = m.get("id") if isinstance(m, dict) else getattr(m, "id", None)
                if mid:
                    out.append(mid)
            return out
        except Exception:
            return []

    def _vector_query(
        self, qvec: List[float], top_k: int, file_names: Optional[List[str]] = None,
        payload_filters: Optional[Dict[str, Any]] = None,
    ) -> List[Dict[str, Any]]:
        """Backend-agnostic vector search returning [{text, metadata, score}].

        Optional scoping:
          - file_names: restrict to those documents.
          - payload_filters: {key: value | [values]} on the scoped-metadata payload
            (e.g. {"doc_type": "Witness Statement", "project": "Edinburgh Tram"}).
            Multiple keys are AND-ed; a list value matches ANY of its entries.

        Works with both Pinecone (direct .query) and Qdrant (query_points + filter).
        """
        chunks: List[Dict[str, Any]] = []
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                must = []
                if file_names:
                    must.append(qmodels.FieldCondition(
                        key="file_name", match=qmodels.MatchAny(any=list(file_names))))
                for key, val in (payload_filters or {}).items():
                    if isinstance(val, (list, tuple, set)):
                        must.append(qmodels.FieldCondition(
                            key=key, match=qmodels.MatchAny(any=list(val))))
                    else:
                        must.append(qmodels.FieldCondition(
                            key=key, match=qmodels.MatchValue(value=val)))
                _corpus = _current_user_corpus()  # per-user isolation (named-doc path)
                if _corpus:
                    must.append(qmodels.FieldCondition(
                        key="corpus", match=qmodels.MatchValue(value=_corpus)))
                flt = qmodels.Filter(must=must) if must else None
                res = self.qdrant_client.query_points(
                    collection_name=QDRANT_COLLECTION,
                    query=qvec,
                    limit=top_k,
                    query_filter=flt,
                    with_payload=True,
                )
                for p in res.points:
                    md = p.payload or {}
                    chunks.append({"text": self._chunk_text_from_payload(md),
                                   "metadata": md, "score": p.score})
            else:
                pc_filter: Dict[str, Any] = {}
                if file_names:
                    pc_filter["file_name"] = {"$in": list(file_names)}
                for key, val in (payload_filters or {}).items():
                    pc_filter[key] = {"$in": list(val)} if isinstance(val, (list, tuple, set)) else {"$eq": val}
                kwargs = dict(vector=qvec, top_k=top_k, include_metadata=True,
                              namespace="__default__")
                if pc_filter:
                    kwargs["filter"] = pc_filter
                res = self.pinecone_index.query(**kwargs)
                matches = res.get("matches") if isinstance(res, dict) else getattr(res, "matches", [])
                for m in matches or []:
                    md = m.get("metadata") if isinstance(m, dict) else getattr(m, "metadata", {}) or {}
                    score = m.get("score") if isinstance(m, dict) else getattr(m, "score", None)
                    chunks.append({"text": self._chunk_text_from_payload(md),
                                   "metadata": md, "score": score})
        except Exception as e:
            logger.warning(f"[VectorQuery] {self.backend} query failed: {e}")
        return chunks

    def fetch_doc_vectors(self, file_name: str, max_chunks: int = 50) -> List[List[float]]:
        """Return raw stored chunk vectors for a document (for centroid pooling).
        Backend-agnostic: Qdrant scroll(with_vectors) or Pinecone zero-vector query."""
        out: List[List[float]] = []
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                flt = qmodels.Filter(must=[qmodels.FieldCondition(
                    key="file_name", match=qmodels.MatchValue(value=file_name))])
                points, _ = self.qdrant_client.scroll(
                    collection_name=QDRANT_COLLECTION,
                    scroll_filter=flt, limit=max_chunks,
                    with_vectors=True, with_payload=False,
                )
                for p in points:
                    vec = p.vector
                    if isinstance(vec, dict):  # named-vector collections
                        vec = next(iter(vec.values()), None)
                    if vec:
                        out.append(list(vec))
            else:
                zero_vec = [0.0] * EMBEDDING_DIMENSION
                res = self.pinecone_index.query(
                    vector=zero_vec, top_k=max_chunks, include_values=True,
                    include_metadata=False, filter={"file_name": {"$eq": file_name}},
                    namespace="",
                )
                matches = res.get("matches") if isinstance(res, dict) else getattr(res, "matches", [])
                for m in matches or []:
                    values = m.get("values") if isinstance(m, dict) else getattr(m, "values", None)
                    if values:
                        out.append(list(values))
        except Exception as e:
            logger.warning(f"[VectorFetch] {self.backend} fetch failed for {file_name!r}: {e}")
        return out

    def _fetch_named_doc_chunks(
        self, question: str, file_names: List[str], top_k: int = 10,
    ) -> List[Dict[str, Any]]:
        """Fetch top-k chunks for the given file_names, ranked by query similarity.
        Returns a list of {text, metadata, score} dicts ordered by descending similarity."""
        if not file_names:
            return []
        try:
            from llama_index.core import Settings
            qvec = Settings.embed_model.get_query_embedding(question)
        except Exception as e:
            logger.warning(f"[NamedDoc] embed failed: {e}")
            return []
        return self._vector_query(qvec, top_k, file_names=file_names)

    def query_named_docs_with_provider(
        self, question: str, provider: str, file_names: List[str],
        top_k_named: int = 10, top_k_semantic: int = 5,
    ) -> dict:
        """Generate an answer for a named-document query.

        Pipeline:
          1) direct Pinecone fetch of top-N chunks within file_names (named context)
          2) unfiltered semantic top-M chunks (for context that the named doc
             may not contain on its own — e.g. cross-references in mail threads)
          3) merge, dedup by file+page, build LLM prompt
          4) synthesize via the requested provider
          5) return {answer, sources}

        Sources order: named-doc chunks first (PDF cited as the user expects),
        semantic chunks after.
        """
        from .llm_client import create_llm

        logger.info(
            f"[NamedDoc] provider={provider} files={file_names} "
            f"top_k_named={top_k_named} top_k_semantic={top_k_semantic}"
        )

        # Jargon normalization (mirror query_with_provider behaviour)
        try:
            from .jargon_manager import get_jargon_manager
            semantic_query = get_jargon_manager().replace_query_terms_with_meanings(question)
            if semantic_query and semantic_query != question:
                logger.info(f"[NamedDoc] jargon semantic query: {semantic_query[:100]}")
                question = semantic_query
        except Exception:
            pass

        named_chunks = self._fetch_named_doc_chunks(
            question, file_names, top_k=top_k_named,
        )

        # Semantic neighbours (no scope) — keep small to leave room for named.
        semantic_chunks: List[Dict[str, Any]] = []
        if top_k_semantic > 0:
            try:
                from llama_index.core import Settings
                qvec = Settings.embed_model.get_query_embedding(question)
                semantic_chunks = self._vector_query(qvec, top_k_semantic)
            except Exception as e:
                logger.warning(f"[NamedDoc] semantic fetch failed: {e}")

        # Merge: named first, then semantic that aren't already in named (dedup
        # by file+page so we don't duplicate the same page from both lanes).
        seen_keys: set = set()
        merged: List[Dict[str, Any]] = []
        for src in (named_chunks + semantic_chunks):
            md = src.get("metadata") or {}
            key = (md.get("file_name", ""), md.get("page_number", -1))
            if key in seen_keys:
                continue
            seen_keys.add(key)
            merged.append(src)

        if not merged:
            return {
                "answer": "No relevant content found for the named document.",
                "sources": [],
            }

        # Build LLM prompt
        context_parts = []
        for i, c in enumerate(merged, 1):
            md = c.get("metadata") or {}
            fname = md.get("file_name", "Unknown")
            page = md.get("page_number", "?")
            text = (c.get("text") or "")[:1500]
            context_parts.append(f"[Source {i}: {fname}, p.{page}]\n{text}")
        context = "\n\n---\n\n".join(context_parts)
        prompt = (
            "Based on the following document excerpts, answer the user's question. "
            "Use information from the provided sources. Cite the relevant filename "
            "and page number naturally where appropriate.\n\n"
            f"SOURCES:\n{context}\n\n"
            f"QUESTION: {question}\n\n"
            "ANSWER:"
        )

        try:
            llm, _ = create_llm(provider)
            response = llm.complete(prompt)
            answer = getattr(response, "text", str(response))
        except Exception as e:
            logger.error(f"[NamedDoc] {provider} synthesis failed: {e}")
            answer = ""

        # Build sources list (named first, then semantic)
        sources: List[Dict[str, Any]] = []
        for c in merged:
            md = c.get("metadata") or {}
            text = (c.get("text") or "").strip()
            sentences = text.replace("\n", " ").split(". ")
            highlight = ". ".join(sentences[:3])[:300]
            try:
                page = int(md.get("page_number", 1))
            except (TypeError, ValueError):
                page = 1
            try:
                total = int(md.get("total_pages", 1))
            except (TypeError, ValueError):
                total = 1
            sources.append({
                "doc_id": md.get("doc_id", ""),
                "file_name": md.get("file_name", ""),
                "file_path": md.get("file_path", ""),
                "page_number": page,
                "total_pages": total,
                "score": round(c.get("score") or 0.0, 3),
                "text_snippet": text[:500],
                "highlight_text": highlight,
            })
        return {"answer": answer, "sources": sources}

    def query_named_docs_dual(
        self, question: str, file_names: List[str],
        top_k_named: int = 10, top_k_semantic: int = 5,
    ) -> dict:
        """Multi-provider variant of query_named_docs_with_provider."""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from .config import LLM_PROVIDERS

        results: Dict[str, Any] = {}

        def _q(prov: str):
            return prov, self.query_named_docs_with_provider(
                question, prov, file_names,
                top_k_named=top_k_named, top_k_semantic=top_k_semantic,
            )

        with ThreadPoolExecutor(max_workers=len(LLM_PROVIDERS)) as ex:
            futs = {ex.submit(_q, p): p for p in LLM_PROVIDERS}
            for fut in as_completed(futs):
                try:
                    prov, res = fut.result()
                    results[prov] = res
                except Exception as e:
                    prov = futs[fut]
                    logger.error(f"[NamedDoc] {prov} thread failed: {e}")
                    results[prov] = {
                        "answer": "Provider failed.", "sources": [],
                        "_error": str(e),
                    }
        return results

    def query_dual(self, question: str, top_k: int = 10,
                   doc_ids: Optional[List[str]] = None,
                   file_names: Optional[List[str]] = None) -> dict:
        """Query with both OpenAI and Claude in parallel."""
        from concurrent.futures import ThreadPoolExecutor, as_completed
        from .config import LLM_PROVIDERS

        results = {}

        def _query_provider(prov):
            return prov, self.query_with_provider(
                question, prov, top_k, doc_ids=doc_ids, file_names=file_names
            )

        with ThreadPoolExecutor(max_workers=len(LLM_PROVIDERS)) as executor:
            futures = {executor.submit(_query_provider, p): p for p in LLM_PROVIDERS}
            for future in as_completed(futures):
                try:
                    prov, result = future.result()
                    results[prov] = result
                except Exception as e:
                    prov = futures[future]
                    logger.error(f"   [{prov}] Query failed: {e}")
                    # Attempt fallback to gemini for failed providers
                    if prov != "gemini" and "gemini" in LLM_PROVIDERS:
                        try:
                            _, fallback_result = _query_provider("gemini")
                            fallback_result["_fallback_from"] = prov
                            results[prov] = fallback_result
                            logger.info(f"   [{prov}] Fallback to gemini succeeded")
                        except Exception:
                            results[prov] = {
                                "answer": "This model is temporarily unavailable. Please try again.",
                                "sources": [],
                            }
                    else:
                        results[prov] = {
                            "answer": "This model is temporarily unavailable. Please try again.",
                            "sources": [],
                        }

        return results

    def get_page_content(self, file_path: str, page_num: int) -> Optional[str]:
        """Get full content of a specific page for preview."""
        try:
            doc = fitz.open(file_path)
            if 1 <= page_num <= len(doc):
                text = doc[page_num - 1].get_text()
                doc.close()
                return text
            doc.close()
        except Exception as e:
            logger.error(f"Error reading page: {e}")
        return None

    def clear_index(self):
        """Clear all vectors from the index."""
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                logger.info("[Qdrant] Clearing entire collection...")
                # Recreate is the simplest way to truly empty a collection.
                self.qdrant_client.delete_collection(QDRANT_COLLECTION)
                self.qdrant_client.create_collection(
                    collection_name=QDRANT_COLLECTION,
                    vectors_config=qmodels.VectorParams(
                        size=EMBEDDING_DIMENSION,
                        distance=qmodels.Distance.COSINE,
                    ),
                )
            else:
                log_pinecone("Clearing entire index...")
                self.pinecone_index.delete(delete_all=True)
            self.documents = []
            self.file_registry = {}
            self.index = None
            logger.info("✅ Index cleared")
        except Exception as e:
            logger.error(f"Error clearing index: {e}")

    def clear_file(self, file_name: str):
        """Clear vectors for a specific file."""
        doc_id = generate_doc_id(file_name)
        try:
            if self.backend == "qdrant":
                from qdrant_client.http import models as qmodels
                flt = qmodels.Filter(must=[
                    qmodels.FieldCondition(
                        key="doc_id",
                        match=qmodels.MatchValue(value=doc_id),
                    )
                ])
                self.qdrant_client.delete(
                    collection_name=QDRANT_COLLECTION,
                    points_selector=qmodels.FilterSelector(filter=flt),
                )
            else:
                self.pinecone_index.delete(filter={"doc_id": {"$eq": doc_id}})
            if file_name in self.file_registry:
                del self.file_registry[file_name]
            self.documents = [d for d in self.documents if d.metadata.get("file_name") != file_name]
            # Keep the lexical chunk store in sync with the vector store.
            try:
                from .chunk_store import get_chunk_store
                get_chunk_store().delete_by_file(file_name)
            except Exception as e:
                logger.debug(f"Chunk store delete skipped: {e}")
            logger.info(f"Cleared: {file_name}")
        except Exception as e:
            logger.error(f"Error clearing file: {e}")

    def get_index_stats(self) -> Dict[str, Any]:
        """Get index statistics."""
        if self.backend == "qdrant":
            try:
                info = self.qdrant_client.get_collection(QDRANT_COLLECTION)
                total = info.points_count or 0
            except Exception:
                total = 0
        else:
            stats = self.pinecone_index.describe_index_stats()
            total = stats.get("total_vector_count", 0)
        return {
            "total_vectors": total,
            "files_indexed": len(self.file_registry),
            "files": list(self.file_registry.keys()),
            "backend": self.backend,
        }


# Singleton
_document_rag: Optional[DocumentRAG] = None


def get_document_rag() -> DocumentRAG:
    """Get or create DocumentRAG singleton."""
    global _document_rag
    if _document_rag is None:
        _document_rag = DocumentRAG()
    return _document_rag
