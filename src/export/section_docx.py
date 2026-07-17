"""Report section → DOCX, built from the structured input.

Unlike the PDF path this does not reuse compose_section: python-docx has no
HTML importer, and parsing back the HTML we just generated would be a lossy
round trip through a format neither side wants.

The cost is a second renderer that must agree with the templater. The caps
below (200 rows / 30 sources / 15 caveats) are copied from html_section, and a
parity test asserts they stay equal — that is the drift this file risks.
"""

from __future__ import annotations

import io
import re
from typing import Any, Dict, List, Optional

# Mirrors html_section.compose_section's slices. Keep in sync (test enforces).
MAX_ROWS = 200
MAX_SOURCES = 30
MAX_CAVEATS = 15

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_PARA_NO_RE = re.compile(r"^(\d+\.\d+\.\d+)\s+")


def _add_markdown_paragraph(doc, text: str) -> None:
    """Minimal markdown → docx runs: **bold** and a leading paragraph number.
    Same subset html_section._md_paragraphs supports; anything else is text."""
    para = doc.add_paragraph()
    m = _PARA_NO_RE.match(text)
    if m:
        run = para.add_run(m.group(1) + " ")
        run.bold = True
        text = text[m.end():]
    pos = 0
    for match in _BOLD_RE.finditer(text):
        if match.start() > pos:
            para.add_run(text[pos:match.start()])
        para.add_run(match.group(1)).bold = True
        pos = match.end()
    if pos < len(text):
        para.add_run(text[pos:])


def render_section_docx(title: str, narrative_markdown: str,
                        tables: Optional[List[Dict[str, Any]]] = None,
                        sources: Optional[List[Dict[str, Any]]] = None,
                        caveats: Optional[List[str]] = None,
                        validation: Optional[Dict[str, Any]] = None,
                        section_number: str = "") -> bytes:
    """Render a report section to DOCX bytes. Signature mirrors compose_section."""
    from docx import Document

    doc = Document()
    heading = f"{section_number + '. ' if section_number else ''}{title}"
    doc.add_heading(heading, level=1)

    for para in re.split(r"\n\s*\n", (narrative_markdown or "").strip()):
        if para.strip():
            _add_markdown_paragraph(doc, re.sub(r"^(#+\s*)", "", para.strip()))

    for t in tables or []:
        if t.get("title"):
            doc.add_heading(str(t["title"]), level=3)
        cols = [str(c) for c in (t.get("columns") or [])]
        rows = (t.get("rows") or [])[:MAX_ROWS]
        if not cols:
            continue
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        for i, c in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = c
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row[:len(cols)]):
                cells[i].text = "" if val is None else str(val)

    if sources:
        doc.add_heading("Sources", level=3)
        for s in sources[:MAX_SOURCES]:
            doc.add_paragraph(
                f"{s.get('file_name', '?')}, p.{s.get('page_number', 1)}",
                style="List Bullet")

    if caveats:
        doc.add_heading("Caveats", level=3)
        for c in caveats[:MAX_CAVEATS]:
            doc.add_paragraph(str(c), style="List Bullet")

    if validation:
        bits = []
        for name, v in (validation or {}).items():
            state = (v.get("post") or v.get("status") or "?") \
                if isinstance(v, dict) else str(v)
            bits.append(f"{name}: {state}")
        doc.add_paragraph("Validation — " + " · ".join(bits))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
