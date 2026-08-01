"""Deterministic Word renderers shared by AI Chronology and Forensic reports."""

from __future__ import annotations

import io
from typing import Dict, Iterable, List, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from .evidence_model import ChronologyEntry, EvidenceItem, ReportAudit, VerifiedClaim, evidence_map
from .word_footnotes import FootnoteRegistry, attach_footnote_part


def _base_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08
    return doc


def _heading(doc: Document, text: str, *, size: int = 12, bold: bool = True,
             centered: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    return p


def _claim_paragraph(
    doc: Document,
    prefix: str,
    claims: Iterable[VerifiedClaim],
    sources: Dict[str, EvidenceItem],
    exhibits: Dict[str, int],
    footnotes: FootnoteRegistry,
    unresolved: List[str],
):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.add_run(prefix).bold = True
    for index, claim in enumerate(claims):
        if not claim.supported or not claim.text.strip():
            continue
        if index or prefix:
            p.add_run(" ")
        p.add_run(claim.text.strip())
        valid = []
        for source_id in claim.source_ids:
            source = sources.get(source_id)
            if source is None:
                unresolved.append(source_id)
                continue
            valid.append(source)
        for source in valid:
            exhibit = exhibits.setdefault(source.source_id, len(exhibits) + 1)
            note = source.footnote_text(
                exhibit,
                claim.inference_basis if claim.is_inference else "",
            )
            footnotes.add(p, note)
    return p


def build_ai_chronology_docx(
    *,
    project_name: str,
    issue_number: int,
    title: str,
    entries: List[ChronologyEntry],
    evidence: List[EvidenceItem],
    audit_metadata: Dict[str, str] | None = None,
) -> Tuple[bytes, ReportAudit]:
    """Render the sample chronology contract with real claim-level footnotes."""
    doc = _base_document()
    _heading(doc, project_name.upper(), size=12, centered=True)
    _heading(doc, "Delay and Prolongation – Chronology of Events", size=12, centered=True)
    _heading(doc, f"{issue_number}. {title}", size=11, centered=True)

    footnotes = FootnoteRegistry()
    sources = evidence_map(evidence)
    exhibits: Dict[str, int] = {}
    unresolved: List[str] = []

    ordered = sorted(entries, key=lambda e: (e.event_date or "9999-99-99", e.entry_ref))
    for position, entry in enumerate(ordered, 1):
        # Numbering is a rendering concern.  Model-provided references are
        # deliberately ignored so an otherwise valid draft cannot escape the
        # sample documents' 6.<issue>.<entry> contract.
        ref = f"6.{issue_number}.{position}"
        date = entry.event_date or "Date not established"
        if entry.date_precision != "exact" and not date.endswith("*"):
            date += "*"
        prefix = f"{ref}  {date} –"
        claims = [c for c in entry.claims if c.supported]
        if not claims:
            claims = [VerifiedClaim(
                text="Record not established from the documents reviewed.",
                source_ids=[], supported=True, confidence="low",
            )]
        _claim_paragraph(doc, prefix, claims, sources, exhibits, footnotes, unresolved)

        # conflicting_positions is retained as structured review metadata.  It
        # is not rendered as free text because every factual sentence in the
        # issued document must be represented by a sourced VerifiedClaim.

    if audit_metadata:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run("AI draft audit · " + " · ".join(
            f"{k}={v}" for k, v in audit_metadata.items() if v
        ))
        run.font.size = Pt(7)
        run.italic = True

    raw = io.BytesIO()
    doc.save(raw)
    blob = attach_footnote_part(raw.getvalue(), footnotes.texts)
    audit = ReportAudit(
        footnote_references=len(footnotes.texts),
        footnote_records=len(footnotes.texts),
        unique_source_ids=len(exhibits),
        unresolved_source_ids=sorted(set(unresolved)),
    )
    return blob, audit


FORENSIC_SECTIONS = (
    "Instructions and scope",
    "Documents reviewed",
    "Methodology",
    "Factual chronology",
    "Issue-by-issue analysis",
    "Parties’ positions",
    "Contradictions and missing records",
    "Findings",
    "Assumptions and limitations",
)


def build_forensic_report_docx(
    *,
    project_name: str,
    title: str,
    sections: Dict[str, List[VerifiedClaim]],
    evidence: List[EvidenceItem],
    status: str = "Draft",
    audit_metadata: Dict[str, str] | None = None,
) -> Tuple[bytes, ReportAudit]:
    doc = _base_document()
    _heading(doc, project_name.upper(), size=12, centered=True)
    _heading(doc, "Forensic Report", size=14, centered=True)
    _heading(doc, title, size=11, centered=True)
    _heading(doc, status.upper(), size=9, centered=True)

    footnotes = FootnoteRegistry()
    sources = evidence_map(evidence)
    exhibits: Dict[str, int] = {}
    unresolved: List[str] = []
    for section_number, name in enumerate(FORENSIC_SECTIONS, 1):
        _heading(doc, f"{section_number}. {name}", size=11)
        claims = sections.get(name) or []
        if not claims:
            doc.add_paragraph("Record not established from the documents reviewed.")
            continue
        for claim_number, claim in enumerate(claims, 1):
            _claim_paragraph(
                doc, f"{section_number}.{claim_number}", [claim], sources,
                exhibits, footnotes, unresolved,
            )
            if claim.counter_source_ids:
                counter = VerifiedClaim(
                    text="Countervailing evidence was identified.",
                    source_ids=claim.counter_source_ids,
                    confidence=claim.confidence,
                )
                _claim_paragraph(doc, "Counter-evidence:", [counter], sources,
                                 exhibits, footnotes, unresolved)
            if claim.missing_records:
                doc.add_paragraph("Missing records: " + "; ".join(claim.missing_records))

    if audit_metadata:
        doc.add_paragraph("AI draft audit · " + " · ".join(
            f"{k}={v}" for k, v in audit_metadata.items() if v
        ))
    raw = io.BytesIO(); doc.save(raw)
    blob = attach_footnote_part(raw.getvalue(), footnotes.texts)
    audit = ReportAudit(
        footnote_references=len(footnotes.texts),
        footnote_records=len(footnotes.texts),
        unique_source_ids=len(exhibits),
        unresolved_source_ids=sorted(set(unresolved)),
    )
    return blob, audit


__all__ = ["FORENSIC_SECTIONS", "build_ai_chronology_docx", "build_forensic_report_docx"]
